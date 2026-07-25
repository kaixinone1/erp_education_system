"""深度分析Word模板结构 - 逐格逐段落"""
from docx import Document
from docx.oxml.ns import qn
import re

template_path = r"D:\erp_thirteen\数据库信息\模板\职工退休呈报表.docx"
doc = Document(template_path)

print("=" * 80)
print("SECTION 结构")
print("=" * 80)
for i, s in enumerate(doc.sections):
    st = {0:"CONTINUOUS", 1:"NEW_COLUMN", 2:"NEW_PAGE", 3:"EVEN_PAGE", 4:"ODD_PAGE"}.get(s.start_type, str(s.start_type))
    w = s.page_width / 36000
    h = s.page_height / 36000
    print(f"Section {i}: {w:.0f}x{h:.0f}mm start={st}")
    print(f"  margins: L={s.left_margin/36000:.0f} R={s.right_margin/36000:.0f} T={s.top_margin/36000:.0f} B={s.bottom_margin/36000:.0f}")

print()
print("=" * 80)
print("所有段落（含空段落） 位置和内容")
print("=" * 80)
for i, p in enumerate(doc.paragraphs):
    text = p.text
    style = p.style.name if p.style else "None"
    alignment = str(p.alignment) if p.alignment else "None"
    # 检查段落XML中是否有分页符
    xml = p._element.xml
    has_page_break = 'w:lastRenderedPageBreak' in xml or 'w:pageBreakBefore' in xml or '<w:br w:type="page"/>' in xml
    has_section_break = 'w:sectPr' in xml
    print(f"Para {i:>3}: style={style:<20} align={alignment:<15} text=\"{text[:60]}\"", end="")
    if has_page_break:
        print(" [PAGE_BREAK]", end="")
    if has_section_break:
        print(" [SECTION_BREAK]", end="")
    print()

print()
print("=" * 80)
print("表格分析 - 逐格")
print("=" * 80)
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti}: {len(table.rows)}行 x {len(table.columns)}列 ---")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            if not text:
                continue
            phs = re.findall(r'\{\{(.+?)\}\}', text)
            if phs:
                print(f"  [{ri},{ci}] {{占位符: {phs}}} -> \"{text[:120]}\"")
            else:
                print(f"  [{ri},{ci}] \"{text[:120]}\"")

print()
print("=" * 80)
print("表格顺序确认（XML中body内顺序）")
print("=" * 80)
body = doc.element.body
table_count = 0
para_count = 0
for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'tbl':
        print(f"  TABLE {table_count} (在body中第{para_count + table_count}个元素)")
        table_count += 1
    elif tag == 'p':
        text = ''.join(child.itertext()).strip()
        if text:
            print(f"  PARA \"{text[:80]}\"")
        para_count += 1
    elif tag == 'sectPr':
        print(f"  SECTION_BREAK")
    else:
        print(f"  {tag}")