#!/usr/bin/env python3
"""检查模板文件中的占位符"""
from docx import Document
import os

def check():
    template_path = r"D:\erp_thirteen\tp_education_system\backend\uploads\templates\20260208_201747_职工退休呈报表.docx"
    
    if not os.path.exists(template_path):
        print(f"模板文件不存在: {template_path}")
        return
    
    doc = Document(template_path)
    
    print("模板内容:")
    print("=" * 50)
    
    # 提取所有文本
    all_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text)
            print(f"段落: {para.text}")
    
    # 检查表格
    for i, table in enumerate(doc.tables):
        print(f"\n表格 {i+1}:")
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            print(f"  {' | '.join(row_text)}")
    
    print("\n" + "=" * 50)
    print("\n检查占位符格式:")
    
    # 查找 {{xxx}} 格式的占位符
    import re
    full_text = '\n'.join(all_text)
    placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)
    
    if placeholders:
        print(f"找到 {len(placeholders)} 个占位符:")
        for ph in placeholders:
            print(f"  - {{ {ph} }}")
    else:
        print("未找到 {{xxx}} 格式的占位符")
        print("\n模板中应该使用 {{字段名}} 格式作为占位符，例如:")
        print("  - {{teacher_name}}")
        print("  - {{gender}}")
        print("  - {{birth_date}}")

if __name__ == '__main__':
    check()
