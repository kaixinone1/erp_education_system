#!/usr/bin/env python3
"""
模板字段映射API - 通用版本
用于管理模板占位符与中间表字段的映射关系
支持自动映射、智能推荐、多格式占位符
"""
from fastapi import APIRouter, HTTPException, Body
from typing import List, Dict, Any, Optional
import json
import os
import sys
import psycopg2
from datetime import datetime

# 添加services目录到路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from services.universal_placeholder_extractor import extract_fields
from services.field_mapping_service import (
    get_table_fields_from_db,
    apply_system_field_names,
    auto_map_fields,
    get_intermediate_tables as get_tables,
    save_field_mapping,
    load_field_mapping,
    load_mapping_config
)

router = APIRouter(prefix="/api/template-field-mapping", tags=["template-field-mapping"])

# 数据库连接配置
DATABASE_CONFIG = {
    'host': 'localhost',
    'port': 5432,
    'database': 'taiping_education',
    'user': 'taiping_user',
    'password': 'taiping_password'
}

# 配置文件路径
CONFIG_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'config')
SCHEMA_FILE = os.path.join(CONFIG_DIR, 'merged_schema_mappings.json')
TABLE_MAPPINGS_FILE = os.path.join(CONFIG_DIR, 'table_name_mappings.json')


def get_db_connection():
    """获取数据库连接"""
    return psycopg2.connect(**DATABASE_CONFIG)


def read_json_file(file_path: str) -> dict:
    """读取JSON文件"""
    try:
        if os.path.exists(file_path):
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    except Exception as e:
        print(f"读取文件失败 {file_path}: {e}")
        return {}


@router.get("/intermediate-tables")
async def get_intermediate_tables():
    """获取所有中间表列表（通用版本）"""
    try:
        # 使用通用服务获取中间表列表
        tables = get_tables()
        
        # 转换为前端期望的格式
        result = []
        for table in tables:
            result.append({
                "name": table.get("name"),
                "name_cn": table.get("label", table.get("name")),
                "type": table.get("type", "intermediate")
            })
        
        return {
            "status": "success",
            "tables": result
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取中间表列表失败: {str(e)}")


@router.get("/table-fields/{table_name}")
async def get_table_fields(table_name: str):
    """获取指定中间表的所有字段（通用版本）"""
    try:
        # 使用通用服务获取表字段
        fields_data = get_table_fields_from_db(table_name)
        fields_data = apply_system_field_names(fields_data)
        
        # 转换为前端期望的格式
        fields = []
        for field in fields_data:
            fields.append({
                "name": field.get("name"),
                "name_cn": field.get("label", field.get("name")),
                "type": field.get("type", "VARCHAR")
            })
        
        return {
            "status": "success",
            "table_name": table_name,
            "table_name_cn": table_name,
            "fields": fields
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取表字段失败: {str(e)}")


@router.get("/template-placeholders/{template_id}")
async def get_template_placeholders(template_id: str):
    """获取模板中的所有占位符（通用版本）"""
    try:
        # 支持字符串类型的模板ID
        template_id_str = template_id
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 从document_templates表获取模板内容（支持字符串ID）
        cursor.execute("""
            SELECT template_name, file_path 
            FROM document_templates 
            WHERE template_id = %s
        """, (template_id_str,))
        
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="模板不存在")
        
        template_name, file_path = row
        
        # 使用通用占位符提取服务
        placeholders = []
        if file_path and os.path.exists(file_path):
            try:
                # 获取文件扩展名
                file_ext = os.path.splitext(file_path)[1].lower()
                # 使用通用提取器
                fields = extract_fields(file_path, file_ext)
                
                # 去重并转换为前端格式
                seen = set()
                for field in fields:
                    name = field.get('name', '')
                    if name and name not in seen:
                        seen.add(name)
                        placeholders.append({
                            "name": name,
                            "name_cn": field.get('label', name)
                        })
            except Exception as e:
                print(f"解析模板文件失败: {e}")
        
        cursor.close()
        conn.close()
        
        return {
            "status": "success",
            "template_id": template_id,
            "template_name": template_name,
            "placeholders": placeholders
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取模板占位符失败: {str(e)}")


@router.post("/save-mapping")
async def save_field_mapping(data: dict = Body(...)):
    """保存字段映射关系（增量更新）"""
    try:
        template_id = data.get("template_id")
        template_name = data.get("template_name")
        intermediate_table = data.get("intermediate_table")
        intermediate_table_cn = data.get("intermediate_table_cn")
        mappings = data.get("mappings", [])  # [{placeholder, field, field_cn}, ...]
        
        if not template_id or not intermediate_table:
            raise HTTPException(status_code=400, detail="缺少必要参数")
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 获取现有映射
        cursor.execute("""
            SELECT placeholder_name, intermediate_field
            FROM template_field_mapping
            WHERE template_id = %s
        """, (template_id,))
        
        existing_mappings = {row[0]: row[1] for row in cursor.fetchall()}
        
        # 统计变更
        updated_count = 0
        inserted_count = 0
        unchanged_count = 0
        
        for mapping in mappings:
            placeholder = mapping.get("placeholder")
            field = mapping.get("field")
            field_cn = mapping.get("field_cn")
            
            if placeholder in existing_mappings:
                if existing_mappings[placeholder] != field:
                    # 更新现有映射
                    cursor.execute("""
                        UPDATE template_field_mapping
                        SET intermediate_table = %s,
                            intermediate_table_cn = %s,
                            intermediate_field = %s,
                            intermediate_field_cn = %s,
                            updated_at = CURRENT_TIMESTAMP
                        WHERE template_id = %s AND placeholder_name = %s
                    """, (
                        intermediate_table,
                        intermediate_table_cn,
                        field,
                        field_cn,
                        template_id,
                        placeholder
                    ))
                    updated_count += 1
                else:
                    unchanged_count += 1
            else:
                # 插入新映射
                cursor.execute("""
                    INSERT INTO template_field_mapping 
                    (template_id, template_name, placeholder_name, 
                     intermediate_table, intermediate_table_cn,
                     intermediate_field, intermediate_field_cn)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (
                    template_id,
                    template_name,
                    placeholder,
                    intermediate_table,
                    intermediate_table_cn,
                    field,
                    field_cn
                ))
                inserted_count += 1
        
        conn.commit()
        cursor.close()
        conn.close()
        
        return {
            "status": "success",
            "message": f"保存完成：新增 {inserted_count} 个，更新 {updated_count} 个，未变更 {unchanged_count} 个"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"保存映射失败: {str(e)}")


@router.get("/get-mapping/{template_id}")
async def get_field_mapping(template_id: str):
    """获取模板的字段映射关系"""
    try:
        # 支持字符串类型的模板ID
        template_id_str = template_id
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT id, template_id, template_name, placeholder_name,
                   intermediate_table, intermediate_table_cn,
                   intermediate_field, intermediate_field_cn, is_active
            FROM template_field_mapping
            WHERE template_id = %s
            ORDER BY id
        """, (template_id_str,))
        
        rows = cursor.fetchall()
        
        mappings = []
        intermediate_table = ""
        intermediate_table_cn = ""
        
        for row in rows:
            mappings.append({
                "id": row[0],
                "placeholder": row[3],
                "field": row[6],
                "field_cn": row[7],
                "is_active": row[8]
            })
            if not intermediate_table:
                intermediate_table = row[4]
                intermediate_table_cn = row[5]
        
        cursor.close()
        conn.close()
        
        return {
            "status": "success",
            "template_id": template_id,
            "intermediate_table": intermediate_table,
            "intermediate_table_cn": intermediate_table_cn,
            "mappings": mappings
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取映射失败: {str(e)}")


@router.get("/fill-data/{template_id}")
async def get_fill_data(template_id: str, teacher_id: int):
    """获取填报数据（强制字段映射版本）
    
    所有模板都通过中间表获取数据，所有占位符都必须配置字段映射：
    1. 获取模板信息
    2. 从模板中提取所有 {{占位符}}
    3. 读取字段映射配置（占位符 → 中间表字段）
    4. 检查所有占位符是否都有映射配置
    5. 根据映射从中间表查询数据
    6. 返回填充数据
    """
    print(f"【DEBUG】get_fill_data 被调用: template_id={template_id}, teacher_id={teacher_id}")
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 1. 获取模板信息
        cursor.execute("""
            SELECT file_path, intermediate_table 
            FROM document_templates 
            WHERE template_id = %s
        """, (template_id,))
        
        template_row = cursor.fetchone()
        if not template_row:
            print(f"【DEBUG】模板不存在: {template_id}")
            raise HTTPException(status_code=404, detail="模板不存在")
        
        file_path = template_row[0]
        default_table = template_row[1]  # 模板关联的默认中间表
        print(f"【DEBUG】模板文件路径: {file_path}, 默认表: {default_table}")
        
        # 2. 从模板中提取占位符
        file_format = file_path.split('.')[-1].lower() if '.' in file_path else 'html'
        placeholders = _extract_placeholders_from_file(file_path, file_format)
        print(f"【DEBUG】提取到的占位符数量: {len(placeholders)}")
        print(f"【DEBUG】前5个占位符: {placeholders[:5]}")
        
        if not placeholders:
            print(f"【DEBUG】没有提取到占位符")
            return {
                "status": "success",
                "template_id": template_id,
                "teacher_id": teacher_id,
                "data": {}
            }
        
        # 3. 读取字段映射配置（占位符 → 中间表字段）
        cursor.execute("""
            SELECT placeholder_name, intermediate_table, intermediate_field, aggregate_func, filter_condition
            FROM template_field_mapping
            WHERE template_id = %s AND is_active = true
        """, (template_id,))
        
        mappings = {}
        rows = cursor.fetchall()
        print(f"【DEBUG】数据库中的字段映射数量: {len(rows)}")
        
        for row in rows:
            placeholder_name = row[0]
            intermediate_table = row[1] or default_table
            intermediate_field = row[2]
            aggregate_func = row[3] if len(row) > 3 else ''
            filter_condition = row[4] if len(row) > 4 else ''
            if intermediate_field:  # 必须有字段映射
                mappings[placeholder_name] = {
                    'table': intermediate_table,
                    'field': intermediate_field,
                    'aggregate_func': aggregate_func,
                    'filter_condition': filter_condition
                }
        
        print(f"【DEBUG】构建的 mappings 数量: {len(mappings)}")
        print(f"【DEBUG】mappings 的 keys: {list(mappings.keys())[:5]}")
        
        # 4. 检查所有占位符是否都有映射配置
        unmapped_placeholders = []
        for placeholder in placeholders:
            field_name = placeholder.strip('{}')
            if field_name not in mappings:
                unmapped_placeholders.append(field_name)
        
        print(f"【DEBUG】未映射的占位符: {unmapped_placeholders}")
        
        if unmapped_placeholders:
            print(f"【DEBUG】存在未映射占位符，返回错误")
            return {
                "status": "error",
                "template_id": template_id,
                "teacher_id": teacher_id,
                "message": f"以下占位符未配置字段映射: {', '.join(unmapped_placeholders)}",
                "unmapped_placeholders": unmapped_placeholders,
                "data": {}
            }
        
        # 5. 从中间表查询数据
        data = {}
        for placeholder in placeholders:
            field_name = placeholder.strip('{}')
            mapping = mappings[field_name]
            table_name = mapping['table']
            field = mapping['field']
            aggregate_func = mapping.get('aggregate_func', '')
            filter_condition = mapping.get('filter_condition', '')
            
            try:
                # 验证表名和字段名，防止SQL注入
                # 只允许字母、数字、下划线和中文字符
                import re
                if not re.match(r'^[\w\u4e00-\u9fa5]+$', table_name):
                    print(f"表名包含非法字符: {table_name}")
                    data[field_name] = ''
                    continue
                if not re.match(r'^[\w\u4e00-\u9fa5]+$', field):
                    print(f"字段名包含非法字符: {field}")
                    data[field_name] = ''
                    continue
                
                # 构建WHERE条件
                where_conditions = [f'teacher_id = %s']
                params = [teacher_id]
                
                # 添加自定义过滤条件
                if filter_condition:
                    # 支持多种条件格式：
                    # 1. 单条件: 字段名='值'
                    # 2. 多条件: 字段名='值' AND 字段2='值2'
                    # 3. OR条件: 字段名='值1' OR 字段名='值2'
                    # 安全验证：只允许中文、字母、数字、下划线、单引号、=、AND、OR、空格
                    safe_pattern = r"^[\w\u4e00-\u9fa5\s'\"=]+(?:\s+(?:AND|OR)\s+[\w\u4e00-\u9fa5\s'\"=]+)*$"
                    if re.match(safe_pattern, filter_condition.strip(), re.IGNORECASE):
                        where_conditions.append(filter_condition.strip())
                    else:
                        print(f"过滤条件格式不安全，已忽略: {filter_condition}")
                
                where_clause = ' AND '.join(where_conditions)
                
                # 判断是否使用聚合函数
                if aggregate_func and aggregate_func.upper() in ('COUNT', 'SUM', 'MAX', 'MIN', 'AVG'):
                    # 聚合查询：统计符合条件的记录数/总和等
                    safe_field = f'"{field}"' if field != '*' else field
                    query = f'SELECT {aggregate_func}({safe_field}) FROM "{table_name}" WHERE {where_clause}'
                    cursor.execute(query, tuple(params))
                else:
                    # 普通查询
                    query = f'SELECT "{field}" FROM "{table_name}" WHERE {where_clause}'
                    cursor.execute(query, tuple(params))
                
                row = cursor.fetchone()
                if row and row[0] is not None:
                    data[field_name] = row[0]
                else:
                    data[field_name] = ''
            except Exception as e:
                print(f"查询字段失败 {field_name} from {table_name}.{field}: {e}")
                data[field_name] = ''
                # 回滚事务，避免影响后续查询
                try:
                    conn.rollback()
                except:
                    pass
        
        cursor.close()
        conn.close()
        
        # 6. 自动计算小计字段
        # 规则：如果字段名以"小计"结尾，尝试计算 人数 * 月工资标准
        for field_name in list(data.keys()):
            if field_name.endswith('小计'):
                # 提取前缀，如 "正高级小计" -> "正高级"
                prefix = field_name[:-2]  # 去掉"小计"
                
                # 寻找对应的人数和工资字段
                count_field = f"{prefix}人数"
                salary_field = f"{prefix}月工资标准"
                
                # 如果人数和工资字段都存在，计算小计
                if count_field in data and salary_field in data:
                    try:
                        count_val = float(data[count_field]) if data[count_field] else 0
                        salary_val = float(data[salary_field]) if data[salary_field] else 0
                        data[field_name] = count_val * salary_val
                        print(f"【DEBUG】自动计算 {field_name} = {count_val} * {salary_val} = {data[field_name]}")
                    except (ValueError, TypeError) as e:
                        print(f"【DEBUG】计算 {field_name} 失败: {e}")
                        data[field_name] = ''
        
        # 7. 自动计算合计字段
        # 规则：如果字段名包含"合计"或"总计"，尝试汇总相关字段
        for field_name in list(data.keys()):
            if '合计' in field_name or '总计' in field_name:
                # 提取类别前缀，如 "行政管理人员合计" -> 找所有"行政管理人员"相关的小计
                # 或者找所有小计字段求和
                total = 0
                has_value = False
                
                # 尝试找到所有小计字段并求和
                for key in data.keys():
                    if key.endswith('小计') and data[key]:
                        try:
                            val = float(data[key])
                            total += val
                            has_value = True
                        except (ValueError, TypeError):
                            pass
                
                if has_value:
                    data[field_name] = total
                    print(f"【DEBUG】自动计算 {field_name} = {total}")
        
        # 8. 自动计算绩效人数（所有类别人数之和）
        if '绩效人数' in data:
            total_count = 0
            has_value = False
            for key in data.keys():
                if key.endswith('人数') and key != '绩效人数' and data[key]:
                    try:
                        val = float(data[key])
                        total_count += val
                        has_value = True
                    except (ValueError, TypeError):
                        pass
            if has_value:
                data['绩效人数'] = int(total_count)
                print(f"【DEBUG】自动计算 绩效人数 = {total_count}")
        
        return {
            "status": "success",
            "template_id": template_id,
            "teacher_id": teacher_id,
            "data": data
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取填报数据失败: {str(e)}")


def _extract_placeholders_from_file(file_path: str, file_format: str) -> list:
    """从文件中提取 {{占位符}}"""
    placeholders = []
    
    if not os.path.exists(file_path):
        return placeholders
    
    try:
        if file_format == 'docx':
            # Word文档
            from docx import Document
            doc = Document(file_path)
            
            # 从段落中提取
            for para in doc.paragraphs:
                placeholders.extend(_extract_placeholders_from_text(para.text))
            
            # 从表格中提取
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        placeholders.extend(_extract_placeholders_from_text(cell.text))
        
        elif file_format in ('html', 'htm'):
            # HTML文件 - 尝试多种编码
            content = None
            for encoding in ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content:
                placeholders.extend(_extract_placeholders_from_text(content))
            else:
                print(f"无法读取HTML文件: {file_path}")
        
        elif file_format == 'xlsx':
            # Excel文件
            from openpyxl import load_workbook
            wb = load_workbook(file_path)
            ws = wb.active
            
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value:
                        placeholders.extend(_extract_placeholders_from_text(str(cell.value)))
        
        # 去重
        return list(set(placeholders))
    
    except Exception as e:
        print(f"提取占位符失败: {e}")
        return placeholders


def _extract_placeholders_from_text(text: str) -> list:
    """从文本中提取 {{占位符}}，并清理HTML标签"""
    import re
    from html.parser import HTMLParser
    
    try:
        # 提取 {{...}} 中的内容
        pattern = r'\{\{([^}]+)\}\}'
        matches = re.findall(pattern, text)
        
        # 清理HTML标签，只保留纯文本
        class MLStripper(HTMLParser):
            def __init__(self):
                super().__init__()
                self.reset()
                self.fed = []
            def handle_data(self, d):
                self.fed.append(d)
            def get_data(self):
                return ''.join(self.fed)
        
        def strip_tags(html):
            if not html:
                return ''
            s = MLStripper()
            try:
                s.feed(html)
                return s.get_data()
            except Exception as e:
                # 如果解析失败，使用正则表达式移除标签
                print(f"HTML解析失败，使用正则移除标签: {e}")
                return re.sub(r'<[^>]+>', '', html)
        
        cleaned_matches = []
        for match in matches:
            try:
                # 清理HTML标签
                clean_text = strip_tags(match).strip()
                # 移除多余的空白字符
                clean_text = re.sub(r'\s+', '', clean_text)
                if clean_text:  # 只保留非空的
                    # 验证占位符内容，只允许字母、数字、下划线和中文字符
                    if re.match(r'^[\w\u4e00-\u9fa5]+$', clean_text):
                        cleaned_matches.append('{{' + clean_text + '}}')
                    else:
                        print(f"占位符包含非法字符，跳过: {clean_text}")
            except Exception as e:
                print(f"处理占位符时出错: {e}, match: {match}")
                continue
        
        return cleaned_matches
    except Exception as e:
        print(f"提取占位符时出错: {e}")
        return []


@router.get("/preview/{template_id}")
async def preview_template_with_mapping(template_id: str, teacher_id: int = 0, mode: str = "fill", debug: bool = False):
    """预览填充后的模板（使用字段映射配置）
    
    参数:
        template_id: 模板ID
        teacher_id: 教师ID（mode=fill时必填，mode=preview时可选）
        mode: 模式，fill=填报模式（填充数据），preview=预览模式（保留占位符）
        debug: 是否返回调试信息（替换统计）
    """
    try:
        # 打印调试信息
        print(f"【API】preview_template_with_mapping called: template_id={template_id}, teacher_id={teacher_id}, mode={mode}")
        
        # 使用新的占位符模板引擎
        from services.placeholder_template_engine import PlaceholderTemplateEngine
        engine = PlaceholderTemplateEngine(get_db_connection)
        
        # 根据模式生成文档
        if mode == "preview":
            # 预览模式：返回原始模板，不填充数据
            html_content = engine.get_raw_template(template_id)
        else:
            # 填报模式：填充数据
            print(f"【API】生成文档: template_id={template_id}, teacher_id={teacher_id}")
            html_content = engine.generate_document(template_id, teacher_id)
            print(f"【API】文档生成完成，长度={len(html_content)}")
        
        # 如果需要调试信息，返回JSON格式
        if debug and mode == "fill":
            import re
            import json
            
            # 获取模板配置
            config = engine.get_field_mapping(template_id)
            
            # 读取原始模板
            raw_html = engine.get_raw_template(template_id)
            
            # 提取所有占位符
            all_placeholders = re.findall(r'\{\{([^{}]+)\}\}', raw_html)
            all_placeholders = [p.strip() for p in all_placeholders]
            total_placeholders = len(all_placeholders)
            unique_placeholders = list(set(all_placeholders))
            
            # 检查哪些占位符已被替换
            remaining = re.findall(r'\{\{([^{}]+)\}\}', html_content)
            remaining = [p.strip() for p in remaining]
            remaining_unique = list(set(remaining))
            
            # 计算替换统计
            replaced_count = total_placeholders - len(remaining)
            replaced_unique = len(unique_placeholders) - len(remaining_unique)
            
            # 获取字段映射信息
            mapping_info = {}
            for mapping in config.get('mappings', []):
                placeholder = mapping['placeholder'].strip()
                if placeholder.startswith('{{') and placeholder.endswith('}}'):
                    placeholder = placeholder[2:-2]
                elif placeholder.startswith('{') and placeholder.endswith('}'):
                    placeholder = placeholder[1:-1]
                mapping_info[placeholder] = {
                    'table': mapping.get('table', ''),
                    'field': mapping.get('field', '')
                }
            
            # 构建未替换占位符的详细信息
            unmapped_details = []
            for p in remaining_unique:
                info = mapping_info.get(p, {})
                unmapped_details.append({
                    'placeholder': p,
                    'mapped': p in mapping_info,
                    'table': info.get('table', ''),
                    'field': info.get('field', '')
                })
            
            debug_info = {
                'status': 'success',
                'template_id': template_id,
                'teacher_id': teacher_id,
                'statistics': {
                    'total_placeholders': total_placeholders,
                    'unique_placeholders': len(unique_placeholders),
                    'replaced_count': replaced_count,
                    'replaced_unique': replaced_unique,
                    'remaining_count': len(remaining),
                    'remaining_unique': len(remaining_unique),
                    'replacement_rate': f"{replaced_unique}/{len(unique_placeholders)} ({replaced_unique/len(unique_placeholders)*100:.1f}%)" if unique_placeholders else "N/A"
                },
                'all_placeholders': unique_placeholders,
                'unmapped_placeholders': unmapped_details,
                'html_preview': html_content[:2000] + '...' if len(html_content) > 2000 else html_content
            }
            
            return debug_info
        
        from fastapi.responses import HTMLResponse
        return HTMLResponse(content=html_content)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"预览生成失败: {str(e)}")


@router.post("/preview-with-data/{template_id}")
async def preview_with_override_data(template_id: str, request: dict = Body(...)):
    """使用覆盖数据预览模板"""
    try:
        teacher_id = request.get("teacher_id", 0)
        override_data = request.get("override_data", {})
        
        # 使用新的占位符模板引擎
        from services.placeholder_template_engine import PlaceholderTemplateEngine
        engine = PlaceholderTemplateEngine(get_db_connection)
        
        # 获取模板配置
        config = engine.get_field_mapping(template_id)
        template_path = config['file_path']
        mappings = config['mappings']
        
        if not os.path.exists(template_path):
            raise ValueError(f"模板文件不存在: {template_path}")
        
        # 获取基础数据
        base_data = {}
        for mapping in mappings:
            placeholder = mapping['placeholder']
            table = mapping['table']
            field = mapping['field']
            value = engine.get_data_from_table(table, field, teacher_id)
            base_data[placeholder] = value
        
        # 使用覆盖数据更新
        base_data.update(override_data)
        
        # 读取模板文件
        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030']
        content = None
        for encoding in encodings:
            try:
                with open(template_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    break
            except:
                continue
        
        if not content:
            raise ValueError("无法读取模板文件")
        
        # 替换占位符
        import re
        for placeholder, value in base_data.items():
            pattern = r'\{\{' + re.escape(placeholder) + r'\}\}'
            content = re.sub(pattern, value, content)
        
        from fastapi.responses import HTMLResponse
        return HTMLResponse(content=content)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"预览生成失败: {str(e)}")


@router.post("/save-html/{template_id}")
async def save_html_content(template_id: str, request: dict = Body(...)):
    """
    保存编辑后的HTML内容
    将编辑后的内容保存为新的模板文件或覆盖原文件
    """
    try:
        teacher_id = request.get("teacher_id", 0)
        html_content = request.get("html_content", "")
        
        if not html_content:
            raise HTTPException(status_code=400, detail="HTML内容不能为空")
        
        # 获取模板信息
        conn = get_db_connection()
        cursor = conn.cursor()
        
        try:
            # 获取模板基本信息（支持id或template_id查询）
            template_row = None
            
            # 统一使用template_id字段查询（支持字符串ID）
            cursor.execute("""
                SELECT id, template_id, template_name, file_path
                FROM document_templates
                WHERE template_id = %s
            """, (template_id,))
            template_row = cursor.fetchone()
            
            if not template_row:
                raise HTTPException(status_code=404, detail=f"模板不存在: {template_id}")
            
            actual_template_id = template_row[0]
            template_name = template_row[2]
            original_file_path = template_row[3]
            
            # 直接覆盖原文件
            import os
            
            if original_file_path and os.path.exists(original_file_path):
                # 覆盖原文件
                file_path = original_file_path
            else:
                # 如果原文件路径不存在，使用默认路径
                save_dir = os.path.join(os.path.dirname(__file__), '..', 'templates')
                file_path = os.path.join(save_dir, template_name)
            
            # 确保目录存在
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # 写入文件（覆盖原文件）
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return {
                "status": "success",
                "message": "保存成功",
                "file_name": template_name,
                "file_path": file_path
            }
            
        finally:
            cursor.close()
            conn.close()
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"保存失败: {str(e)}")
