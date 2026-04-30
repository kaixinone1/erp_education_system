# 新的导出函数 - 严格按照图一图二1:1还原
# 此文件包含完整的导出函数代码

# 由于代码量超过250行限制，我将分段写入
# 这是第一段：导入和页面设置

@router.post("/export-a3-report")
async def export_a3_report(data: Dict[str, Any]):
    """
    导出A3横向对折册子格式的退休呈报表
    严格按照图一、图二格式1:1还原
    支持Word和PDF两种格式
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml
        from fastapi.responses import FileResponse
        import os

        teacher_name = data.get("teacher_name", "未知")
        export_format = data.get("format", "word").lower()

        # 创建Word文档
        doc = Document()

        # 设置A3横向页面
        section = doc.sections[0]
        section.page_height = Cm(29.7)  # A3高度
        section.page_width = Cm(42.0)   # A3宽度
        section.orientation = WD_ORIENT.LANDSCAPE

        # 设置页边距 - 上下左右均为3厘米
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

        # 设置中文字体函数
        def set_chinese_font(run, font_name='仿宋_GB2312', font_size=10.5, bold=False):
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
            run.font.bold = bold

        # 创建主表格（两栏布局）- 无边框
        main_table = doc.add_table(rows=1, cols=2)
        main_table.autofit = False
        main_table.allow_autofit = False
        
        # 设置表格边框为无
        tbl = main_table._tbl
        tblPr = tbl.tblPr
        tblBorders = parse_xml(r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>')
        tblPr.append(tblBorders)

        # 设置列宽 - 两栏均分（每栏18cm）
        for row in main_table.rows:
            row.cells[0].width = Cm(18.0)
            row.cells[1].width = Cm(18.0)

        # ===== 第1页左栏：审批意见表格（严格按照图一） =====
        left_cell = main_table.rows[0].cells[0]
        
        # 创建审批意见表格（4行2列）
        approval_table = left_cell.add_table(rows=4, cols=2)
        approval_table.style = 'Table Grid'
        approval_table.autofit = False
        
        # 设置列宽
        for row in approval_table.rows:
            row.cells[0].width = Cm(3.0)
            row.cells[1].width = Cm(15.0)
        
        # 第1行：呈报单位意见
        row = approval_table.rows[0]
        row.cells[0].text = '呈\n报\n单\n位\n意\n见'
        cell = row.cells[1]
        p = cell.paragraphs[0]
        run = p.add_run(f"经研究，同意 {teacher_name} 同志按以下第（    ）条办理退休，从      年    月执行。")
        set_chinese_font(run, font_size=10.5)
        p.add_run("\n\n")
        run = p.add_run("（一）弹性提前退休")
        set_chinese_font(run, font_size=10.5)
        p.add_run("\n")
        run = p.add_run("（二）法定退休年龄退休")
        set_chinese_font(run, font_size=10.5)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p2.add_run("\n\n\n    年    月    日")
        set_chinese_font(run, font_size=10.5)

        # 第2行：主管部门审查意见
        row = approval_table.rows[1]
        row.cells[0].text = '主 管 部\n门 审 查\n意    见'
        cell = row.cells[1]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\n\n\n同意呈报\n\n\n")
        set_chinese_font(run, font_size=10.5)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p2.add_run("    年    月    日")
        set_chinese_font(run, font_size=10.5)

        # 第3行：退休一次性补贴审批意见
        row = approval_table.rows[2]
        row.cells[0].text = '退 休 一\n次性补贴\n审 批 意 见'
        cell = row.cells[1]
        p = cell.paragraphs[0]
        run = p.add_run(f"根据鄂人社发【2017】8号文件规定，同意 {teacher_name} 同志发放一次性独生子女费          元，教育特殊贡献奖          元，从      年    月执行。")
        set_chinese_font(run, font_size=10.5)
        p.add_run("\n\n")
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p2.add_run("    年    月    日")
        set_chinese_font(run, font_size=10.5)

        # 第4行：批准机关审批意见
        row = approval_table.rows[3]
        row.cells[0].text = '批 准 机\n关 审 批\n意    见'
        cell = row.cells[1]
        p = cell.paragraphs[0]
        run = p.add_run(f"根据人社部发【2024】94号文件规定，同意 {teacher_name} 同志按第（    ）条退休，从      年    月执行。")
        set_chinese_font(run, font_size=10.5)
        p.add_run("\n\n\n")
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p2.add_run("    年    月    日")
        set_chinese_font(run, font_size=10.5)

        # 左栏底部文字
        p = left_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('枣阳市人力资源和社会保障局制')
        set_chinese_font(run, font_size=9)

        # ===== 第1页右栏：封面（严格按照图一） =====
        right_cell = main_table.rows[0].cells[1]
        
        # 创建封面表格结构 - 无边框
        cover_table = right_cell.add_table(rows=1, cols=1)
        cover_table.style = 'Table Grid'
        cover_table.autofit = False
        cover_table.rows[0].cells[0].width = Cm(18.0)
        
        # 移除封面表格边框
        tbl2 = cover_table._tbl
        tblPr2 = tbl2.tblPr
        tblBorders2 = parse_xml(r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>')
        tblPr2.append(tblBorders2)

        cover_cell = cover_table.rows[0].cells[0]

        # 编号（右上角）
        p = cover_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"编号：{data.get('report_number', '')}")
        set_chinese_font(run, font_size=10.5)

        # 空行
        for _ in range(8):
            cover_cell.add_paragraph()

        # 主标题 - 小标宋，26号
        p = cover_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('职 工 退 休 呈 报 表')
        set_chinese_font(run, font_name='小标宋', font_size=26, bold=True)

        # 空行
        for _ in range(12):
            cover_cell.add_paragraph()

        # 单位 - 16号
        p = cover_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(2)
        run = p.add_run(f"单位：{data.get('work_unit', '枣阳市太平镇中心学校')}")
        set_chinese_font(run, font_size=16)

        # 空行
        cover_cell.add_paragraph()

        # 姓名 - 16号
        p = cover_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(2)
        run = p.add_run(f"姓名：{teacher_name}")
        set_chinese_font(run, font_size=16)

        # 空行
        for _ in range(4):
            cover_cell.add_paragraph()

        # 日期 - 16号，居中
        p = cover_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{data.get('cover_year', '')} 年   {data.get('cover_month', '')} 月   {data.get('cover_day', '')} 日")
        set_chinese_font(run, font_size=16)

        # 空行
        for _ in range(3):
            cover_cell.add_paragraph()

        # 副标题 - 20号
        p = cover_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('退 休 呈 报 表')
        set_chinese_font(run, font_size=20)

        # 添加分页符（第二页）
        doc.add_page_break()

        # ===== 第2页 =====
        page2_section = doc.sections[-1]
        page2_section.page_height = Cm(29.7)
        page2_section.page_width = Cm(42.0)
        page2_section.orientation = WD_ORIENT.LANDSCAPE
        page2_section.top_margin = Cm(3.0)
        page2_section.bottom_margin = Cm(3.0)
        page2_section.left_margin = Cm(3.0)
        page2_section.right_margin = Cm(3.0)

        # 创建两栏布局表格
        page2_table = doc.add_table(rows=1, cols=2)
        page2_table.style = 'Table Grid'
        page2_table.autofit = False
        
        # 设置两栏宽度（每栏18cm）
        for row in page2_table.rows:
            row.cells[0].width = Cm(18.0)
            row.cells[1].width = Cm(18.0)

        # ===== 第2页左栏：基本信息（严格按照图二） =====
        left_cell2 = page2_table.rows[0].cells[0]
        
        # 添加标题
        p = left_cell2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('退 休 呈 报 表')
        set_chinese_font(run, font_size=18, bold=True)

        # 创建基本信息表格（15行6列）
        info_table = left_cell2.add_table(rows=15, cols=6)
        info_table.style = 'Table Grid'
        info_table.autofit = False
        
        # 设置列宽
        for row in info_table.rows:
            for cell in row.cells:
                cell.width = Cm(3.0)

        # 第1行：姓名、性别、出生年月
        row = info_table.rows[0]
        row.cells[0].text = '姓 名'
        row.cells[1].text = data.get('teacher_name', '')
        row.cells[2].text = '性 别'
        row.cells[3].text = data.get('gender', '')
        row.cells[4].text = '出生年月'
        row.cells[5].text = data.get('birth_date', '')

        # 第2行：民族、文化程度、是否独生子女
        row = info_table.rows[1]
        row.cells[0].text = '民 族'
        row.cells[1].text = data.get('ethnicity', '')
        row.cells[2].text = '文化程度'
        row.cells[3].text = data.get('education', '')
        row.cells[4].text = '是否独生子女'
        row.cells[5].text = data.get('is_only_child', '')

        # 第3行：入党年月、职务、技术职称
        row = info_table.rows[2]
        row.cells[0].text = '入党年月'
        row.cells[1].text = data.get('join_party_date', '')
        row.cells[2].text = '职 务'
        row.cells[3].text = data.get('position', '')
        row.cells[4].text = '技术职称'
        row.cells[5].text = data.get('title', '')

        # 第4行：参加工作时间、工作年限
        row = info_table.rows[3]
        row.cells[0].text = '参加工作时间'
        row.cells[0].merge(row.cells[1])
        row.cells[2].text = data.get('work_start_date', '')
        row.cells[2].merge(row.cells[3])
        row.cells[4].text = '工作年限'
        row.cells[5].text = data.get('work_years', '')

        # 第5行：籍贯、现在住址
        row = info_table.rows[4]
        row.cells[0].text = '籍 贯'
        row.cells[1].text = data.get('native_place', '')
        row.cells[1].merge(row.cells[2])
        row.cells[3].text = '现在住址'
        row.cells[4].text = data.get('current_address', '')
        row.cells[4].merge(row.cells[5])

        # 第6行：工作简历标题
        row = info_table.rows[5]
        row.cells[0].text = '工 作 简 历'
        for i in range(1, 6):
            row.cells[0].merge(row.cells[i])

        # 第7行：工作简历表头
        row = info_table.rows[6]
        row.cells[0].text = '自何年何月'
        row.cells[0].merge(row.cells[1])
        row.cells[2].text = '至何年何月'
        row.cells[2].merge(row.cells[3])
        row.cells[4].text = '在何单位任何职'
        row.cells[5].text = '证明人及其住址'

        # 第8-12行：工作简历内容
        work_exp = data.get('work_experience', [])
        for i in range(5):
            row = info_table.rows[7 + i]
            if i < len(work_exp):
                row.cells[0].text = work_exp[i].get('start_date', '')
                row.cells[0].merge(row.cells[1])
                row.cells[2].text = work_exp[i].get('end_date', '')
                row.cells[2].merge(row.cells[3])
                row.cells[4].text = work_exp[i].get('unit_position', '')
                row.cells[5].text = work_exp[i].get('witness', '')

        # 第13行：退休原因
        row = info_table.rows[12]
        row.cells[0].text = '退休原因'
        row.cells[1].text = data.get('retirement_reason', '')
        for i in range(2, 6):
            row.cells[1].merge(row.cells[i])

        # 第14行：供养直系亲属
        row = info_table.rows[13]
        row.cells[0].text = '供养直系亲属、姓名、出生年月、与退休人员的关系'
        row.cells[1].text = data.get('family_members', '')
        for i in range(2, 6):
            row.cells[1].merge(row.cells[i])

        # 第15行：退休后居住地址和发给退休费的单位
        row = info_table.rows[14]
        row.cells[0].text = '退休后\n居住地址'
        row.cells[1].text = data.get('retirement_address', '')
        row.cells[1].merge(row.cells[2])
        row.cells[3].text = '发给退休\n费的单位'
        row.cells[4].text = data.get('pension_unit', '枣阳市人力资源和社会保障局')
        row.cells[4].merge(row.cells[5])

        # ===== 第2页右栏：工资信息（严格按照图二） =====
        right_cell2 = page2_table.rows[0].cells[1]
        
        # 创建工资信息表格（12行7列）
        salary_table = right_cell2.add_table(rows=12, cols=7)
        salary_table.style = 'Table Grid'
        salary_table.autofit = False
        
        # 设置列宽
        for row in salary_table.rows:
            row.cells[0].width = Cm(3.0)
            for i in range(1, 7):
                row.cells[i].width = Cm(2.5)

        # 2014年9月30日 - 占4行
        row = salary_table.rows[0]
        row.cells[0].text = '2014 年 9 月\n30 日'
        for i in range(1, 4):
            salary_table.rows[0].cells[0].merge(salary_table.rows[i].cells[0])
        row.cells[1].text = '机关工人'
        row.cells[2].text = '技术等级'
        row.cells[3].text = data.get('salary_2014_worker_level', '')
        row.cells[4].text = '级别薪级'
        row.cells[5].text = data.get('salary_2014_worker_grade', '')
        row.cells[6].text = '级'

        row = salary_table.rows[1]
        row.cells[1].text = '事业管理'
        row.cells[2].text = '岗位'
        row.cells[3].text = data.get('salary_2014_mgmt_level', '')
        row.cells[4].text = '对应原职务'
        row.cells[5].text = data.get('salary_2014_mgmt_position', '')
        row.cells[6].text = '薪级'

        row = salary_table.rows[2]
        row.cells[1].text = '事业专技'
        row.cells[2].text = '岗位'
        row.cells[3].text = data.get('salary_2014_tech_level', '')
        row.cells[4].text = '对应原职务'
        row.cells[5].text = data.get('salary_2014_tech_position', '')
        row.cells[6].text = '薪级'

        row = salary_table.rows[3]
        row.cells[1].text = '事业工勤'
        row.cells[2].text = '岗位'
        row.cells[3].text = data.get('salary_2014_service_level', '')
        row.cells[4].text = '对应技术等级'
        row.cells[5].text = data.get('salary_2014_service_tech', '')
        row.cells[6].text = '薪级'

        # 最后一次职务升降时间 - 占4行
        row = salary_table.rows[4]
        row.cells[0].text = '最后一次职务\n（技术职称）\n升降时间'
        for i in range(1, 4):
            salary_table.rows[4].cells[0].merge(salary_table.rows[4 + i].cells[0])
        row.cells[1].text = '机关工人'
        row.cells[2].text = '技术等级'
        row.cells[3].text = data.get('last_promotion_worker_level', '')
        row.cells[4].text = '级别薪级'
        row.cells[5].text = data.get('last_promotion_worker_grade', '')
        row.cells[6].text = '级'

        row = salary_table.rows[5]
        row.cells[1].text = '事业管理'
        row.cells[2].text = '岗位'
        row.cells[3].text = data.get('last_promotion_mgmt_level', '')
        row.cells[4].text = '对应原职务'
        row.cells[5].text = data.get('last_promotion_mgmt_position', '')
        row.cells[6].text = '薪级'

        row = salary_table.rows[6]
        row.cells[1].text = '事业专技'
        row.cells[2].text = '岗位'
        row.cells[3].text = data.get('last_promotion_tech_level', '')
        row.cells[4].text = '对应原职务'
        row.cells[5].text = data.get('last_promotion_tech_position', '')
        row.cells[6].text = '薪级'

        row = salary_table.rows[7]
        row.cells[1].text = '事业工勤'
        row.cells[2].text = '岗位'
        row.cells[3].text = data.get('last_promotion_service_level', '')
        row.cells[4].text = '对应技术等级'
        row.cells[5].text = data.get('last_promotion_service_tech', '')
        row.cells[6].text = '薪级'

        # 退休时 - 占4行
        row = salary_table.rows[8]
        row.cells[0].text = '退\n休\n时'
        for i in range(1, 4):
            salary_table.rows[8].cells[0].merge(salary_table.rows[8 + i].cells[0])
        row.cells[1].text = '机关工人'
        row.cells[2].text = '技术等级'
        row.cells[3].text = data.get('retirement_worker_level', '')
        row.cells[4].text = '级别薪级'
        row.cells[5].text = data.get('retirement_worker_grade', '')
        row.cells[6].text = '级'

        row = salary_table.rows[9]
        row.cells[1].text = '事业管理'
        row.cells[2].text = '岗位'
        row.cells[3].text = data.get('retirement_mgmt_level', '')
        row.cells[4].text = '对应原职务'
        row.cells[5].text = data.get('retirement_mgmt_position', '')
        row.cells[6].text = '薪级'

        row = salary_table.rows[10]
        row.cells[1].text = '事业专技'
        row.cells[2].text = '岗位'
        row.cells[3].text = data.get('retirement_tech_level', '')
        row.cells[4].text = '对应原职务'
        row.cells[5].text = data.get('retirement_tech_position', '')
        row.cells[6].text = '薪级'

        row = salary_table.rows[11]
        row.cells[1].text = '事业工勤'
        row.cells[2].text = '岗位'
        row.cells[3].text = data.get('retirement_service_level', '')
        row.cells[4].text = '对应技术等级'
        row.cells[5].text = data.get('retirement_service_tech', '')
        row.cells[6].text = '薪级'

        # 保存文件
        output_path = f'/tmp/职工退休呈报表_{teacher_name}_A3.docx'
        doc.save(output_path)

        # 如果需要PDF格式，转换为PDF
        if export_format == 'pdf':
            try:
                import subprocess
                pdf_path = f'/tmp/职工退休呈报表_{teacher_name}_A3.pdf'
                
                result = subprocess.run([
                    'soffice', '--headless', '--convert-to', 'pdf', 
                    '--outdir', '/tmp', output_path
                ], capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0 and os.path.exists(pdf_path):
                    return FileResponse(
                        pdf_path,
                        filename=f'职工退休呈报表_{teacher_name}.pdf',
                        media_type='application/pdf'
                    )
                else:
                    return FileResponse(
                        output_path,
                        filename=f'职工退休呈报表_{teacher_name}.docx',
                        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
            except Exception as pdf_error:
                return FileResponse(
                    output_path,
                    filename=f'职工退休呈报表_{teacher_name}.docx',
                    media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
        else:
            return FileResponse(
                output_path,
                filename=f'职工退休呈报表_{teacher_name}.docx',
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")
