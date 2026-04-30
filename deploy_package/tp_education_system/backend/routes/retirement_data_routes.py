"""
退休呈报表数据管理API
"""
from fastapi import APIRouter, HTTPException, Query
from typing import List, Dict, Any, Optional
from datetime import datetime, date
import psycopg2
import sys
import os

# 添加 utils 目录到路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from utils.dict_utils import get_education_name
from utils.retirement_calculator import calculate_retirement_info

router = APIRouter(prefix="/api/retirement-data", tags=["retirement-data"])

DATABASE_CONFIG = {
    "host": "localhost",
    "port": "5432",
    "database": "taiping_education",
    "user": "taiping_user",
    "password": "taiping_password"
}


def get_db_connection():
    """获取数据库连接"""
    return psycopg2.connect(**DATABASE_CONFIG)


def extract_birth_date_from_id_card(id_card: str) -> Optional[str]:
    """从身份证号提取出生日期"""
    if not id_card or len(id_card) != 18:
        return None
    try:
        year = id_card[6:10]
        month = id_card[10:12]
        day = id_card[12:14]
        return f"{year}-{month}-{day}"
    except:
        return None


def extract_gender_from_id_card(id_card: str) -> Optional[str]:
    """从身份证号提取性别（奇数男，偶数女）"""
    if not id_card or len(id_card) != 18:
        return None
    try:
        gender_code = int(id_card[16])
        return "男" if gender_code % 2 == 1 else "女"
    except:
        return None


def calculate_work_years(work_start_date: str) -> int:
    """计算工作年限"""
    if not work_start_date:
        return 0
    try:
        start = datetime.strptime(work_start_date, "%Y-%m-%d")
        today = datetime.now()
        years = today.year - start.year
        if (today.month, today.day) < (start.month, start.day):
            years -= 1
        return years
    except:
        return 0


@router.post("/collect/{teacher_id}")
async def collect_retirement_data(teacher_id: int):
    """
    汇集退休呈报表数据到中间表
    从教师基础信息表、学历表等多表查询数据
    """
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # 1. 查询教师基础信息
        cursor.execute("""
            SELECT id, name, id_card, birth_date, ethnicity, 
                   native_place, work_start_date, archive_birth_date,
                   employment_status, contact_phone
            FROM teacher_basic_info 
            WHERE id = %s
        """, (teacher_id,))
        
        teacher_row = cursor.fetchone()
        if not teacher_row:
            raise HTTPException(status_code=404, detail="教师不存在")
        
        # 2. 查询最高学历信息
        cursor.execute("""
            SELECT education_level, graduation_date, school_name, major
            FROM teacher_education_record
            WHERE teacher_id = %s
            ORDER BY graduation_date DESC
            LIMIT 1
        """, (teacher_id,))
        
        education_row = cursor.fetchone()
        
        # 3. 处理数据
        id_card = teacher_row[2] or ''
        
        # 出生日期：优先档案出生日期，否则从身份证提取
        birth_date = teacher_row[7] or extract_birth_date_from_id_card(id_card)
        
        # 性别：从身份证提取
        gender = extract_gender_from_id_card(id_card)
        
        # 工作年限
        work_years = calculate_work_years(teacher_row[6])
        
        # 4. 插入或更新中间表
        cursor.execute("""
            INSERT INTO retirement_report_data (
                teacher_id, name, id_card, gender, birth_date, 
                ethnicity, education, work_start_date, work_years,
                native_place, current_address, status, created_at, updated_at
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(), NOW())
            ON CONFLICT (teacher_id) DO UPDATE SET
                name = EXCLUDED.name,
                id_card = EXCLUDED.id_card,
                gender = EXCLUDED.gender,
                birth_date = EXCLUDED.birth_date,
                ethnicity = EXCLUDED.ethnicity,
                education = EXCLUDED.education,
                work_start_date = EXCLUDED.work_start_date,
                work_years = EXCLUDED.work_years,
                native_place = EXCLUDED.native_place,
                current_address = EXCLUDED.current_address,
                status = 'pending',
                updated_at = NOW()
        """, (
            teacher_id,
            teacher_row[1],  # name
            id_card,
            gender,
            birth_date,
            teacher_row[4],  # ethnicity
            get_education_name(education_row[0], DATABASE_CONFIG) if education_row else None,  # education - 转换为中文
            teacher_row[6],  # work_start_date
            work_years,
            teacher_row[5],  # native_place
            None,  # current_address
            'pending'
        ))
        
        conn.commit()
        
        return {
            "status": "success",
            "message": "退休呈报表数据已汇集到中间表",
            "data": {
                "teacher_id": teacher_id,
                "name": teacher_row[1],
                "gender": gender,
                "birth_date": birth_date,
                "work_years": work_years
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"数据汇集失败: {str(e)}")
    finally:
        cursor.close()
        conn.close()


@router.get("/list")
async def get_retirement_data_list(
    status: Optional[str] = Query(None, description="状态筛选"),
    keyword: Optional[str] = Query(None, description="关键词搜索")
):
    """获取退休呈报表数据列表"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        sql = """
            SELECT 
                r.id, r.teacher_id, r.name, r.gender, r.id_card,
                r.birth_date, r.ethnicity, r.work_years, r.status,
                r.created_at, r.updated_at
            FROM retirement_report_data r
            WHERE 1=1
        """
        params = []
        
        if status:
            sql += " AND r.status = %s"
            params.append(status)
        
        if keyword:
            sql += " AND (r.name LIKE %s OR r.id_card LIKE %s)"
            params.extend([f"%{keyword}%", f"%{keyword}%"])
        
        sql += " ORDER BY r.created_at DESC"
        
        cursor.execute(sql, params)
        rows = cursor.fetchall()
        
        result = []
        for row in rows:
            result.append({
                "id": row[0],
                "teacher_id": row[1],
                "name": row[2],
                "gender": row[3],
                "id_card": row[4],
                "birth_date": str(row[5]) if row[5] else None,
                "ethnicity": row[6],
                "work_years": row[7],
                "status": row[8],
                "created_at": str(row[9]) if row[9] else None,
                "updated_at": str(row[10]) if row[10] else None
            })
        
        return {
            "status": "success",
            "data": result,
            "total": len(result)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"查询失败: {str(e)}")
    finally:
        cursor.close()
        conn.close()


@router.get("/detail/{teacher_id}")
async def get_retirement_data_detail(teacher_id: int):
    """获取单个教师的退休呈报表详细数据"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            SELECT * FROM retirement_report_data 
            WHERE teacher_id = %s
        """, (teacher_id,))
        
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="数据不存在")
        
        # 获取列名
        columns = [desc[0] for desc in cursor.description]
        
        # 构建字典
        result = {}
        for i, col in enumerate(columns):
            value = row[i]
            if isinstance(value, datetime):
                value = str(value)
            result[col] = value
        
        return {
            "status": "success",
            "data": result
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"查询失败: {str(e)}")
    finally:
        cursor.close()
        conn.close()


@router.put("/update/{teacher_id}")
async def update_retirement_data(teacher_id: int, data: Dict[str, Any]):
    """更新退休呈报表数据"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # 构建更新SQL - 包含所有模板字段
        allowed_fields = [
            'name', 'gender', 'id_card', 'birth_date', 'ethnicity', 'education',
            'work_start_date', 'work_years', 'native_place', 'current_address',
            'job_title', 'tech_title', 'position',
            'retirement_reason', 'retirement_address', 'retirement_date', 'retirement_type',
            'post_20140930_specialty', 'post_20140930_specialty_level', 'post_20140930_specialty_salary',
            'post_20140930_worker', 'post_20140930_worker_level', 'post_20140930_worker_salary',
            'last_promotion_date', 'last_position', 'last_position_level', 'last_position_salary',
            'last_worker', 'last_worker_level', 'last_worker_salary',
            'retirement_specialty', 'retirement_specialty_level', 'retirement_specialty_salary',
            'retirement_worker', 'retirement_worker_level', 'retirement_worker_salary',
            'work_history', 'family_support', 'pension_unit',
            'unit_opinion', 'unit_opinion_date', 'dept_opinion', 'dept_opinion_date',
            'approval_opinion', 'approval_date', 'status'
        ]
        
        updates = []
        params = []
        
        for field in allowed_fields:
            if field in data:
                updates.append(f"{field} = %s")
                params.append(data[field])
        
        if not updates:
            raise HTTPException(status_code=400, detail="没有要更新的字段")
        
        params.append(teacher_id)
        
        sql = f"""
            UPDATE retirement_report_data 
            SET {', '.join(updates)}, updated_at = NOW()
            WHERE teacher_id = %s
        """
        
        cursor.execute(sql, params)
        conn.commit()
        
        return {
            "status": "success",
            "message": "数据已更新"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"更新失败: {str(e)}")
    finally:
        cursor.close()
        conn.close()


@router.post("/calculate-retirement/{teacher_id}")
async def calculate_retirement(
    teacher_id: int,
    custom_retirement_date: Optional[str] = None
):
    """
    计算退休日期和工作年限
    
    Args:
        teacher_id: 教师ID
        custom_retirement_date: 自定义退休日期（可选，格式：YYYY-MM-DD）
    
    Returns:
        退休计算信息
    """
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # 获取教师数据
        cursor.execute("""
            SELECT birth_date, gender, "个人身份", "参加工作时间"
            FROM retirement_report_data
            WHERE teacher_id = %s
        """, (teacher_id,))
        
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="教师数据不存在")
        
        birth_date, gender, personal_identity, work_start_date = row
        
        if not birth_date or not gender or not work_start_date:
            raise HTTPException(status_code=400, detail="缺少必要的计算参数（出生日期、性别、参加工作时间）")
        
        # 解析日期
        birth_date_obj = birth_date if isinstance(birth_date, date) else datetime.strptime(birth_date, "%Y-%m-%d").date()
        work_start_date_obj = work_start_date if isinstance(work_start_date, date) else datetime.strptime(work_start_date, "%Y-%m-%d").date()
        
        # 解析自定义退休日期
        custom_date_obj = None
        if custom_retirement_date:
            custom_date_obj = datetime.strptime(custom_retirement_date, "%Y-%m-%d").date()
        
        # 计算退休信息
        result = calculate_retirement_info(
            birth_date=birth_date_obj,
            gender=gender,
            personal_identity=personal_identity or "干部",
            work_start_date=work_start_date_obj,
            custom_retirement_date=custom_date_obj
        )
        
        # 转换为可序列化的格式
        return {
            "status": "success",
            "data": {
                "birth_date": str(result["birth_date"]),
                "gender": result["gender"],
                "personal_identity": result["personal_identity"],
                "work_start_date": str(result["work_start_date"]),
                "original_retirement_date": str(result["original_retirement_date"]),
                "delay_months": result["delay_months"],
                "calculated_retirement_date": str(result["calculated_retirement_date"]),
                "actual_retirement_date": str(result["actual_retirement_date"]),
                "work_years": result["work_years"]
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"计算失败: {str(e)}")
    finally:
        cursor.close()
        conn.close()


@router.post("/save-retirement-calculation/{teacher_id}")
async def save_retirement_calculation(
    teacher_id: int,
    data: dict
):
    """
    保存退休计算结果到中间表
    
    Args:
        teacher_id: 教师ID
        data: 包含 retirement_date 和 work_years 的字典
    
    Returns:
        保存结果
    """
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        retirement_date = data.get("retirement_date")
        work_years = data.get("work_years")
        
        if not retirement_date or work_years is None:
            raise HTTPException(status_code=400, detail="缺少退休日期或工作年限")
        
        cursor.execute("""
            UPDATE retirement_report_data 
            SET "退休时间" = %s, "工作年限" = %s, updated_at = NOW()
            WHERE teacher_id = %s
        """, (retirement_date, work_years, teacher_id))
        
        conn.commit()
        
        return {
            "status": "success",
            "message": "退休计算结果已保存"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"保存失败: {str(e)}")
    finally:
        cursor.close()
        conn.close()


@router.get("/generate-report/{teacher_id}")
async def generate_report(
    teacher_id: int,
    format: str = Query("html", description="报表格式: html, excel, word, pdf")
):
    """生成退休呈报表"""
    try:
        # 获取教师数据
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM retirement_report_data WHERE teacher_id = %s", (teacher_id,))
        row = cursor.fetchone()
        
        if not row:
            raise HTTPException(status_code=404, detail="数据不存在")
        
        # 获取列名和数据
        columns = [desc[0] for desc in cursor.description]
        data = dict(zip(columns, row))
        
        cursor.close()
        conn.close()
        
        # 根据格式生成报表
        if format == 'html':
            return generate_html_report(data)
        elif format == 'excel':
            return generate_excel_report(data)
        elif format == 'word':
            return generate_word_report(data)
        elif format == 'pdf':
            return generate_pdf_report(data)
        else:
            raise HTTPException(status_code=400, detail="不支持的格式")
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成报表失败: {str(e)}")


def generate_html_report(data: dict):
    """生成HTML格式报表"""
    from fastapi.responses import HTMLResponse
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>职工退休呈报表</title>
        <style>
            body {{ font-family: SimSun, serif; font-size: 14px; }}
            table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
            td, th {{ border: 1px solid #000; padding: 10px; text-align: left; }}
            .title {{ text-align: center; font-size: 22px; font-weight: bold; margin: 20px 0; }}
            .header {{ background-color: #f0f0f0; font-weight: bold; }}
        </style>
    </head>
    <body>
        <div class="title">职工退休呈报表</div>
        <table>
            <tr>
                <td class="header">姓名</td>
                <td>{data.get('name', '')}</td>
                <td class="header">性别</td>
                <td>{data.get('gender', '')}</td>
                <td class="header">身份证号</td>
                <td>{data.get('id_card', '')}</td>
            </tr>
            <tr>
                <td class="header">出生日期</td>
                <td>{data.get('birth_date', '')}</td>
                <td class="header">民族</td>
                <td>{data.get('ethnicity', '')}</td>
                <td class="header">文化程度</td>
                <td>{data.get('education', '')}</td>
            </tr>
            <tr>
                <td class="header">参加工作时间</td>
                <td>{data.get('work_start_date', '')}</td>
                <td class="header">工作年限</td>
                <td>{data.get('work_years', '')}</td>
                <td class="header">籍贯</td>
                <td>{data.get('native_place', '')}</td>
            </tr>
            <tr>
                <td class="header">现住址</td>
                <td colspan="5">{data.get('current_address', '')}</td>
            </tr>
            <tr>
                <td class="header">退休后居住地址</td>
                <td colspan="5">{data.get('retirement_address', '')}</td>
            </tr>
            <tr>
                <td class="header">退休原因</td>
                <td colspan="5">{data.get('retirement_reason', '')}</td>
            </tr>
            <tr>
                <td class="header">单位意见</td>
                <td colspan="5">{data.get('unit_opinion', '')}</td>
            </tr>
        </table>
    </body>
    </html>
    """
    
    return HTMLResponse(content=html_content)


def generate_excel_report(data: dict):
    """生成Excel格式报表"""
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side
    from fastapi.responses import StreamingResponse
    import io
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "职工退休呈报表"
    
    # 标题
    ws.merge_cells('A1:F1')
    ws['A1'] = '职工退休呈报表'
    ws['A1'].font = Font(size=16, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    
    # 数据行
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # 数据
    row_data = [
        ['姓名', data.get('name', ''), '性别', data.get('gender', ''), '身份证号', data.get('id_card', '')],
        ['出生日期', data.get('birth_date', ''), '民族', data.get('ethnicity', ''), '文化程度', data.get('education', '')],
        ['参加工作时间', data.get('work_start_date', ''), '工作年限', data.get('work_years', ''), '籍贯', data.get('native_place', '')],
        ['现住址', data.get('current_address', ''), '', '', '', ''],
        ['退休后居住地址', data.get('retirement_address', ''), '', '', '', ''],
        ['退休原因', data.get('retirement_reason', ''), '', '', '', ''],
        ['单位意见', data.get('unit_opinion', ''), '', '', '', '']
    ]
    
    for i, row in enumerate(row_data, start=3):
        for j, value in enumerate(row, start=1):
            cell = ws.cell(row=i, column=j, value=value)
            cell.border = border
            cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
    
    # 合并单元格
    for row in [6, 7, 8, 9]:
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
    
    # 保存到内存
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="职工退休呈报表_{data.get("name", "" )}.xlsx"'}
    )


def generate_word_report(data: dict):
    """生成Word格式报表"""
    from docx import Document
    from fastapi.responses import StreamingResponse
    import io
    
    doc = Document()
    
    # 标题
    title = doc.add_heading('职工退休呈报表', 0)
    title.alignment = 1  # 居中
    
    # 添加表格
    table = doc.add_table(rows=7, cols=6)
    table.style = 'Table Grid'
    
    # 填充数据
    cells_data = [
        ['姓名', data.get('name', ''), '性别', data.get('gender', ''), '身份证号', data.get('id_card', '')],
        ['出生日期', data.get('birth_date', ''), '民族', data.get('ethnicity', ''), '文化程度', data.get('education', '')],
        ['参加工作时间', data.get('work_start_date', ''), '工作年限', data.get('work_years', ''), '籍贯', data.get('native_place', '')],
        ['现住址', data.get('current_address', ''), '', '', '', ''],
        ['退休后居住地址', data.get('retirement_address', ''), '', '', '', ''],
        ['退休原因', data.get('retirement_reason', ''), '', '', '', ''],
        ['单位意见', data.get('unit_opinion', ''), '', '', '', '']
    ]
    
    for i, row_data in enumerate(cells_data):
        row = table.rows[i]
        for j, value in enumerate(row_data):
            row.cells[j].text = str(value) if value else ''
    
    # 合并单元格
    for row_idx in [3, 4, 5, 6]:
        for col_idx in range(1, 6):
            table.rows[row_idx].cells[col_idx].merge(table.rows[row_idx].cells[1])
    
    # 保存到内存
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="职工退休呈报表_{data.get("name", "" )}.docx"'}
    )


def generate_pdf_report(data: dict):
    """生成PDF格式报表"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from fastapi.responses import StreamingResponse
    import io
    
    # 注册中文字体
    try:
        pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))
    except:
        pass
    
    # 创建PDF
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4)
    
    # 样式
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName='SimSun',
        fontSize=18,
        alignment=1,  # 居中
        spaceAfter=30
    )
    
    # 内容
    elements = []
    
    # 标题
    elements.append(Paragraph('职工退休呈报表', title_style))
    elements.append(Spacer(1, 20))
    
    # 表格数据
    table_data = [
        ['姓名', data.get('name', ''), '性别', data.get('gender', ''), '身份证号', data.get('id_card', '')],
        ['出生日期', data.get('birth_date', ''), '民族', data.get('ethnicity', ''), '文化程度', data.get('education', '')],
        ['参加工作时间', data.get('work_start_date', ''), '工作年限', data.get('work_years', ''), '籍贯', data.get('native_place', '')],
        ['现住址', data.get('current_address', ''), '', '', '', ''],
        ['退休后居住地址', data.get('retirement_address', ''), '', '', '', ''],
        ['退休原因', data.get('retirement_reason', ''), '', '', '', ''],
        ['单位意见', data.get('unit_opinion', ''), '', '', '', '']
    ]
    
    # 创建表格
    table = Table(table_data, colWidths=[80, 100, 80, 100, 80, 100])
    
    # 表格样式
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, -1), 'SimSun'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BOX', (0, 0), (-1, -1), 1, colors.black),
        ('SPAN', (1, 3), (5, 3)),  # 现住址合并
        ('SPAN', (1, 4), (5, 4)),  # 退休后居住地址合并
        ('SPAN', (1, 5), (5, 5)),  # 退休原因合并
        ('SPAN', (1, 6), (5, 6)),  # 单位意见合并
    ]))
    
    elements.append(table)
    
    # 生成PDF
    doc.build(elements)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type='application/pdf',
        headers={'Content-Disposition': f'attachment; filename="职工退休呈报表_{data.get("name", "" )}.pdf"'}
    )
