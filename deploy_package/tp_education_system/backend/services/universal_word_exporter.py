"""
通用Word导出器 - 100%保留模板格式
核心原则：只替换占位符，不做任何其他修改
"""
import shutil
from docx import Document


def export_word(template_path: str, output_path: str, data: dict) -> str:
    """
    导出Word文件 - 100%保留模板格式
    
    Args:
        template_path: 模板文件路径
        output_path: 输出文件路径
        data: 填充数据 {字段名: 值}
    
    Returns:
        输出文件路径
    """
    # 1. 复制模板文件（保留所有格式）
    shutil.copy(template_path, output_path)
    
    # 2. 打开文档
    doc = Document(output_path)
    
    # 3. 替换所有 {{占位符}}
    # 3.1 替换段落中的占位符
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, data)
    
    # 3.2 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, data)
    
    # 4. 保存（保留所有页面设置）
    doc.save(output_path)
    
    return output_path


def _replace_in_paragraph(paragraph, data):
    """
    在段落中替换占位符
    只修改包含占位符的run，保留所有格式
    规则：有值则替换，无值（None或空字符串）则保留原占位符
    """
    # 检查段落中是否包含任何占位符
    paragraph_text = paragraph.text
    has_placeholder = False
    for key in data.keys():
        if f'{{{{{key}}}}}' in paragraph_text:
            has_placeholder = True
            break
    
    if not has_placeholder:
        return
    
    # 遍历所有runs，只替换包含占位符的run
    for run in paragraph.runs:
        run_text = run.text
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in run_text:
                # 判断是否有有效值
                if value is not None and str(value).strip() != '':
                    # 有值，进行替换，保留所有格式属性
                    run.text = run_text.replace(placeholder, str(value))
                # 无值时保留原占位符，不做任何替换


# 保持向后兼容
export_report = export_word
