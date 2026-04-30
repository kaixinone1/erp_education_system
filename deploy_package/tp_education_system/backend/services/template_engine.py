"""
模板填充引擎 - 驱动数据注入与格式保真导出
核心服务：读取模板、填充数据、生成文档
"""
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, Any, List, Optional
import json
import os
import shutil
from datetime import datetime


class TemplateEngine:
    """模板填充引擎"""
    
    def __init__(self, db_connection_func):
        """
        初始化模板引擎
        :param db_connection_func: 数据库连接函数
        """
        self.get_db_connection = db_connection_func
    
    def get_template_config(self, template_id: str) -> Dict[str, Any]:
        """
        获取模板配置
        :param template_id: 模板ID
        :return: 模板配置信息
        """
        conn = self.get_db_connection()
        cursor = conn.cursor()
        
        try:
            # 获取模板基本信息
            cursor.execute("""
                SELECT template_id, template_name, file_path, description
                FROM document_templates
                WHERE template_id = %s
            """, (template_id,))
            
            template_row = cursor.fetchone()
            if not template_row:
                raise ValueError(f"模板不存在: {template_id}")
            
            # 获取字段标记配置
            cursor.execute("""
                SELECT field_name, field_label, field_type, position_type, 
                       position_data, data_source, default_value
                FROM template_field_mappings
                WHERE template_id = %s
                ORDER BY sort_order
            """, (template_id,))
            
            field_rows = cursor.fetchall()
            fields = []
            for row in field_rows:
                fields.append({
                    'field_name': row[0],
                    'field_label': row[1],
                    'field_type': row[2],
                    'position_type': row[3],
                    'position_data': row[4] if isinstance(row[4], dict) else json.loads(row[4]),
                    'data_source': row[5],
                    'default_value': row[6]
                })
            
            return {
                'template_id': template_row[0],
                'template_name': template_row[1],
                'file_path': template_row[2],
                'description': template_row[3],
                'fields': fields
            }
            
        finally:
            cursor.close()
            conn.close()
    
    def get_data_by_source(self, data_source: str, business_id: int) -> str:
        """
        根据数据源获取数据
        :param data_source: 数据源，格式如 "teacher_basic_info.name"
        :param business_id: 业务对象ID
        :return: 数据值
        """
        if not data_source or '.' not in data_source:
            return ''
        
        table_name, column_name = data_source.split('.', 1)
        
        conn = self.get_db_connection()
        cursor = conn.cursor()
        
        try:
            # 构建查询SQL
            cursor.execute(f"""
                SELECT {column_name} FROM {table_name} WHERE id = %s
            """, (business_id,))
            
            result = cursor.fetchone()
            if result and result[0]:
                return str(result[0])
            return ''
            
        except Exception as e:
            print(f"获取数据失败 {data_source}: {e}")
            return ''
        finally:
            cursor.close()
            conn.close()
    
    def fill_template(self, template_id: str, business_id: int, 
                     override_data: Dict[str, str] = None) -> str:
        """
        填充模板
        :param template_id: 模板ID
        :param business_id: 业务对象ID（如教师ID）
        :param override_data: 覆盖数据（如表单填写的数据）
        :return: 生成的文件路径
        """
        # 获取模板配置
        config = self.get_template_config(template_id)
        template_path = config['file_path']
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        # 打开模板文档
        doc = Document(template_path)
        
        # 遍历所有字段并填充
        for field in config['fields']:
            field_name = field['field_name']
            data_source = field.get('data_source', '')
            default_value = field.get('default_value', '')
            position_type = field['position_type']
            position_data = field['position_data']
            
            # 获取字段值（优先级：覆盖数据 > 数据库 > 默认值）
            field_value = ''
            if override_data and field_name in override_data:
                field_value = override_data[field_name]
            elif data_source:
                field_value = self.get_data_by_source(data_source, business_id)
            
            if not field_value:
                field_value = default_value
            
            # 根据位置类型填充
            if position_type == 'table':
                self._fill_table_cell(doc, position_data, field_value)
            elif position_type == 'paragraph':
                self._fill_paragraph(doc, position_data, field_value)
            elif position_type == 'header':
                self._fill_header(doc, position_data, field_value)
            elif position_type == 'footer':
                self._fill_footer(doc, position_data, field_value)
        
        # 生成输出文件路径
        output_dir = r'd:\erp_thirteen\tp_education_system\backend\uploads\generated'
        os.makedirs(output_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"{template_id}_{business_id}_{timestamp}.docx"
        output_path = os.path.join(output_dir, output_filename)
        
        # 保存文档
        doc.save(output_path)
        
        # 记录使用记录
        self._record_usage(template_id, business_id, output_path, output_filename)
        
        return output_path
    
    def _fill_table_cell(self, doc: Document, position_data: Dict, value: str):
        """填充表格单元格"""
        table_idx = position_data.get('table_index', 0)
        row_idx = position_data.get('row_index', 0)
        cell_idx = position_data.get('cell_index', 0)
        
        if table_idx < len(doc.tables):
            table = doc.tables[table_idx]
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                if cell_idx < len(row.cells):
                    cell = row.cells[cell_idx]
                    # 保留第一个run的格式
                    if cell.paragraphs:
                        para = cell.paragraphs[0]
                        if para.runs:
                            para.runs[0].text = value
                        else:
                            para.text = value
    
    def _fill_paragraph(self, doc: Document, position_data: Dict, value: str):
        """填充段落"""
        para_idx = position_data.get('paragraph_index', 0)
        
        if para_idx < len(doc.paragraphs):
            para = doc.paragraphs[para_idx]
            # 保留第一个run的格式
            if para.runs:
                para.runs[0].text = value
            else:
                para.text = value
    
    def _fill_header(self, doc: Document, position_data: Dict, value: str):
        """填充页眉"""
        # 实现页眉填充逻辑
        pass
    
    def _fill_footer(self, doc: Document, position_data: Dict, value: str):
        """填充页脚"""
        # 实现页脚填充逻辑
        pass
    
    def _record_usage(self, template_id: str, business_id: int, 
                     file_path: str, file_name: str):
        """记录模板使用记录"""
        conn = self.get_db_connection()
        cursor = conn.cursor()
        
        try:
            cursor.execute("""
                INSERT INTO template_usage_records 
                (template_id, business_type, business_id, teacher_id, 
                 generated_file_path, generated_file_name)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (template_id, 'retirement_report', business_id, business_id, 
                  file_path, file_name))
            
            conn.commit()
        finally:
            cursor.close()
            conn.close()
    
    def preview_template(self, template_id: str, business_id: int) -> str:
        """
        预览模板（使用示例数据填充）
        :param template_id: 模板ID
        :param business_id: 业务对象ID
        :return: 预览文件路径
        """
        # 获取模板配置
        config = self.get_template_config(template_id)
        
        # 构建示例数据
        sample_data = {}
        for field in config['fields']:
            field_name = field['field_name']
            default_value = field.get('default_value', '')
            data_source = field.get('data_source', '')
            
            # 如果有数据源，尝试获取真实数据
            if data_source:
                value = self.get_data_by_source(data_source, business_id)
                if value:
                    sample_data[field_name] = value
                else:
                    sample_data[field_name] = default_value or f"[{field['field_label']}]"
            else:
                sample_data[field_name] = default_value or f"[{field['field_label']}]"
        
        # 使用示例数据填充
        return self.fill_template(template_id, business_id, sample_data)


# 便捷函数 - 这些函数在路由中被调用，database模块会在路由上下文中导入
def fill_template(template_id: str, business_id: int, 
                  override_data: Dict[str, str] = None) -> str:
    """
    填充模板的便捷函数
    """
    import sys
    import os
    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    from database import get_db_connection
    
    engine = TemplateEngine(get_db_connection)
    return engine.fill_template(template_id, business_id, override_data)


def preview_template(template_id: str, business_id: int) -> str:
    """
    预览模板的便捷函数
    """
    import sys
    import os
    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    from database import get_db_connection
    
    engine = TemplateEngine(get_db_connection)
    return engine.preview_template(template_id, business_id)
