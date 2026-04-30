"""
智能模板引擎 - 自动识别占位符并匹配数据源
"""
from docx import Document
from typing import Dict, Any, List, Tuple
import re
import os


class SmartTemplateEngine:
    """智能模板引擎 - 自动识别和填充"""

    # 字段映射表 - 占位符名称到数据库字段的映射
    FIELD_MAPPING = {
        # 教师基本信息
        '姓名': 'teacher_basic_info.name',
        '教师姓名': 'teacher_basic_info.name',
        'name': 'teacher_basic_info.name',
        '身份证号': 'teacher_basic_info.id_card',
        '身份证号码': 'teacher_basic_info.id_card',
        'id_card': 'teacher_basic_info.id_card',
        '性别': 'teacher_basic_info.gender',
        '男女': 'teacher_basic_info.gender',
        '出生日期': 'teacher_basic_info.archive_birth_date',
        '出生年月': 'teacher_basic_info.archive_birth_date',
        '民族': 'teacher_basic_info.ethnicity',
        '籍贯': 'teacher_basic_info.native_place',
        '参加工作时间': 'teacher_basic_info.work_start_date',
        '工作日期': 'teacher_basic_info.work_start_date',
        '联系电话': 'teacher_basic_info.contact_phone',
        '电话': 'teacher_basic_info.contact_phone',
        '手机号': 'teacher_basic_info.contact_phone',

        # 退休相关
        '退休日期': 'retirement_cert_records.retirement_date',
        '退休时间': 'retirement_cert_records.retirement_date',
        '应退休时间': 'retirement_cert_records.retirement_date',
        '申请日期': 'retirement_report_form.application_date',

        # 通用占位符
        '日期': 'current_date',
        '当前日期': 'current_date',
        'today': 'current_date',
        '年份': 'current_year',
        '年': 'current_year',
        '月': 'current_month',
        '日': 'current_day',
    }

    @staticmethod
    def extract_placeholders(file_path: str) -> List[Dict[str, Any]]:
        """
        从文档中提取占位符
        支持的格式: {{字段名}}、[[字段名]]、【字段名】、[字段名]、<字段名>
        """
        if not os.path.exists(file_path):
            return []

        doc = Document(file_path)
        placeholders = []
        placeholder_pattern = r'[\{\[【\<]([^\}\]】\>]+)[\}\]】\>]'

        # 扫描段落
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            matches = re.finditer(placeholder_pattern, text)
            for match in matches:
                placeholder_name = match.group(1).strip()
                placeholders.append({
                    'placeholder': placeholder_name,
                    'position_type': 'paragraph',
                    'position_data': {'paragraph_index': para_idx},
                    'matched_field': SmartTemplateEngine.FIELD_MAPPING.get(placeholder_name, '')
                })

        # 扫描表格
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text
                    matches = re.finditer(placeholder_pattern, text)
                    for match in matches:
                        placeholder_name = match.group(1).strip()
                        placeholders.append({
                            'placeholder': placeholder_name,
                            'position_type': 'table',
                            'position_data': {
                                'table_index': table_idx,
                                'row_index': row_idx,
                                'cell_index': cell_idx
                            },
                            'matched_field': SmartTemplateEngine.FIELD_MAPPING.get(placeholder_name, '')
                        })

        # 去重
        seen = set()
        unique_placeholders = []
        for p in placeholders:
            key = (p['placeholder'], p['position_type'], str(p['position_data']))
            if key not in seen:
                seen.add(key)
                unique_placeholders.append(p)

        return unique_placeholders

    @staticmethod
    def auto_generate_fields(template_id: str, file_path: str, db_connection_func) -> List[Dict[str, Any]]:
        """
        自动生成字段配置
        """
        placeholders = SmartTemplateEngine.extract_placeholders(file_path)

        fields = []
        for idx, p in enumerate(placeholders):
            field_name = p['placeholder'].replace(' ', '_').replace('-', '_')
            field_label = p['placeholder']

            fields.append({
                'field_name': field_name,
                'field_label': field_label,
                'field_type': 'text',
                'position_type': p['position_type'],
                'position_data': p['position_data'],
                'data_source': p['matched_field'],
                'default_value': '',
                'sort_order': idx
            })

        # 保存到数据库
        if fields:
            conn = db_connection_func()
            cursor = conn.cursor()

            # 清除旧配置
            cursor.execute("""
                DELETE FROM template_field_mappings WHERE template_id = %s
            """, (template_id,))

            # 插入新配置
            for field in fields:
                cursor.execute("""
                    INSERT INTO template_field_mappings
                    (template_id, field_name, field_label, field_type, position_type,
                     position_data, data_source, default_value, sort_order)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    template_id,
                    field['field_name'],
                    field['field_label'],
                    field['field_type'],
                    field['position_type'],
                    json.dumps(field['position_data']),
                    field['data_source'],
                    field['default_value'],
                    field['sort_order']
                ))

            conn.commit()
            cursor.close()
            conn.close()

        return fields

    @staticmethod
    def fill_document(template_path: str, data: Dict[str, Any], output_path: str):
        """
        填充文档 - 替换所有占位符
        """
        doc = Document(template_path)
        placeholder_pattern = r'[\{\[【\<]([^\}\]】\>]+)[\}\]】\>]'

        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if re.search(placeholder_pattern, run.text):
                    for key, value in data.items():
                        # 替换各种格式的占位符
                        patterns = [
                            f'{{{{{key}}}}}',
                            f'[[{key}]]',
                            f'【{key}】',
                            f'<{key}>',
                            f'[{key}]'
                        ]
                        for pattern in patterns:
                            run.text = run.text.replace(pattern, str(value))

        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if re.search(placeholder_pattern, run.text):
                                for key, value in data.items():
                                    patterns = [
                                        f'{{{{{key}}}}}',
                                        f'[[{key}]]',
                                        f'【{key}】',
                                        f'<{key}>',
                                        f'[{key}]'
                                    ]
                                    for pattern in patterns:
                                        run.text = run.text.replace(pattern, str(value))

        doc.save(output_path)
        return output_path


# 便捷函数
def analyze_template(template_id: str, file_path: str, db_connection_func) -> Dict[str, Any]:
    """
    分析模板，提取占位符并生成字段配置
    """
    engine = SmartTemplateEngine()
    placeholders = engine.extract_placeholders(file_path)
    fields = engine.auto_generate_fields(template_id, file_path, db_connection_func)

    return {
        'placeholders_found': len(placeholders),
        'fields_generated': len(fields),
        'fields': fields,
        'unmatched_fields': [p['placeholder'] for p in placeholders if not p['matched_field']]
    }
