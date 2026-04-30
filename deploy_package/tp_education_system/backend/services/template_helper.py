"""
模板辅助工具 - 帮助用户分析文档结构
"""
from docx import Document
from typing import Dict, Any, List
import os


def analyze_document_structure(file_path: str) -> Dict[str, Any]:
    """
    分析文档结构，返回详细的表格和段落信息
    """
    if not os.path.exists(file_path):
        return {'error': '文件不存在'}

    doc = Document(file_path)

    # 分析表格
    tables_info = []
    for idx, table in enumerate(doc.tables):
        table_data = {
            '序号': idx + 1,  # 从1开始显示
            '行数': len(table.rows),
            '列数': len(table.rows[0].cells) if table.rows else 0,
            '内容预览': []
        }

        # 获取前3行内容作为预览
        for row_idx, row in enumerate(table.rows[:3]):
            row_data = []
            for cell in row.cells:
                text = cell.text.strip()[:20]  # 限制长度
                row_data.append(text if text else '(空白)')
            table_data['内容预览'].append(f"第{row_idx + 1}行: {', '.join(row_data)}")

        tables_info.append(table_data)

    # 分析段落
    paragraphs_info = []
    for idx, para in enumerate(doc.paragraphs[:10]):  # 只显示前10个非空段落
        text = para.text.strip()
        if text:
            paragraphs_info.append({
                '序号': idx + 1,
                '内容': text[:50] + '...' if len(text) > 50 else text
            })

    return {
        '表格总数': len(doc.tables),
        '段落总数': len(doc.paragraphs),
        '表格详情': tables_info,
        '段落预览': paragraphs_info
    }


def print_document_structure(file_path: str):
    """
    打印文档结构信息
    """
    result = analyze_document_structure(file_path)

    print("=" * 60)
    print("文档结构分析")
    print("=" * 60)
    print(f"表格总数: {result['表格总数']}")
    print(f"段落总数: {result['段落总数']}")
    print()

    if result['表格详情']:
        print("表格详情:")
        print("-" * 60)
        for table in result['表格详情']:
            print(f"\n【表格 {table['序号']}】 - {table['行数']}行 x {table['列数']}列")
            for preview in table['内容预览']:
                print(f"  {preview}")

    print("\n" + "=" * 60)


# 便捷函数
def get_table_index_by_content(file_path: str, keyword: str) -> int:
    """
    根据内容关键词查找表格序号
    """
    doc = Document(file_path)

    for idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                if keyword in cell.text:
                    return idx + 1  # 返回从1开始的序号

    return -1  # 未找到


def find_cell_position(file_path: str, label_text: str) -> Dict[str, Any]:
    """
    查找包含特定标签的单元格位置
    """
    doc = Document(file_path)

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if label_text in cell.text:
                    return {
                        '表格序号': table_idx + 1,
                        '行号': row_idx + 1,
                        '列号': cell_idx + 1,
                        '单元格内容': cell.text.strip()
                    }

    return None
