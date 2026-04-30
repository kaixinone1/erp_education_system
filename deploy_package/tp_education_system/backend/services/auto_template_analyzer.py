"""
自动模板分析器 - 无需占位符，智能识别文档结构
"""
from docx import Document
from typing import Dict, Any, List, Tuple
import re
import os


class AutoTemplateAnalyzer:
    """自动模板分析器 - 分析文档结构，智能识别填充位置"""

    # 常见字段关键词映射
    FIELD_KEYWORDS = {
        # 教师基本信息字段
        '姓名': {'field': 'teacher_basic_info.name', 'type': 'text'},
        '教师姓名': {'field': 'teacher_basic_info.name', 'type': 'text'},
        '性别': {'field': 'teacher_basic_info.gender', 'type': 'text'},
        '男女': {'field': 'teacher_basic_info.gender', 'type': 'text'},
        '出生日期': {'field': 'teacher_basic_info.archive_birth_date', 'type': 'date'},
        '出生年月': {'field': 'teacher_basic_info.archive_birth_date', 'type': 'date'},
        '出生': {'field': 'teacher_basic_info.archive_birth_date', 'type': 'date'},
        '身份证号': {'field': 'teacher_basic_info.id_card', 'type': 'text'},
        '身份证号码': {'field': 'teacher_basic_info.id_card', 'type': 'text'},
        '身份证': {'field': 'teacher_basic_info.id_card', 'type': 'text'},
        '民族': {'field': 'teacher_basic_info.ethnicity', 'type': 'text'},
        '籍贯': {'field': 'teacher_basic_info.native_place', 'type': 'text'},
        '参加工作时间': {'field': 'teacher_basic_info.work_start_date', 'type': 'date'},
        '工作时间': {'field': 'teacher_basic_info.work_start_date', 'type': 'date'},
        '联系电话': {'field': 'teacher_basic_info.contact_phone', 'type': 'text'},
        '电话': {'field': 'teacher_basic_info.contact_phone', 'type': 'text'},
        '手机': {'field': 'teacher_basic_info.contact_phone', 'type': 'text'},

        # 退休相关
        '退休日期': {'field': 'retirement_cert_records.retirement_date', 'type': 'date'},
        '退休时间': {'field': 'retirement_cert_records.retirement_date', 'type': 'date'},
        '应退休时间': {'field': 'retirement_cert_records.retirement_date', 'type': 'date'},
        '申请日期': {'field': 'current_date', 'type': 'date'},
        '申请时间': {'field': 'current_date', 'type': 'date'},

        # 通用
        '日期': {'field': 'current_date', 'type': 'date'},
        '填表日期': {'field': 'current_date', 'type': 'date'},
    }

    @staticmethod
    def analyze_document(file_path: str) -> Dict[str, Any]:
        """
        分析文档结构，识别所有可能的填充位置
        """
        if not os.path.exists(file_path):
            return {'tables': [], 'paragraphs': []}

        doc = Document(file_path)

        # 分析表格
        tables = AutoTemplateAnalyzer._analyze_tables(doc)

        # 分析段落
        paragraphs = AutoTemplateAnalyzer._analyze_paragraphs(doc)

        return {
            'tables': tables,
            'paragraphs': paragraphs,
            'suggested_fields': AutoTemplateAnalyzer._suggest_fields(tables, paragraphs)
        }

    @staticmethod
    def _analyze_tables(doc: Document) -> List[Dict]:
        """分析所有表格结构"""
        tables = []

        for table_idx, table in enumerate(doc.tables):
            table_info = {
                'index': table_idx,
                'rows': len(table.rows),
                'cols': len(table.rows[0].cells) if table.rows else 0,
                'cells': []
            }

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()

                    # 判断单元格类型
                    cell_type = 'empty'
                    if text:
                        # 检查是否是标签（有冒号、是关键词等）
                        if any(keyword in text for keyword in AutoTemplateAnalyzer.FIELD_KEYWORDS.keys()):
                            cell_type = 'label'
                        elif re.match(r'^[\u4e00-\u9fa5]+[:：]', text):
                            cell_type = 'label'
                        else:
                            cell_type = 'content'

                    table_info['cells'].append({
                        'row': row_idx,
                        'col': cell_idx,
                        'text': text,
                        'type': cell_type
                    })

            tables.append(table_info)

        return tables

    @staticmethod
    def _analyze_paragraphs(doc: Document) -> List[Dict]:
        """分析所有段落"""
        paragraphs = []

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 识别段落类型
            para_type = 'normal'
            if re.match(r'^[\u4e00-\u9fa5]+[:：]', text):
                para_type = 'label'
            elif any(keyword in text for keyword in AutoTemplateAnalyzer.FIELD_KEYWORDS.keys()):
                para_type = 'potential_field'

            paragraphs.append({
                'index': para_idx,
                'text': text,
                'type': para_type
            })

        return paragraphs

    @staticmethod
    def _suggest_fields(tables: List, paragraphs: List) -> List[Dict]:
        """根据分析结果建议字段"""
        suggested_fields = []

        # 从表格中识别字段
        for table in tables:
            for cell in table['cells']:
                if cell['type'] == 'label':
                    text = cell['text']

                    # 查找匹配的关键词
                    for keyword, config in AutoTemplateAnalyzer.FIELD_KEYWORDS.items():
                        if keyword in text:
                            # 找到标签，检查相邻单元格
                            next_cell = AutoTemplateAnalyzer._find_next_empty_cell(
                                table['cells'], cell['row'], cell['col']
                            )

                            if next_cell:
                                suggested_fields.append({
                                    'field_label': keyword,
                                    'field_name': config['field'].replace('.', '_'),
                                    'data_source': config['field'],
                                    'field_type': config['type'],
                                    'position_type': 'table',
                                    'position_data': {
                                        'table_index': table['index'],
                                        'row_index': next_cell['row'],
                                        'cell_index': next_cell['col']
                                    },
                                    'confidence': 'high' if keyword == text.replace(':', '').replace('：', '') else 'medium'
                                })
                            break

        return suggested_fields

    @staticmethod
    def _find_next_empty_cell(cells: List, row: int, col: int) -> Dict:
        """找到标签旁边的空白单元格"""
        # 先找同一行的下一个单元格
        for cell in cells:
            if cell['row'] == row and cell['col'] == col + 1:
                if cell['type'] == 'empty' or not cell['text']:
                    return cell

        # 再找同一列的下一个单元格
        for cell in cells:
            if cell['row'] == row + 1 and cell['col'] == col:
                if cell['type'] == 'empty' or not cell['text']:
                    return cell

        return None

    @staticmethod
    def auto_fill_template(template_path: str, data: Dict[str, Any], output_path: str):
        """
        自动填充模板 - 根据识别的位置填充数据
        """
        doc = Document(template_path)
        analysis = AutoTemplateAnalyzer.analyze_document(template_path)

        # 填充表格
        for field in analysis['suggested_fields']:
            if field['position_type'] == 'table':
                pos = field['position_data']
                table_idx = pos['table_index']
                row_idx = pos['row_index']
                cell_idx = pos['cell_index']

                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    if row_idx < len(table.rows):
                        row = table.rows[row_idx]
                        if cell_idx < len(row.cells):
                            cell = row.cells[cell_idx]

                            # 获取数据值
                            data_key = field['data_source'].split('.')[-1]
                            value = data.get(data_key, '')
                            cell.text = str(value) if value else ''

        doc.save(output_path)
        return output_path


# 便捷函数
def analyze_and_generate_fields(template_id: str, file_path: str, db_connection_func) -> Dict[str, Any]:
    """
    分析模板并自动生成字段配置
    """
    import json

    analyzer = AutoTemplateAnalyzer()
    analysis = analyzer.analyze_document(file_path)

    fields = analysis['suggested_fields']

    # 保存到数据库
    if fields:
        conn = db_connection_func()
        cursor = conn.cursor()

        # 清除旧配置
        cursor.execute("""
            DELETE FROM template_field_mappings WHERE template_id = %s
        """, (template_id,))

        # 插入新配置
        for idx, field in enumerate(fields):
            cursor.execute("""
                INSERT INTO template_field_mappings
                (template_id, field_name, field_label, field_type, position_type,
                 position_data, data_source, default_value, sort_order)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                template_id,
                field['field_name'],
                field['field_label'],
                field['field_type'],
                field['position_type'],
                json.dumps(field['position_data']),
                field['data_source'],
                '',
                idx
            ))

        conn.commit()
        cursor.close()
        conn.close()

    return {
        'tables_found': len(analysis['tables']),
        'paragraphs_found': len(analysis['paragraphs']),
        'fields_generated': len(fields),
        'fields': fields,
        'document_structure': analysis
    }
