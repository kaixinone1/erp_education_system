"""
通用占位符提取服务
支持多种文件格式和占位符格式
"""
import re
import json
import os
from typing import List, Dict, Any, Optional
from bs4 import BeautifulSoup

# 加载配置
CONFIG_FILE = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'config', 'template_config.json')


def load_config() -> Dict[str, Any]:
    """加载模板配置"""
    try:
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"加载配置失败: {e}")
        return {
            "placeholder_formats": [{"pattern": r"\{([^{}]+)\}", "name": "curly_braces"}],
            "excluded_patterns": ["mso-", "style", "font-"],
            "system_fields": {}
        }


def is_valid_placeholder(name: str, config: Dict[str, Any]) -> bool:
    """
    验证占位符名称是否有效
    排除CSS代码和其他无效内容
    """
    name = name.strip()
    
    # 排除空字符串
    if not name:
        return False
    
    # 排除纯数字
    if name.isdigit():
        return False
    
    # 排除包含特殊字符的
    if ':' in name or ';' in name or '{' in name or '}' in name:
        return False
    
    # 排除太长的（CSS代码通常很长）
    if len(name) > 50:
        return False
    
    # 排除CSS相关内容
    excluded = config.get("excluded_patterns", [])
    for pattern in excluded:
        if pattern in name.lower():
            return False
    
    # 只保留中文、英文、数字、下划线、空格
    if not re.match(r'^[\u4e00-\u9fa5a-zA-Z0-9_\s]+$', name):
        return False
    
    return True


def extract_placeholders_from_text(text: str, config: Optional[Dict[str, Any]] = None) -> List[str]:
    """
    从文本中提取所有占位符
    支持多种占位符格式
    """
    if config is None:
        config = load_config()
    
    placeholders = []
    seen = set()
    
    # 遍历所有配置的占位符格式
    for format_config in config.get("placeholder_formats", []):
        pattern = format_config.get("pattern", r"\{([^{}]+)\}")
        try:
            matches = re.findall(pattern, text)
            for match in matches:
                if is_valid_placeholder(match, config) and match not in seen:
                    seen.add(match)
                    placeholders.append(match)
        except re.error as e:
            print(f"正则表达式错误: {pattern}, {e}")
            continue
    
    return placeholders


def extract_from_html(file_path: str) -> List[Dict[str, Any]]:
    """从HTML文件中提取占位符"""
    config = load_config()
    fields = []
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 方法1: 从HTML标签文本中提取
        soup = BeautifulSoup(content, 'html.parser')
        
        # 提取所有文本节点中的占位符
        for idx, element in enumerate(soup.find_all(text=True)):
            text = str(element)
            placeholders = extract_placeholders_from_text(text, config)
            for ph in placeholders:
                fields.append({
                    'name': ph,
                    'label': ph,
                    'field_type': 'text',
                    'position_data': {'type': 'text', 'index': idx},
                    'source': 'html_text'
                })
        
        # 方法2: 从input元素中提取
        for idx, inp in enumerate(soup.find_all('input')):
            field_name = inp.get('name') or inp.get('id')
            if field_name:
                fields.append({
                    'name': field_name,
                    'label': field_name,
                    'field_type': inp.get('type', 'text'),
                    'position_data': {'type': 'input', 'index': idx},
                    'source': 'html_input'
                })
        
        # 方法3: 直接从原始内容中提取（处理CSS样式中的误匹配）
        raw_placeholders = extract_placeholders_from_text(content, config)
        existing_names = {f['name'] for f in fields}
        for ph in raw_placeholders:
            if ph not in existing_names:
                fields.append({
                    'name': ph,
                    'label': ph,
                    'field_type': 'text',
                    'position_data': {'type': 'raw', 'index': 0},
                    'source': 'raw_content'
                })
        
    except Exception as e:
        print(f"提取HTML字段失败: {e}")
    
    return fields


def extract_from_word(file_path: str) -> List[Dict[str, Any]]:
    """从Word文件中提取占位符"""
    config = load_config()
    fields = []
    
    try:
        from docx import Document
        doc = Document(file_path)
        
        # 提取所有段落中的占位符
        for idx, para in enumerate(doc.paragraphs):
            text = para.text
            placeholders = extract_placeholders_from_text(text, config)
            for ph in placeholders:
                fields.append({
                    'name': ph,
                    'label': ph,
                    'field_type': 'text',
                    'position_data': {'type': 'paragraph', 'index': idx},
                    'source': 'word_paragraph'
                })
        
        # 提取表格中的占位符
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text
                    placeholders = extract_placeholders_from_text(text, config)
                    for ph in placeholders:
                        # 检查是否已存在
                        existing = next((f for f in fields if f['name'] == ph), None)
                        if not existing:
                            fields.append({
                                'name': ph,
                                'label': ph,
                                'field_type': 'text',
                                'position_data': {
                                    'type': 'table_cell',
                                    'table': table_idx,
                                    'row': row_idx,
                                    'cell': cell_idx
                                },
                                'source': 'word_table'
                            })
        
    except Exception as e:
        print(f"提取Word字段失败: {e}")
    
    return fields


def extract_from_excel(file_path: str) -> List[Dict[str, Any]]:
    """从Excel文件中提取占位符"""
    config = load_config()
    fields = []
    
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        
        for sheet_idx, sheet in enumerate(wb.worksheets):
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value:
                        text = str(cell.value)
                        placeholders = extract_placeholders_from_text(text, config)
                        for ph in placeholders:
                            # 检查是否已存在
                            existing = next((f for f in fields if f['name'] == ph), None)
                            if not existing:
                                fields.append({
                                    'name': ph,
                                    'label': ph,
                                    'field_type': 'text',
                                    'position_data': {
                                        'type': 'excel_cell',
                                        'sheet': sheet_idx,
                                        'row': cell.row,
                                        'column': cell.column
                                    },
                                    'source': 'excel_cell'
                                })
        
    except Exception as e:
        print(f"提取Excel字段失败: {e}")
    
    return fields


def extract_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    """从PDF文件中提取占位符"""
    config = load_config()
    fields = []
    
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                placeholders = extract_placeholders_from_text(text, config)
                for ph in placeholders:
                    fields.append({
                        'name': ph,
                        'label': ph,
                        'field_type': 'text',
                        'position_data': {'type': 'pdf_page', 'page': page_idx},
                        'source': 'pdf_text'
                    })
        
    except Exception as e:
        print(f"提取PDF字段失败: {e}")
    
    return fields


def extract_fields(file_path: str, file_ext: str) -> List[Dict[str, Any]]:
    """
    通用字段提取入口
    根据文件扩展名自动选择对应的提取器
    """
    config = load_config()
    file_types = config.get("file_types", {})
    
    # 获取文件类型配置
    type_config = file_types.get(file_ext.lower(), {})
    handler = type_config.get("handler", "")
    
    # 根据handler选择提取器
    extractors = {
        "html": extract_from_html,
        "word": extract_from_word,
        "excel": extract_from_excel,
        "pdf": extract_from_pdf
    }
    
    extractor = extractors.get(handler, extract_from_html)
    return extractor(file_path)


# 向后兼容的别名
extract_html_fields = extract_from_html
extract_word_fields = extract_from_word
extract_excel_fields = extract_from_excel
extract_pdf_fields = extract_from_pdf
