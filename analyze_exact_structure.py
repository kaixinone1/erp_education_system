#!/usr/bin/env python3
"""
精确分析职工退休呈报表Word模板结构
完全按照原表格式
"""
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

def analyze_exact_structure():
    template_path = r"D:\erp_thirteen\tp_education_system\backend\uploads\templates\20260208_201747_职工退休呈报表.docx"
    
    doc = Document(template_path)
    
    print("=" * 100)
    print("职工退休呈报表 - 精确结构分析")
    print("=" * 100)
    
    # 分析所有段落
    print("\n【所有段落】")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            alignment = para.alignment
            style = para.style.name if para.style else "None"
            print(f"  段落{i}: [{style}] 对齐={alignment}")
            print(f"    内容: {text[:100]}")
    
    # 详细分析每个表格
    print("\n【表格详细分析】")
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n{'='*80}")
        print(f"表格 {table_idx + 1}")
        print(f"{'='*80}")
        print(f"总行数: {len(table.rows)}")
        print(f"总列数: {len(table.columns)}")
        
        # 分析表格结构
        for row_idx, row in enumerate(table.rows):
            print(f"\n  第{row_idx + 1}行:")
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                # 获取单元格合并信息
                tc = cell._tc
                grid_span = tc.grid_span if hasattr(tc, 'grid_span') else 1
                
                print(f"    单元格({row_idx},{col_idx}): {text[:60]}")
                if grid_span > 1:
                    print(f"      -> 合并单元格，跨度={grid_span}")
    
    # 生成HTML结构建议
    print("\n" + "=" * 100)
    print("【HTML结构建议】")
    print("=" * 100)
    
    html_structure = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {
            font-family: 'SimSun', '宋体', serif;
            font-size: 14pt;
            line-height: 1.8;
            padding: 40px;
        }
        
        /* 标题 */
        .report-title {
            text-align: center;
            font-size: 26pt;
            font-weight: bold;
            letter-spacing: 12px;
            margin-bottom: 10px;
        }
        
        .report-subtitle {
            text-align: center;
            font-size: 12pt;
            color: #333;
            margin-bottom: 30px;
        }
        
        /* 表格样式 */
        table {
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 20px;
        }
        
        td, th {
            border: 1px solid #000;
            padding: 8px 12px;
            vertical-align: middle;
        }
        
        .label {
            background-color: #f5f5f5;
            font-weight: normal;
            text-align: center;
            width: 15%;
        }
        
        .value {
            text-align: left;
            width: 35%;
        }
        
        /* 审批意见区域 */
        .approval-table td {
            padding: 15px;
        }
        
        .approval-title {
            background-color: #f5f5f5;
            text-align: center;
            font-weight: bold;
            width: 15%;
        }
        
        .approval-content {
            text-align: left;
            line-height: 2;
        }
        
        .date-line {
            text-align: right;
            margin-top: 20px;
        }
        
        .date-line span {
            display: inline-block;
            min-width: 50px;
            border-bottom: 1px solid #000;
            text-align: center;
            margin: 0 5px;
        }
    </style>
</head>
<body>
    <!-- 标题 -->
    <div class="report-title">职 工 退 休 呈 报 表</div>
    <div class="report-subtitle">枣阳市人力资源和社会保障局制</div>
    
    <!-- 基本信息表格 -->
    <table>
        <tr>
            <td class="label">姓名</td>
            <td class="value">{{teacher_name}}</td>
            <td class="label">性别</td>
            <td class="value">{{gender}}</td>
            <td class="label">出生年月</td>
            <td class="value">{{birth_date}}</td>
        </tr>
        <tr>
            <td class="label">身份证号</td>
            <td class="value" colspan="3">{{id_card}}</td>
            <td class="label">民族</td>
            <td class="value">{{ethnicity}}</td>
        </tr>
        <tr>
            <td class="label">籍贯</td>
            <td class="value" colspan="5">{{native_place}}</td>
        </tr>
    </table>
    
    <!-- 教育信息 -->
    <table>
        <tr>
            <td class="label">文化程度</td>
            <td class="value">{{education}}</td>
            <td class="label">毕业学校</td>
            <td class="value">{{graduation_school}}</td>
        </tr>
        <tr>
            <td class="label">专业</td>
            <td class="value">{{major}}</td>
            <td class="label">学位</td>
            <td class="value">{{degree}}</td>
        </tr>
    </table>
    
    <!-- 工作信息 -->
    <table>
        <tr>
            <td class="label">参加工作时间</td>
            <td class="value">{{work_start_date}}</td>
            <td class="label">工作单位</td>
            <td class="value">{{work_unit}}</td>
        </tr>
        <tr>
            <td class="label">职务</td>
            <td class="value">{{position}}</td>
            <td class="label">职称</td>
            <td class="value">{{title}}</td>
        </tr>
    </table>
    
    <!-- 退休信息 -->
    <table>
        <tr>
            <td class="label">年龄</td>
            <td class="value">{{age}}</td>
            <td class="label">退休日期</td>
            <td class="value">{{retirement_date}}</td>
        </tr>
    </table>
    
    <!-- 审批意见 -->
    <table class="approval-table">
        <tr>
            <td class="approval-title">呈报单位意见</td>
            <td class="approval-content">
                <p>经研究，同意 {{teacher_name}} 同志按以下第（&nbsp;&nbsp;&nbsp;&nbsp;）条办理退休，从 {{retirement_date}} 执行。</p>
                <p>（一）弹性提前退休</p>
                <p>（二）法定退休年龄退休</p>
                <div class="date-line">
                    <span>{{year}}</span>年<span>{{month}}</span>月<span>{{day}}</span>日
                </div>
            </td>
        </tr>
        <tr>
            <td class="approval-title">主管部门审查意见</td>
            <td class="approval-content">
                <p>同意呈报</p>
                <div class="date-line">
                    <span>{{year}}</span>年<span>{{month}}</span>月<span>{{day}}</span>日
                </div>
            </td>
        </tr>
        <tr>
            <td class="approval-title">退休一次性补贴审批意见</td>
            <td class="approval-content">
                <p>根据鄂人社发【2017】8号文件规定，同意 {{teacher_name}} 同志发放一次性独生子女费 {{only_child_fee}} 元，教育特殊贡献奖 {{education_contribution_award}} 元，从 {{retirement_date}} 执行。</p>
                <div class="date-line">
                    <span>{{year}}</span>年<span>{{month}}</span>月<span>{{day}}</span>日
                </div>
            </td>
        </tr>
        <tr>
            <td class="approval-title">批准机关审批意见</td>
            <td class="approval-content">
                <p>根据人社部发【2024】94号文件规定，同意 {{teacher_name}} 同志按第（&nbsp;&nbsp;&nbsp;&nbsp;）条退休，从 {{retirement_date}} 执行。</p>
                <div class="date-line">
                    <span>{{year}}</span>年<span>{{month}}</span>月<span>{{day}}</span>日
                </div>
            </td>
        </tr>
    </table>
</body>
</html>
"""
    
    print(html_structure)
    
    print("\n" + "=" * 100)
    print("分析完成！")
    print("=" * 100)

if __name__ == '__main__':
    analyze_exact_structure()
