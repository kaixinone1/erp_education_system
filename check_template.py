#!/usr/bin/env python3
"""检查模板文件中的文本内容"""

from docx import Document

def check_template():
    template_path = r'd:\erp_thirteen\职工退休呈报表.docx'
    doc = Document(template_path)
    
    print("=== 模板文件中的所有段落 ===")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"段落 {i}: {repr(text)}")
    
    print("\n=== 模板文件中的所有表格单元格 ===")
    for table_idx, table in enumerate(doc.tables):
        print(f"\n表格 {table_idx}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    print(f"  行{row_idx} 列{cell_idx}: {repr(text[:100])}")

if __name__ == '__main__':
    check_template()
