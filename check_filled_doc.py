#!/usr/bin/env python3
"""检查填充后的文档"""
from docx import Document

def check():
    try:
        doc = Document('test_filled_report.docx')
        
        print("=== 填充后的文档内容 ===")
        
        # 提取所有文本
        for para in doc.paragraphs:
            if para.text.strip():
                print(f"段落: {para.text}")
        
        # 检查表格
        for i, table in enumerate(doc.tables):
            print(f"\n表格 {i+1}:")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                print(f"  {' | '.join(row_text)}")
                
    except Exception as e:
        print(f"错误: {e}")

if __name__ == '__main__':
    check()
