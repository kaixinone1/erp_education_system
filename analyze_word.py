from docx import Document
import re

# 读取Word文件
doc = Document('职工退休呈报表.docx')

print('=== Word文件分析 ===')
print(f'段落数量: {len(doc.paragraphs)}')
print(f'表格数量: {len(doc.tables)}')

# 获取页面设置
section = doc.sections[0]
print('\n=== 页面设置 ===')
print(f'页面宽度: {section.page_width.cm:.2f} cm ({section.page_width.inches:.2f} inches)')
print(f'页面高度: {section.page_height.cm:.2f} cm ({section.page_height.inches:.2f} inches)')
print(f'左边距: {section.left_margin.cm:.2f} cm')
print(f'右边距: {section.right_margin.cm:.2f} cm')
print(f'上边距: {section.top_margin.cm:.2f} cm')
print(f'下边距: {section.bottom_margin.cm:.2f} cm')

# 判断纸张大小和方向
width_cm = section.page_width.cm
height_cm = section.page_height.cm
is_landscape = width_cm > height_cm
paper_size = 'A3' if max(width_cm, height_cm) >= 40 else 'A4'
print(f'纸张大小: {paper_size}')
print(f'纸张方向: {"横向" if is_landscape else "纵向"}')

# 查找所有占位符
print('\n=== 查找占位符 ===')
all_text = []
for para in doc.paragraphs:
    all_text.append(para.text)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            all_text.append(cell.text)

full_text = '\n'.join(all_text)
placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)

print(f'找到 {len(placeholders)} 个占位符:')
for i, ph in enumerate(placeholders, 1):
    print(f'{i}. {{{{{ph}}}}}')
