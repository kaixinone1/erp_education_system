import sys, os
sys.path.insert(0, r'd:\erp_thirteen\tp_education_system\backend')

# 直接用fill_word_template和_build_word_fill_data模拟用户实际看到的效果
from services.word_template_filler import fill_word_template
from routes.universal_template_routes import _build_word_fill_data

test_data = {
    '姓名': '张三', '性别': '男', '出生日期': '1965-03-15',
    '民族': '汉族', '是否独生子女': '是',
    '入党年月': '1990-07', '职务': '教师', '技术职称': '高级教师',
    '参加工作时间': '1985-09-01', '工作年限': '40',
    '籍贯': '湖北枣阳', '现住址': '枣阳市太平镇', '退休原因': '正常退休',
    '直系亲属信息': '', '退休后居住地址': '', '发给退休费的单位': '',
    '自何年何月': '1985-09', '至何年何月': '2025-03',
    '所在单位及职务': '太平镇第一中学 教师', '证明人及其住址': '',
    '退休时事业专技岗位8': '五级专技', '对应原职务8': '副高级', '薪级8': '35级',
    '事业专技岗位2': '十二级专技', '对应原职务2': '初级', '薪级2': '8级',
    '事业专技岗位5': '八级专技', '对应原职务5': '中级', '薪级5': '22级',
    '最后一次职务升降时间': '2018-06-15',
    '退休方式序号': '二', '审批退休方式序号': '二',
    '退休执行年月': '2025年5月', '审批退休执行年月': '2025年5月',
    '独生子女费金额': '3500', '特殊贡献奖金额': '0', '补贴执行年月': '2025年5月',
}

fill_data = _build_word_fill_data(test_data, '')
template = r'D:\erp_thirteen\数据库信息\模板\职工退休呈报表.docx'
output = r'd:\erp_thirteen\_test_retirement.docx'

fill_word_template(template, output, fill_data)

from docx import Document
doc = Document(output)
table = doc.tables[0]

print('=== 用户下载的文件内容 ===')
print(f'文件: {output}')
print(f'表格数: {len(doc.tables)}')
print(f'审批表: {len(table.rows)}行 x {len(table.columns)}列')

names = ['呈报单位', '主管部门', '一次性补贴', '批准机关']
for ri in [0, 1, 2, 3]:
    cell = table.rows[ri].cells[1]
    print(f'\n行{ri+1}（{names[ri]}）共{len(cell.paragraphs)}段:')
    date_count = 0
    for pi, p in enumerate(cell.paragraphs):
        txt = p.text
        is_date = '年' in txt and '月' in txt and '日' in txt and '执行' not in txt
        if is_date:
            date_count += 1
        if txt.strip():
            print(f'  P{pi+1}: {repr(txt[:120])} {"<--落款" if is_date else ""}')
        else:
            print(f'  P{pi+1}: (空)')
    if date_count > 1:
        print(f'  *** 重复！{date_count}个落款日期 ***')
    elif date_count == 0:
        print(f'  *** 缺失 ***')
    else:
        print(f'  正常: 1个落款日期')

print(f'\n文件: {output}')