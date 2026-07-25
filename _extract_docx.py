import sys
sys.path.append(r'd:\erp_thirteen\tp_education_system\backend')
from docx import Document
from lxml import etree

template_path = r"D:\erp_thirteen\数据库信息\模板\职工退休呈报表.docx"
doc = Document(template_path)

# 检查审批表各单元格的段落
table = doc.tables[0]
for ri, row in enumerate(table.rows):
    cell = row.cells[1]  # 列2
    print(f"\n=== 行{ri+1}列2 ===")
    for pi, para in enumerate(cell.paragraphs):
        print(f"  段落{pi+1}: [{para.text}] (runs: {len(para.runs)})")
        for runi, run in enumerate(para.runs):
            print(f"    Run{runi}: text=[{run.text}] font={run.font.name} size={run.font.size} underline={run.underline}")

# 检查原始XML中行1列2的段落
print("\n=== 行1列2 原始XML ===")
cell1 = table.rows[0].cells[1]
xml = etree.tostring(cell1._element, pretty_print=True, encoding='unicode')
print(xml[:3000])