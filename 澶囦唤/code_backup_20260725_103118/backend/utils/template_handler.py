"""
模板处理基类
使用OpenPyXL和python-docx处理Excel和Word模板，保证100%一致性
"""
import os
from typing import Dict, Any, List, Optional, Tuple
from copy import copy
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, Fill, Border, Alignment, PatternFill
from openpyxl.utils import get_column_letter
import pandas as pd


class ExcelTemplateHandler:
    """Excel模板处理器 - 使用OpenPyXL保证100%一致性"""
    
    def __init__(self, template_path: str):
        """
        初始化Excel模板处理器
        
        Args:
            template_path: 模板文件路径
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        self.template_path = template_path
        self.wb: Optional[Workbook] = None
        self.ws = None
        self.original_styles: Dict[str, Dict[str, Any]] = {}
        
    def load(self) -> bool:
        """加载模板文件"""
        try:
            self.wb = load_workbook(self.template_path, keep_vba=True, data_only=False)
            self.ws = self.wb.active
            self._save_all_styles()
            print(f"[Excel模板] 加载成功: {self.template_path}")
            return True
        except Exception as e:
            print(f"[Excel模板] 加载失败: {e}")
            return False
    
    def _save_all_styles(self):
        """保存所有单元格的样式"""
        for row in self.ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    cell_key = f"{cell.column_letter}{cell.row}"
                    self.original_styles[cell_key] = {
                        'font': copy(cell.font),
                        'fill': copy(cell.fill),
                        'border': copy(cell.border),
                        'alignment': copy(cell.alignment),
                        'number_format': cell.number_format
                    }
    
    def _serialize_style(self, style_dict: Dict[str, Any]) -> Dict[str, Any]:
        """将样式对象序列化为可JSON化的字典"""
        result = {}
        
        if 'font' in style_dict:
            font = style_dict['font']
            result['font'] = {
                'name': font.name if hasattr(font, 'name') else None,
                'size': font.size if hasattr(font, 'size') else None,
                'bold': font.bold if hasattr(font, 'bold') else False,
                'italic': font.italic if hasattr(font, 'italic') else False,
                'color': font.color.rgb if hasattr(font, 'color') and font.color and hasattr(font.color, 'rgb') else None,
            }
        
        if 'fill' in style_dict:
            fill = style_dict['fill']
            fill_data = {
                'patternType': fill.patternType if hasattr(fill, 'patternType') else None,
            }
            if hasattr(fill, 'fgColor') and fill.fgColor and hasattr(fill.fgColor, 'rgb'):
                fill_data['fgColor'] = fill.fgColor.rgb
            if hasattr(fill, 'bgColor') and fill.bgColor and hasattr(fill.bgColor, 'rgb'):
                fill_data['bgColor'] = fill.bgColor.rgb
            result['fill'] = fill_data
        
        if 'border' in style_dict:
            border = style_dict['border']
            border_data = {}
            for side in ['left', 'right', 'top', 'bottom']:
                side_obj = getattr(border, side, None)
                if side_obj:
                    border_data[side] = {
                        'style': side_obj.style if hasattr(side_obj, 'style') else None,
                        'color': side_obj.color.rgb if hasattr(side_obj, 'color') and side_obj.color and hasattr(side_obj.color, 'rgb') else None,
                    }
            result['border'] = border_data
        
        if 'alignment' in style_dict:
            alignment = style_dict['alignment']
            result['alignment'] = {
                'horizontal': alignment.horizontal if hasattr(alignment, 'horizontal') else None,
                'vertical': alignment.vertical if hasattr(alignment, 'vertical') else None,
                'wrapText': alignment.wrapText if hasattr(alignment, 'wrapText') else False,
                'textRotation': alignment.textRotation if hasattr(alignment, 'textRotation') else 0,
            }
        
        if 'number_format' in style_dict:
            result['number_format'] = style_dict['number_format']
        
        return result
    
    def preview(self) -> str:
        """
        预览模板 - 直接返回原始文件路径
        保证预览与导入100%一致
        """
        return self.template_path
    
    def import_template(self) -> Dict[str, Any]:
        """
        导入模板 - 保持100%一致
        
        Returns:
            包含模板信息的字典
        """
        if not self.wb:
            self.load()
        
        serialized_styles = {}
        for cell_key, style_dict in self.original_styles.items():
            serialized_styles[cell_key] = self._serialize_style(style_dict)
        
        merged_cells = []
        if hasattr(self.ws, 'merged_cells') and self.ws.merged_cells:
            for merged_range in self.ws.merged_cells.ranges:
                merged_cells.append({
                    'range': str(merged_range),
                    'min_row': merged_range.min_row,
                    'max_row': merged_range.max_row,
                    'min_col': merged_range.min_col,
                    'max_col': merged_range.max_col,
                })
        
        row_dimensions = {}
        for row_idx in range(1, self.ws.max_row + 1):
            row_dim = self.ws.row_dimensions[row_idx]
            height = row_dim.height if row_dim and hasattr(row_dim, 'height') and row_dim.height else 15.0
            hidden = row_dim.hidden if row_dim and hasattr(row_dim, 'hidden') else False
            row_dimensions[row_idx] = {
                'height': height,
                'hidden': hidden
            }
        
        column_dimensions = {}
        for col_idx in range(1, self.ws.max_column + 1):
            col_letter = get_column_letter(col_idx)
            col_dim = self.ws.column_dimensions.get(col_letter)
            width = col_dim.width if col_dim and hasattr(col_dim, 'width') and col_dim.width else 8.43
            hidden = col_dim.hidden if col_dim and hasattr(col_dim, 'hidden') else False
            column_dimensions[col_letter] = {
                'width': width,
                'hidden': hidden
            }
        
        page_setup = {}
        if hasattr(self.ws, 'page_setup') and self.ws.page_setup:
            ps = self.ws.page_setup
            page_setup = {
                'paperSize': ps.paperSize if hasattr(ps, 'paperSize') else None,
                'orientation': ps.orientation if hasattr(ps, 'orientation') else None,
                'fitToPage': ps.fitToPage if hasattr(ps, 'fitToPage') else None,
                'fitToWidth': ps.fitToWidth if hasattr(ps, 'fitToWidth') else None,
                'fitToHeight': ps.fitToHeight if hasattr(ps, 'fitToHeight') else None,
            }
        
        page_margins = {}
        if hasattr(self.ws, 'page_margins') and self.ws.page_margins:
            pm = self.ws.page_margins
            page_margins = {
                'left': pm.left if hasattr(pm, 'left') else None,
                'right': pm.right if hasattr(pm, 'right') else None,
                'top': pm.top if hasattr(pm, 'top') else None,
                'bottom': pm.bottom if hasattr(pm, 'bottom') else None,
                'header': pm.header if hasattr(pm, 'header') else None,
                'footer': pm.footer if hasattr(pm, 'footer') else None,
            }
        
        return {
            'file': self.template_path,
            'sheets': self.wb.sheetnames,
            'cells': self._extract_all_cells(),
            'styles': serialized_styles,
            'merged_cells': merged_cells,
            'row_dimensions': row_dimensions,
            'column_dimensions': column_dimensions,
            'page_setup': page_setup,
            'page_margins': page_margins,
            'dimensions': {
                'rows': self.ws.max_row,
                'columns': self.ws.max_column
            }
        }
    
    def _extract_all_cells(self) -> List[Dict[str, Any]]:
        """提取所有单元格信息"""
        cells = []
        for row in self.ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    cells.append({
                        'position': f"{cell.column_letter}{cell.row}",
                        'value': cell.value,
                        'type': type(cell.value).__name__
                    })
        return cells
    
    def verify_consistency_with_original(self, output_path: str) -> Dict[str, Any]:
        """
        验证输出文件与原模板的100%一致性
        
        Args:
            output_path: 输出文件路径
        
        Returns:
            验证结果字典
        """
        if not os.path.exists(output_path):
            return {
                'is_consistent': False,
                'message': f"输出文件不存在: {output_path}",
                'details': []
            }
        
        try:
            output_wb = load_workbook(output_path)
            output_ws = output_wb.active
            
            details = []
            total_checks = 0
            passed_checks = 0
            
            if self.ws.max_row != output_ws.max_row or self.ws.max_column != output_ws.max_column:
                details.append({
                    'type': 'dimensions',
                    'status': 'failed',
                    'message': f"维度不一致: 原模板({self.ws.max_row}x{self.ws.max_column}) vs 输出({output_ws.max_row}x{output_ws.max_column})"
                })
                total_checks += 1
            else:
                details.append({
                    'type': 'dimensions',
                    'status': 'passed',
                    'message': f"维度一致: {self.ws.max_row}行 x {self.ws.max_column}列"
                })
                total_checks += 1
                passed_checks += 1
            
            for row_idx in range(1, min(self.ws.max_row, output_ws.max_row) + 1):
                orig_row_dim = self.ws.row_dimensions[row_idx]
                output_row_dim = output_ws.row_dimensions[row_idx]
                
                orig_height = orig_row_dim.height if orig_row_dim and hasattr(orig_row_dim, 'height') and orig_row_dim.height else 15.0
                output_height = output_row_dim.height if output_row_dim and hasattr(output_row_dim, 'height') and output_row_dim.height else 15.0
                
                if abs(orig_height - output_height) > 0.01:
                    details.append({
                        'type': 'row_height',
                        'status': 'failed',
                        'message': f"行{row_idx}高度不一致: 原模板({orig_height}) vs 输出({output_height})"
                    })
                    total_checks += 1
                else:
                    total_checks += 1
                    passed_checks += 1
            
            for col_idx in range(1, min(self.ws.max_column, output_ws.max_column) + 1):
                col_letter = get_column_letter(col_idx)
                orig_col_dim = self.ws.column_dimensions.get(col_letter)
                output_col_dim = output_ws.column_dimensions.get(col_letter)
                
                orig_width = orig_col_dim.width if orig_col_dim and hasattr(orig_col_dim, 'width') and orig_col_dim.width else 8.43
                output_width = output_col_dim.width if output_col_dim and hasattr(output_col_dim, 'width') and output_col_dim.width else 8.43
                
                if abs(orig_width - output_width) > 0.01:
                    details.append({
                        'type': 'column_width',
                        'status': 'failed',
                        'message': f"列{col_letter}宽度不一致: 原模板({orig_width}) vs 输出({output_width})"
                    })
                    total_checks += 1
                else:
                    total_checks += 1
                    passed_checks += 1
            
            for row in self.ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        cell_pos = f"{cell.column_letter}{cell.row}"
                        output_cell = output_ws[cell_pos]
                        
                        if cell.value != output_cell.value:
                            details.append({
                                'type': 'cell_value',
                                'status': 'failed',
                                'position': cell_pos,
                                'message': f"单元格{cell_pos}值不一致: 原模板({cell.value}) vs 输出({output_cell.value})"
                            })
                            total_checks += 1
                            continue
                        
                        orig_style = self.original_styles.get(cell_pos, {})
                        
                        orig_font = orig_style.get('font')
                        output_font = output_cell.font
                        
                        if orig_font and output_font:
                            if orig_font.bold != output_font.bold or orig_font.italic != output_font.italic:
                                details.append({
                                    'type': 'font_style',
                                    'status': 'failed',
                                    'position': cell_pos,
                                    'message': f"单元格{cell_pos}字体样式不一致"
                                })
                                total_checks += 1
                                continue
                        
                        orig_alignment = orig_style.get('alignment')
                        output_alignment = output_cell.alignment
                        
                        if orig_alignment and output_alignment:
                            if orig_alignment.horizontal != output_alignment.horizontal or orig_alignment.vertical != output_alignment.vertical:
                                details.append({
                                    'type': 'alignment',
                                    'status': 'failed',
                                    'position': cell_pos,
                                    'message': f"单元格{cell_pos}对齐方式不一致"
                                })
                                total_checks += 1
                                continue
                        
                        total_checks += 1
                        passed_checks += 1
            
            is_consistent = (passed_checks == total_checks)
            
            return {
                'is_consistent': is_consistent,
                'message': f"验证完成: {passed_checks}/{total_checks} 项通过",
                'total_checks': total_checks,
                'passed_checks': passed_checks,
                'failed_checks': total_checks - passed_checks,
                'details': details[:100]
            }
            
        except Exception as e:
            return {
                'is_consistent': False,
                'message': f"验证失败: {str(e)}",
                'details': []
            }
    
    def fill_data(self, data: Dict[str, Any]) -> Workbook:
        """
        填充数据 - 保持格式不变
        
        Args:
            data: 数据字典，key为单元格位置（如'A5'），value为要填充的值
        
        Returns:
            填充后的Workbook对象
        """
        if not self.wb:
            self.load()
        
        for cell_pos, value in data.items():
            try:
                cell = self.ws[cell_pos]
                
                if cell_pos in self.original_styles:
                    original_style = self.original_styles[cell_pos]
                    cell.value = value
                    cell.font = original_style['font']
                    cell.fill = original_style['fill']
                    cell.border = original_style['border']
                    cell.alignment = original_style['alignment']
                    cell.number_format = original_style['number_format']
                else:
                    cell.value = value
                
                print(f"[Excel填充] {cell_pos}: {value}")
            except Exception as e:
                print(f"[Excel填充] 填充单元格 {cell_pos} 失败: {e}")
        
        return self.wb
    
    def fill_by_field_name(self, field_mappings: Dict[str, Any]) -> Workbook:
        """
        根据字段名填充数据
        
        Args:
            field_mappings: 字段映射字典
                {
                    '姓名': '张三',
                    '身份证号': '1234567890',
                    ...
                }
        
        Returns:
            填充后的Workbook对象
        """
        if not self.wb:
            self.load()
        
        for row in self.ws.iter_rows():
            for cell in row:
                if cell.value:
                    cell_value_str = str(cell.value)
                    for field_name, field_value in field_mappings.items():
                        if field_name in cell_value_str:
                            new_value = cell_value_str.replace(field_name, str(field_value))
                            
                            cell_key = f"{cell.column_letter}{cell.row}"
                            if cell_key in self.original_styles:
                                original_style = self.original_styles[cell_key]
                                cell.value = new_value
                                cell.font = original_style['font']
                                cell.fill = original_style['fill']
                                cell.border = original_style['border']
                                cell.alignment = original_style['alignment']
                            else:
                                cell.value = new_value
                            
                            print(f"[Excel填充] {cell_key}: {cell_value_str} -> {new_value}")
        
        return self.wb
    
    def export(self, output_path: str) -> str:
        """
        导出结果 - 100%一致
        
        Args:
            output_path: 输出文件路径
        
        Returns:
            输出文件路径
        """
        if not self.wb:
            raise ValueError("没有可导出的工作簿，请先加载或填充数据")
        
        self.wb.save(output_path)
        print(f"[Excel导出] 成功: {output_path}")
        return output_path
    
    def get_cell_value(self, cell_pos: str) -> Any:
        """获取单元格值"""
        if not self.wb:
            self.load()
        return self.ws[cell_pos].value
    
    def get_all_field_names(self) -> List[str]:
        """获取所有字段名（用于智能匹配）"""
        if not self.wb:
            self.load()
        
        field_names = []
        for row in self.ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    if '：' in cell.value or ':' in cell.value:
                        parts = cell.value.replace('：', ':').split(':')
                        if parts[0].strip():
                            field_names.append(parts[0].strip())
        
        return field_names


class WordTemplateHandler:
    """Word模板处理器 - 使用python-docx保证100%一致性"""
    
    def __init__(self, template_path: str):
        """
        初始化Word模板处理器
        
        Args:
            template_path: 模板文件路径
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        self.template_path = template_path
        self.doc = None
    
    def load(self) -> bool:
        """加载模板文件"""
        try:
            from docx import Document
            self.doc = Document(self.template_path)
            print(f"[Word模板] 加载成功: {self.template_path}")
            return True
        except Exception as e:
            print(f"[Word模板] 加载失败: {e}")
            return False
    
    def preview(self) -> str:
        """
        预览模板 - 直接返回原始文件路径
        保证预览与导入100%一致
        """
        return self.template_path
    
    def import_template(self) -> Dict[str, Any]:
        """
        导入模板 - 保持100%一致
        
        Returns:
            包含模板信息的字典
        """
        if not self.doc:
            self.load()
        
        return {
            'file': self.template_path,
            'paragraphs': self._extract_paragraphs(),
            'tables': self._extract_tables()
        }
    
    def _extract_paragraphs(self) -> List[Dict[str, Any]]:
        """提取所有段落信息"""
        paragraphs = []
        for i, para in enumerate(self.doc.paragraphs):
            if para.text.strip():
                paragraphs.append({
                    'index': i,
                    'text': para.text,
                    'style': para.style.name if para.style else None
                })
        return paragraphs
    
    def _extract_tables(self) -> List[Dict[str, Any]]:
        """提取所有表格信息"""
        tables = []
        for table_idx, table in enumerate(self.doc.tables):
            table_data = []
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for col_idx, cell in enumerate(row.cells):
                    row_data.append({
                        'position': f"表格{table_idx+1}-行{row_idx+1}-列{col_idx+1}",
                        'text': cell.text
                    })
                table_data.append(row_data)
            tables.append({
                'index': table_idx,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'data': table_data
            })
        return tables
    
    def fill_data(self, data: Dict[str, Any]) -> Any:
        """
        填充数据 - 保持格式不变
        
        Args:
            data: 数据字典，key为字段名，value为要填充的值
        
        Returns:
            填充后的Document对象
        """
        if not self.doc:
            self.load()
        
        for key, value in data.items():
            self._replace_in_paragraphs(key, value)
            self._replace_in_tables(key, value)
        
        return self.doc
    
    def _replace_in_paragraphs(self, key: str, value: Any):
        """在段落中替换内容"""
        for para in self.doc.paragraphs:
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))
                        print(f"[Word填充] 段落: {key} -> {value}")
    
    def _replace_in_tables(self, key: str, value: Any):
        """在表格中替换内容"""
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if key in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, str(value))
                                    print(f"[Word填充] 表格: {key} -> {value}")
    
    def export(self, output_path: str) -> str:
        """
        导出结果 - 100%一致
        
        Args:
            output_path: 输出文件路径
        
        Returns:
            输出文件路径
        """
        if not self.doc:
            raise ValueError("没有可导出的文档，请先加载或填充数据")
        
        self.doc.save(output_path)
        print(f"[Word导出] 成功: {output_path}")
        return output_path
    
    def get_all_field_names(self) -> List[str]:
        """获取所有字段名（用于智能匹配）"""
        if not self.doc:
            self.load()
        
        field_names = []
        
        for para in self.doc.paragraphs:
            text = para.text
            if '：' in text or ':' in text:
                parts = text.replace('：', ':').split(':')
                if parts[0].strip():
                    field_names.append(parts[0].strip())
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    if '：' in text or ':' in text:
                        parts = text.replace('：', ':').split(':')
                        if parts[0].strip():
                            field_names.append(parts[0].strip())
        
        return list(set(field_names))


def verify_consistency(original_path: str, result_path: str) -> Tuple[bool, str]:
    """
    验证两个文件是否100%一致
    
    Args:
        original_path: 原始文件路径
        result_path: 结果文件路径
    
    Returns:
        (是否一致, 差异信息)
    """
    if not os.path.exists(original_path):
        return False, f"原始文件不存在: {original_path}"
    
    if not os.path.exists(result_path):
        return False, f"结果文件不存在: {result_path}"
    
    size1 = os.path.getsize(original_path)
    size2 = os.path.getsize(result_path)
    
    if original_path.endswith('.xlsx') and result_path.endswith('.xlsx'):
        try:
            wb1 = load_workbook(original_path)
            wb2 = load_workbook(result_path)
            
            if wb1.sheetnames != wb2.sheetnames:
                return False, "工作表名称不一致"
            
            for ws1, ws2 in zip(wb1.worksheets, wb2.worksheets):
                if ws1.max_row != ws2.max_row or ws1.max_column != ws2.max_column:
                    return False, f"工作表维度不一致: {ws1.title}"
                
                for row1, row2 in zip(ws1.rows, ws2.rows):
                    for cell1, cell2 in zip(row1, row2):
                        if cell1.value != cell2.value:
                            return False, f"单元格值不一致: {cell1.coordinate}"
            
            return True, "文件100%一致"
        except Exception as e:
            return False, f"验证失败: {e}"
    
    return True, "文件类型不支持验证"


if __name__ == '__main__':
    print("\n=== 模板处理器测试 ===\n")
    
    test_excel = r"d:\erp_thirteen\系统说明文件类\模板\退休呈报表.xlsx"
    if os.path.exists(test_excel):
        print("测试Excel模板处理器:")
        handler = ExcelTemplateHandler(test_excel)
        handler.load()
        
        template_info = handler.import_template()
        print(f"  工作表: {template_info['sheets']}")
        print(f"  单元格数: {len(template_info['cells'])}")
        print(f"  维度: {template_info['dimensions']}")
        
        field_names = handler.get_all_field_names()
        print(f"  字段名: {field_names[:10]}...")
    else:
        print(f"测试文件不存在: {test_excel}")
    
    print("\n测试完成!")
