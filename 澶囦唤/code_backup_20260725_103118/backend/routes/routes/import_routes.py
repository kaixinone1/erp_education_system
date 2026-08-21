from fastapi import APIRouter, UploadFile, File, HTTPException, Body, Query
import pandas as pd
import io
import re
from typing import List, Dict, Any, Optional
import sys
import os
from datetime import datetime

# 添加services目录到路径
sys.path.append(os.path.dirname(os.path.dirname(__file__)))
from services.validation_service import ValidationService
from services.import_service import ImportService
from services.universal_import_service_v3 import UniversalImportServiceV3
from services.mapping_optimizer import mapping_optimizer

# 导入元数据引擎
sys.path.append(os.path.dirname(os.path.dirname(__file__)))
from core.metadata_engine import get_metadata_engine
from core.table_name_manager import TableNameManager
from core.field_name_manager import field_name_manager

router = APIRouter(prefix="/api/import", tags=["import"])


def convert_date_format(value: str) -> str:
    """
    智能转换日期格式为 YYYY-MM-DD（统一日期格式管理工具）
    支持：2001-01-01, 2001-1-1, 2001/01/01, 2001/1/1, 2001年01月01日, 2001年1月1日
    支持带时分秒：2001-01-01 10:30:00, 2001-1-1 10:30:00, 2001/01/01 10:30:00 等
    只保留日期部分，去掉时分秒和毫秒
    如果已经是标准格式，直接返回
    """
    if not value or pd.isna(value):
        return None
    
    value_str = str(value).strip()
    if not value_str:
        return None
    
    # 如果已经是标准格式 YYYY-MM-DD，直接返回
    if re.match(r'^\d{4}-\d{2}-\d{2}$', value_str):
        return value_str
    
    try:
        # 匹配 2001-1-1 格式（不补零的日期）
        if re.match(r'^\d{4}-\d{1,2}-\d{1,2}$', value_str):
            year, month, day = value_str.split('-')
            return f"{year}-{int(month):02d}-{int(day):02d}"
        
        # 匹配 2001/1/1 或 2001/01/01 格式
        elif re.match(r'^\d{4}/\d{1,2}/\d{1,2}$', value_str):
            year, month, day = value_str.split('/')
            return f"{year}-{int(month):02d}-{int(day):02d}"
        
        # 匹配 2001年1月1日 或 2001年01月01日 格式
        elif re.match(r'^\d{4}年\d{1,2}月\d{1,2}日$', value_str):
            match = re.match(r'^(\d{4})年(\d{1,2})月(\d{1,2})日$', value_str)
            if match:
                year, month, day = match.groups()
                return f"{year}-{int(month):02d}-{int(day):02d}"
        
        # 匹配带时分秒（含毫秒）的格式，统一转为yyyy-MM-dd
        elif re.match(r'^\d{4}-\d{1,2}-\d{1,2}\s+\d{1,2}:\d{1,2}:\d{1,2}', value_str):
            match = re.match(r'^(\d{4})-(\d{1,2})-(\d{1,2})\s+(\d{1,2}):(\d{1,2}):(\d{1,2})', value_str)
            if match:
                year, month, day = match.group(1), match.group(2), match.group(3)
                return f"{year}-{int(month):02d}-{int(day):02d}"
        
        elif re.match(r'^\d{4}/\d{1,2}/\d{1,2}\s+\d{1,2}:\d{1,2}:\d{1,2}', value_str):
            match = re.match(r'^(\d{4})/(\d{1,2})/(\d{1,2})\s+(\d{1,2}):(\d{1,2}):(\d{1,2})', value_str)
            if match:
                year, month, day = match.group(1), match.group(2), match.group(3)
                return f"{year}-{int(month):02d}-{int(day):02d}"
        
        elif re.match(r'^\d{4}年\d{1,2}月\d{1,2}日\s+\d{1,2}:\d{1,2}:\d{1,2}', value_str):
            match = re.match(r'^(\d{4})年(\d{1,2})月(\d{1,2})日\s+(\d{1,2}):(\d{1,2}):(\d{1,2})', value_str)
            if match:
                year, month, day = match.group(1), match.group(2), match.group(3)
                return f"{year}-{int(month):02d}-{int(day):02d}"
        
        # 尝试用 datetime 解析其他格式（包含毫秒等）
        for fmt in [
            '%Y-%m-%d %H:%M:%S',
            '%Y/%m/%d %H:%M:%S',
            '%Y年%m月%d日 %H:%M:%S',
            '%Y-%m-%d %H:%M:%S.%f', '%Y/%m/%d %H:%M:%S.%f',
            '%Y-%m-%dT%H:%M:%S', '%Y-%m-%dT%H:%M:%S.%f',
            '%Y-%m-%d %H:%M', '%Y/%m/%d %H:%M',
            '%Y-%m-%d',
            '%Y/%m/%d',
            '%Y年%m月%d日',
        ]:
            try:
                dt = datetime.strptime(value_str, fmt)
                return dt.strftime('%Y-%m-%d')
            except ValueError:
                continue
        
        # 如果都无法解析，返回原值
        return value_str
        
    except Exception as e:
        print(f"日期格式转换失败: {value}, 错误: {e}")
        return value_str

# 获取元数据引擎实例
metadata_engine = get_metadata_engine()

# 获取表名管理器实例
table_name_manager = TableNameManager()

# 智能映射规则库：中文字段名 -> 英文字段名和数据类型
FIELD_MAPPING_RULES = {
    # 基础信息字段
    "姓名": {"target_field": "name", "data_type": "VARCHAR", "length": 50},
    "名字": {"target_field": "name", "data_type": "VARCHAR", "length": 50},
    "性别": {"target_field": "gender", "data_type": "VARCHAR", "length": 10},
    "年龄": {"target_field": "age", "data_type": "INTEGER"},
    "出生日期": {"target_field": "birth_date", "data_type": "DATE"},
    "生日": {"target_field": "birth_date", "data_type": "DATE"},
    
    # 身份证相关
    "身份证号": {"target_field": "id_card", "data_type": "VARCHAR", "length": 18},
    "身份证号码": {"target_field": "id_card", "data_type": "VARCHAR", "length": 18},
    "身份证": {"target_field": "id_card", "data_type": "VARCHAR", "length": 18},
    
    # 联系方式
    "电话": {"target_field": "phone", "data_type": "VARCHAR", "length": 20},
    "联系电话": {"target_field": "phone", "data_type": "VARCHAR", "length": 20},
    "手机": {"target_field": "mobile", "data_type": "VARCHAR", "length": 20},
    "手机号码": {"target_field": "mobile", "data_type": "VARCHAR", "length": 20},
    "邮箱": {"target_field": "email", "data_type": "VARCHAR", "length": 100},
    "电子邮箱": {"target_field": "email", "data_type": "VARCHAR", "length": 100},
    
    # 地址信息
    "地址": {"target_field": "address", "data_type": "VARCHAR", "length": 200},
    "家庭地址": {"target_field": "home_address", "data_type": "VARCHAR", "length": 200},
    "工作单位": {"target_field": "work_unit", "data_type": "VARCHAR", "length": 100},
    "单位": {"target_field": "work_unit", "data_type": "VARCHAR", "length": 100},
    
    # 教育信息
    "学历": {"target_field": "education", "data_type": "VARCHAR", "length": 20},
    "学位": {"target_field": "degree", "data_type": "VARCHAR", "length": 20},
    "毕业院校": {"target_field": "school", "data_type": "VARCHAR", "length": 100},
    "专业": {"target_field": "major", "data_type": "VARCHAR", "length": 50},
    
    # 工作信息
    "职称": {"target_field": "title", "data_type": "VARCHAR", "length": 30},
    "职务": {"target_field": "position", "data_type": "VARCHAR", "length": 30},
    "部门": {"target_field": "department", "data_type": "VARCHAR", "length": 50},
    "入职日期": {"target_field": "hire_date", "data_type": "DATE"},
    "参加工作时间": {"target_field": "work_date", "data_type": "DATE"},
    
    # 工资信息
    "基本工资": {"target_field": "base_salary", "data_type": "DECIMAL", "precision": 10, "scale": 2},
    "岗位工资": {"target_field": "position_salary", "data_type": "DECIMAL", "precision": 10, "scale": 2},
    "绩效工资": {"target_field": "performance_salary", "data_type": "DECIMAL", "precision": 10, "scale": 2},
    "工资": {"target_field": "salary", "data_type": "DECIMAL", "precision": 10, "scale": 2},
    
    # 政治面貌
    "政治面貌": {"target_field": "political_status", "data_type": "VARCHAR", "length": 20},
    "党员": {"target_field": "is_party_member", "data_type": "BOOLEAN"},
    "入党日期": {"target_field": "party_date", "data_type": "DATE"},
    
    # 其他常用字段
    "备注": {"target_field": "remark", "data_type": "TEXT"},
    "状态": {"target_field": "status", "data_type": "VARCHAR", "length": 20},
    "编号": {"target_field": "code", "data_type": "VARCHAR", "length": 30},
    "序号": {"target_field": "seq_no", "data_type": "INTEGER"},
}

# 数据类型推断规则
DATA_TYPE_RULES = {
    "INTEGER": {
        "patterns": [r'^\d+$'],
        "examples": ['123', '0', '9999']
    },
    "DECIMAL": {
        "patterns": [r'^\d+\.\d+$'],
        "examples": ['123.45', '0.00', '9999.99']
    },
    "DATE": {
        "patterns": [
            r'^\d{4}-\d{1,2}-\d{1,2}$',  # 2001-01-01 或 2001-1-1
            r'^\d{4}/\d{1,2}/\d{1,2}$',  # 2001/01/01 或 2001/1/1
            r'^\d{4}年\d{1,2}月\d{1,2}日$',  # 2001年01月01日 或 2001年1月1日
            r'^\d{4}-\d{1,2}-\d{1,2}\s+\d{1,2}:\d{1,2}:\d{1,2}',  # 2001-01-01 10:30:00
            r'^\d{4}/\d{1,2}/\d{1,2}\s+\d{1,2}:\d{1,2}:\d{1,2}',  # 2001/01/01 10:30:00
            r'^\d{4}年\d{1,2}月\d{1,2}日\s+\d{1,2}:\d{1,2}:\d{1,2}',  # 2001年01月01日 10:30:00
        ],
        "examples": ['2023-01-15', '2023-1-1', '2023/01/15', '2023/1/1', '2023年01月15日', '2023年1月1日', '2023-01-15 10:30:00']
    },
    "BOOLEAN": {
        "patterns": [r'^(是|否|true|false|1|0|yes|no)$'],
        "examples": ['是', '否', 'true', 'false']
    }
}


def infer_data_type(values: List[Any]) -> Dict[str, Any]:
    """根据数据样本推断数据类型"""
    # 过滤空值
    non_null_values = [str(v) for v in values if pd.notna(v) and str(v).strip()]
    
    if not non_null_values:
        return {"data_type": "VARCHAR", "length": 255}
    
    # 检查是否匹配特定类型规则
    for data_type, rules in DATA_TYPE_RULES.items():
        match_count = 0
        for value in non_null_values:
            for pattern in rules["patterns"]:
                if re.match(pattern, str(value), re.IGNORECASE):
                    match_count += 1
                    break
        
        # 如果80%以上的数据匹配该类型，则使用该类型
        if match_count / len(non_null_values) >= 0.8:
            if data_type == "INTEGER":
                return {"data_type": "INTEGER"}
            elif data_type == "DECIMAL":
                return {"data_type": "DECIMAL", "precision": 10, "scale": 2}
            elif data_type == "DATE":
                return {"data_type": "DATE"}
            elif data_type == "BOOLEAN":
                return {"data_type": "BOOLEAN"}
    
    # 默认为VARCHAR，根据内容长度确定
    max_length = max(len(str(v)) for v in non_null_values)
    # 向上取整到标准长度
    if max_length <= 10:
        length = 10
    elif max_length <= 20:
        length = 20
    elif max_length <= 50:
        length = 50
    elif max_length <= 100:
        length = 100
    elif max_length <= 200:
        length = 200
    else:
        length = 500
    
    return {"data_type": "VARCHAR", "length": length}


def generate_smart_mapping(field_name: str, values: List[Any], module_name: str = "") -> Dict[str, Any]:
    """生成智能字段映射建议 - 使用元数据引擎"""
    # 使用元数据引擎自动映射字段
    mappings = metadata_engine.auto_map_fields([field_name], module_name)
    
    if mappings and len(mappings) > 0:
        mapping = mappings[0]
        
        # 根据数据内容验证和调整类型
        inferred = infer_data_type(values)
        
        return {
            "source_field": field_name,
            "target_field": mapping["target_field"],
            "data_type": inferred.get("data_type", mapping["data_type"]),
            "length": inferred.get("length", 255),
            "confidence": mapping["confidence"],
            "is_required": mapping.get("is_required", False),
            "is_unique": mapping.get("is_unique", False)
        }
    
    # 如果元数据引擎没有返回结果，使用原有逻辑
    # 清理字段名
    clean_field = field_name.strip()
    
    # 首先尝试精确匹配规则库
    if clean_field in FIELD_MAPPING_RULES:
        mapping = FIELD_MAPPING_RULES[clean_field].copy()
        mapping["source_field"] = field_name
        mapping["confidence"] = "high"
        return mapping
    
    # 尝试部分匹配
    for cn_name, mapping_rule in FIELD_MAPPING_RULES.items():
        if cn_name in clean_field or clean_field in cn_name:
            mapping = mapping_rule.copy()
            mapping["source_field"] = field_name
            mapping["confidence"] = "medium"
            return mapping
    
    # 如果没有匹配到规则，根据数据内容推断类型
    inferred = infer_data_type(values)
    
    # 生成英文字段名（使用下划线替换空格）
    target_field = re.sub(r'[^\w\s]', '', clean_field)
    target_field = re.sub(r'\s+', '_', target_field).lower()
    
    if not target_field:
        target_field = f"field_{hash(field_name) % 10000}"
    
    return {
        "source_field": field_name,
        "target_field": target_field,
        "confidence": "low",
        **inferred
    }


def _detect_encoding(contents: bytes) -> str:
    """自动检测文件编码，支持 UTF-8, UTF-8-BOM, GBK, GB2312, GB18030, latin-1"""
    # 按优先级尝试常见编码
    encodings = ['utf-8-sig', 'utf-8', 'gb18030', 'gbk', 'gb2312', 'latin-1']
    for enc in encodings:
        try:
            contents.decode(enc)
            return enc
        except (UnicodeDecodeError, UnicodeError):
            continue
    return 'latin-1'  # 兜底


def _detect_delimiter(text: str) -> str:
    """自动检测CSV/文本文件的分隔符"""
    # 取前几行来检测
    lines = text.split('\n')[:10]
    if not lines:
        return ','
    
    candidates = [',', '\t', ';', '|']
    # 统计每行每个分隔符的出现次数，选最稳定且出现次数最多的
    best_delim = ','
    best_score = -1
    
    for delim in candidates:
        counts = [line.count(delim) for line in lines if line.strip()]
        if not counts:
            continue
        # 要求至少2行且分隔符出现次数一致
        if len(counts) >= 2 and len(set(counts)) == 1 and counts[0] > 0:
            score = counts[0]
            if score > best_score:
                best_score = score
                best_delim = delim
    
    return best_delim


def _parse_excel_auto(contents: bytes, filename: str) -> pd.DataFrame:
    """
    统一Excel解析器，自动适配各种Excel格式。
    支持：.xls, .xlsx, .xlsm, .xlsb, .xlt, .xltx, .xltm, .ods, .et, .csv
    自动探测非空sheet，尝试多种解析引擎。
    """
    filename_lower = filename.lower()
    io_bytes = io.BytesIO(contents)
    
    # 第一步：用xlrd探测.xls文件的sheet结构（找到非空sheet）
    if filename_lower.endswith('.xls') and not filename_lower.endswith('.xlsx'):
        try:
            import xlrd
            wb = xlrd.open_workbook(file_contents=contents, on_demand=True)
            sheet_name = None
            for i in range(wb.nsheets):
                sheet = wb.sheet_by_index(i)
                if sheet.nrows > 0 and sheet.ncols > 0:
                    sheet_name = sheet.name
                    break
            if sheet_name is None:
                raise ValueError("所有工作表均为空")
            # 用xlrd引擎读取找到的非空sheet
            return pd.read_excel(io_bytes, sheet_name=sheet_name)
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"无法解析此.xls文件（{str(e)}）。请将文件另存为.xlsx格式后重新上传。"
            )
    
    # 第二步：处理其他Excel格式（.xlsx, .xlsm, .xlsb, .ods 等）
    # 引擎尝试顺序：openpyxl → calamine
    engines_to_try = ['openpyxl', 'calamine']
    
    for engine in engines_to_try:
        try:
            df = pd.read_excel(io_bytes, engine=engine)
            if df.shape[0] > 0 and df.shape[1] > 0:
                return df
        except Exception:
            io_bytes.seek(0)  # 重置指针，尝试下一个引擎
            continue
    
    # 第三步：尝试读取所有sheet，找到第一个非空的
    try:
        io_bytes.seek(0)
        all_sheets = pd.read_excel(io_bytes, sheet_name=None)
        for name, sheet_df in all_sheets.items():
            if sheet_df.shape[0] > 0 and sheet_df.shape[1] > 0:
                return sheet_df
    except Exception:
        pass
    
    raise HTTPException(
        status_code=400,
        detail=f"无法解析此文件。请确认文件格式正确，或尝试另存为.xlsx格式后重新上传。"
    )


def _parse_csv_auto(contents: bytes, filename: str) -> pd.DataFrame:
    """统一CSV/文本解析器，自动检测编码和分隔符"""
    encoding = _detect_encoding(contents)
    text = contents.decode(encoding)
    delimiter = _detect_delimiter(text)
    
    print(f"[parse-file] CSV解析: 编码={encoding}, 分隔符={repr(delimiter)}")
    
    try:
        df = pd.read_csv(io.BytesIO(contents), encoding=encoding, sep=delimiter,
                         on_bad_lines='skip')
    except Exception:
        # 如果自动检测的编码失败，用latin-1兜底
        df = pd.read_csv(io.BytesIO(contents), encoding='latin-1', sep=delimiter,
                         on_bad_lines='skip')
    
    if df.shape[0] == 0 or df.shape[1] == 0:
        raise HTTPException(
            status_code=400,
            detail="此文件解析后为空。请检查文件内容是否正确。"
        )
    return df


def parse_word_document(file_path: str) -> pd.DataFrame:
    """
    解析Word文档，提取表格数据
    返回DataFrame格式，与Excel解析结果一致
    """
    from docx import Document
    
    doc = Document(file_path)
    
    # 查找文档中的第一个表格
    if not doc.tables:
        raise ValueError("Word文档中没有找到表格")
    
    # 获取第一个表格（通常数据在第一个表格）
    table = doc.tables[0]
    
    # 提取表格数据
    data = []
    for i, row in enumerate(table.rows):
        row_data = [cell.text.strip() for cell in row.cells]
        data.append(row_data)
    
    if not data:
        raise ValueError("表格为空")
    
    # 第一行作为表头
    headers = data[0]
    rows = data[1:]
    
    # 创建DataFrame
    df = pd.DataFrame(rows, columns=headers)
    
    return df


@router.post("/parse-file")
async def parse_file(file: UploadFile = File(...), module_name: str = "", preview_only: bool = Query(True)):
    """解析上传的文件并生成智能映射建议"""
    try:
        # 检查文件类型（不区分大小写，支持更多格式）
        filename_lower = file.filename.lower()
        excel_extensions = ('.xlsx', '.xls', '.xlsm', '.xlsb', '.xlt', '.xltx', '.xltm', '.ods', '.et')
        csv_extensions = ('.csv', '.tsv', '.txt')
        word_extensions = ('.doc', '.docx')
        all_allowed = excel_extensions + csv_extensions + word_extensions
        
        if not filename_lower.endswith(all_allowed):
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件格式。支持的格式：Excel({', '.join(excel_extensions)})、CSV/TSV/TXT、Word(.doc/.docx)"
            )

        # 读取文件内容
        contents = await file.read()

        # 保存上传的文件用于检查（临时）
        upload_dir = os.path.join(os.path.dirname(__file__), '..', 'uploads')
        os.makedirs(upload_dir, exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        saved_filename = f"{timestamp}_{file.filename}"
        saved_path = os.path.join(upload_dir, saved_filename)
        with open(saved_path, 'wb') as f:
            f.write(contents)
        print(f"[parse-file] 文件已保存到: {saved_path}")

        # 根据文件类型选择合适的解析器
        if filename_lower.endswith(excel_extensions):
            df = _parse_excel_auto(contents, file.filename)
        elif filename_lower.endswith(csv_extensions):
            df = _parse_csv_auto(contents, file.filename)
        elif filename_lower.endswith(word_extensions):
            df = parse_word_document(saved_path)
        else:
            raise HTTPException(status_code=400, detail="不支持的文件格式")
        
        # 获取原始字段列表
        original_fields = df.columns.tolist()
        
        # 生成智能映射建议
        suggested_mappings = []
        for field in original_fields:
            values = df[field].tolist()
            mapping = generate_smart_mapping(field, values, module_name)
            suggested_mappings.append(mapping)
        
        # 根据preview_only参数决定返回全部数据还是仅预览
        if preview_only:
            # 获取前10行数据作为预览
            data = df.head(10).to_dict(orient='records')
        else:
            # 返回全部数据
            data = df.to_dict(orient='records')
        
        # 转换数据中的NaN为None，并智能转换日期格式
        # 首先识别哪些字段是日期类型
        date_fields = set()
        for mapping in suggested_mappings:
            if mapping.get('data_type') == 'DATE' or any(
                keyword in mapping.get('source_field', '') 
                for keyword in ['日期', '出生', '毕业', '入党', '参加', '进入']
            ):
                date_fields.add(mapping.get('source_field'))
        
        # 对所有数据进行日期格式转换
        for row in data:
            for key, value in row.items():
                if pd.isna(value):
                    row[key] = None
                # 如果字段是日期类型，转换日期格式
                elif key in date_fields and value is not None:
                    converted = convert_date_format(str(value))
                    if converted:
                        row[key] = converted
        
        return {
            "fields": original_fields,
            "preview_data": data,
            "total_rows": len(df),
            "filename": file.filename,
            "suggested_mappings": suggested_mappings
        }
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"文件解析失败: {str(e)}")


@router.post("/validate")
async def validate_data(
    data: List[Dict[str, Any]] = Body(...),
    field_configs: List[Dict[str, Any]] = Body(...),
    validation_level: int = Body(3),
    reference_data: Optional[Dict[str, List[str]]] = Body(None)
):
    """验证数据（自动分批处理大数据量）"""
    try:
        # 创建验证服务
        validation_service = ValidationService()
        
        # 如果数据量大于1000行，使用分批验证
        if len(data) > 1000:
            result = validation_service.validate_data_batch(
                data=data,
                field_configs=field_configs,
                validation_level=validation_level,
                reference_data=reference_data,
                batch_size=1000
            )
        else:
            # 数据量小于1000行，使用普通验证
            result = validation_service.validate_data(
                data=data,
                field_configs=field_configs,
                validation_level=validation_level,
                reference_data=reference_data
            )
        
        return result
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"数据验证失败: {str(e)}")


@router.post("/translate-table-name")
async def translate_table_name(
    chinese_name: str = Body(...),
    module_name: str = Body("")
):
    """翻译中文表名为英文表名"""
    try:
        english_name = metadata_engine.translate_table_name(chinese_name, module_name)
        return {
            "chinese_name": chinese_name,
            "english_name": english_name,
            "module_name": module_name
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"表名翻译失败: {str(e)}")


@router.post("/translate-field-names")
async def translate_field_names(
    chinese_fields: List[str] = Body(...),
    module_name: str = Body("")
):
    """批量翻译中文字段名为英文字段名"""
    try:
        mappings = metadata_engine.auto_map_fields(chinese_fields, module_name)
        return {
            "mappings": mappings,
            "total": len(mappings)
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"字段翻译失败: {str(e)}")


@router.post("/check-table-name")
async def check_table_name(
    chinese_name: str = Body(...),
    field_configs: List[Dict[str, Any]] = Body(...),
    table_type: str = Body("master")
):
    """
    检查表名是否可用
    返回状态：
    - 'existing': 中文表名已存在，表结构一致，直接使用
    - 'structure_mismatch': 中文表名已存在，但表结构不一致，需要修改中文表名
    - 'name_conflict': 中文表名不重复，但表结构相同，需要用户确认
    - 'new_table': 新表，可以创建
    """
    try:
        status, message, english_name = table_name_manager.check_table_name(
            chinese_name=chinese_name,
            field_configs=field_configs,
            table_type=table_type
        )
        
        return {
            "status": status,
            "message": message,
            "english_name": english_name,
            "chinese_name": chinese_name
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"检查表名失败: {str(e)}")


@router.post("/finalize")
async def finalize_import(
    table_name: str = Body(...),
    field_configs: List[Dict[str, Any]] = Body(...),
    data: List[Dict[str, Any]] = Body(...),
    module_id: str = Body(...),
    module_name: str = Body(...),
    file_name: str = Body(""),
    chinese_title: str = Body(""),
    sub_module_id: str = Body(""),
    sub_module_name: str = Body(""),
    table_type: str = Body("master"),
    parent_table: Optional[str] = Body(None),
    force_use_existing: bool = Body(False),  # 强制使用已存在的表（当表结构相同时）
    force_overwrite: bool = Body(False),  # 强制覆盖已有表（当表结构不一致时）
    analyze_only: bool = Body(False)  # 仅分析不执行，返回数据差异报告
):
    """完成导入 - 原子化操作"""
    try:
        # 0. 首先使用字段名管理器处理字段配置，确保中文字段名映射到唯一的英文字段名
        # 注意：analyze_only 模式下使用原始字段配置（因为需要匹配数据库已有列名）
        processed_field_configs = field_name_manager.process_field_configs(field_configs)
        print(f"字段配置已处理，共 {len(processed_field_configs)} 个字段")
        
        # 用于表名检查的字段配置：始终使用原始配置（需要匹配数据库已有列名）
        check_field_configs = field_configs
        
        # 1. 检查中文表名
        actual_table_name = table_name
        is_existing_table = False  # 标记是否为已有表
        if chinese_title:
            status, message, existing_english_name = table_name_manager.check_table_name(
                chinese_name=chinese_title,
                field_configs=check_field_configs,
                table_type=table_type
            )
            
            # 根据检查结果处理
            if status == 'existing':
                # 中文表名已存在，表结构一致，直接使用现有表
                print(f"使用已存在的表: {existing_english_name}")
                actual_table_name = existing_english_name
                is_existing_table = True
                
            elif status == 'structure_mismatch':
                # 中文表名已存在，但表结构不一致
                if force_overwrite:
                    # 用户确认强制覆盖：删除旧表，使用新结构重建
                    print(f"强制覆盖模式：删除旧表 {existing_english_name}，使用新结构重建")
                    
                    # 获取旧表名
                    old_mapping = table_name_manager.table_name_mappings.get("mappings", {}).get(chinese_title, {})
                    old_english_name = old_mapping.get("english_name", existing_english_name)
                    
                    # 删除旧表
                    from sqlalchemy import create_engine, text as sa_text
                    DATABASE_URL = "postgresql://taiping_user:taiping_password@localhost:5432/taiping_education"
                    engine = create_engine(DATABASE_URL)
                    with engine.connect() as conn:
                        conn.execute(sa_text(f'DROP TABLE IF EXISTS "{old_english_name}" CASCADE'))
                        conn.commit()
                    print(f"已删除旧表: {old_english_name}")
                    
                    # 清除旧映射
                    if chinese_title in table_name_manager.table_name_mappings.get("mappings", {}):
                        del table_name_manager.table_name_mappings["mappings"][chinese_title]
                    if old_english_name in table_name_manager.table_name_mappings.get("reverse_mappings", {}):
                        del table_name_manager.table_name_mappings["reverse_mappings"][old_english_name]
                    table_name_manager._save_table_name_mappings()
                    print(f"已清除旧表名映射: {chinese_title} -> {old_english_name}")
                    
                    # 使用原始表名（新表将使用相同的英文名）
                    actual_table_name = table_name
                    
                    # 注册新映射
                    table_name_manager.register_table_name(
                        chinese_name=chinese_title,
                        english_name=table_name,
                        table_type=table_type,
                        field_configs=processed_field_configs
                    )
                else:
                    # 对于 analyze_only 模式，返回字段差异报告而不是直接报错
                    if analyze_only:
                        # 获取字段差异
                        diff_report = import_service._compare_field_structures(existing_english_name, field_configs)
                        return {
                            "status": "field_diff_detected",
                            "message": "中文表名已存在，但导入字段与现有表结构不一致",
                            "chinese_name": chinese_title,
                            "existing_table": existing_english_name,
                            "field_diff": diff_report,
                            "hint": "需要启用强制覆盖才能继续导入"
                        }
                    else:
                        # 不允许覆盖，提示用户修改中文表名
                        raise HTTPException(
                            status_code=400, 
                            detail=f"中文表名'{chinese_title}'已存在，但表结构不一致。请修改中文表名后重新导入，或启用强制覆盖模式。"
                        )
                
            elif status == 'name_conflict':
                # 中文表名不重复，但表结构相同，需要用户确认
                if not force_use_existing:
                    return {
                        "status": "confirm_required",
                        "message": message,
                        "existing_table": existing_english_name,
                        "suggested_table": table_name,
                        "chinese_name": chinese_title
                    }
                # 用户确认使用现有表
                actual_table_name = existing_english_name
                is_existing_table = True
                # 注册中文表名映射到现有英文表名
                table_name_manager.register_table_name(
                    chinese_name=chinese_title,
                    english_name=existing_english_name,
                    table_type=table_type,
                    field_configs=field_configs
                )
                
            else:  # 'new_table'
                # 新表，注册表名映射（仅在非分析模式下注册）
                actual_table_name = table_name
                if not analyze_only:
                    table_name_manager.register_table_name(
                        chinese_name=chinese_title,
                        english_name=table_name,
                        table_type=table_type,
                        field_configs=field_configs
                    )
        else:
            # 没有中文标题，直接使用传入的表名
            actual_table_name = table_name
        
        # 2. 创建导入服务
        import_service = ImportService()
        
        # 2.5 如果是 analyze_only 模式，仅分析数据差异并返回报告
        if analyze_only and actual_table_name:
            # 使用原始字段配置（因为需要匹配数据库已有列名）
            analysis = import_service.analyze_data_diff(
                table_name=actual_table_name,
                field_configs=field_configs,
                data=data
            )
            return {
                "status": "analyzed",
                "message": "数据差异分析完成，请确认后执行导入",
                "table_name": actual_table_name,
                "chinese_name": chinese_title,
                "analysis": analysis
            }
        
        # 3. 执行导入
        # 关键：已有表使用原始字段配置（匹配数据库列名），新表使用处理后的配置（保证全局唯一）
        import_field_configs = field_configs if is_existing_table else processed_field_configs
        result = import_service.import_data(
            table_name=actual_table_name,
            field_configs=import_field_configs,
            data=data,
            module_id=module_id,
            module_name=module_name,
            table_type=table_type,
            parent_table=parent_table,
            file_name=file_name,
            chinese_title=chinese_title,
            sub_module_id=sub_module_id,
            sub_module_name=sub_module_name
        )
        
        # 添加表名检查信息到结果
        if chinese_title:
            result['chinese_name'] = chinese_title
            result['english_name'] = actual_table_name
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"导入失败: {str(e)}")


@router.post("/import-v3")
async def import_data_v3(
    table_name: str = Body(...),
    field_configs: List[Dict[str, Any]] = Body(...),
    data: List[Dict[str, Any]] = Body(...),
    chinese_title: Optional[str] = Body(None),
    table_type: str = Body("master")
):
    """
    使用V3导入服务导入数据（支持自动字典表管理）
    """
    try:
        # 0. 使用字段名管理器处理字段配置，确保中文字段名映射到唯一的英文字段名
        processed_field_configs = field_name_manager.process_field_configs(field_configs)
        print(f"字段配置已处理，共 {len(processed_field_configs)} 个字段")
        
        # 使用V3导入服务
        import_service = UniversalImportServiceV3()
        
        result = import_service.import_data(
            table_name=table_name,
            field_configs=processed_field_configs,
            data=data,
            auto_manage_dict=True
        )
        
        # 添加中文表名信息
        if chinese_title:
            result['chinese_name'] = chinese_title
            result['english_name'] = table_name
        
        import_service.close()
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"导入失败: {str(e)}")


@router.post("/export-report")
async def export_import_report(
    report_data: Dict[str, Any] = Body(...),
    file_name: str = Body("导入报告")
):
    """
    导出导入报告为Excel文件
    命名格式: 导出文件中文名+YYYY-MM-DD
    """
    try:
        from fastapi.responses import StreamingResponse
        import openpyxl
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

        # 创建Excel工作簿
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "导入报告"

        # 设置标题样式
        title_font = Font(name='微软雅黑', size=14, bold=True, color='FFFFFF')
        title_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        title_alignment = Alignment(horizontal='center', vertical='center')

        # 写入标题
        ws['A1'] = '导入报告'
        ws['A1'].font = title_font
        ws['A1'].fill = title_fill
        ws['A1'].alignment = title_alignment
        ws.merge_cells('A1:D1')
        ws.row_dimensions[1].height = 30

        # 写入导入时间
        ws['A2'] = f'导入时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        ws['A2'].font = Font(name='微软雅黑', size=10)
        ws.merge_cells('A2:D2')

        # 写入基本信息
        row = 4
        headers = ['项目', '内容']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = Font(name='微软雅黑', size=11, bold=True)
            cell.fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 写入数据
        info_items = [
            ('文件名', report_data.get('file_name', '-')),
            ('中文表名', report_data.get('chinese_title', '-')),
            ('英文表名', report_data.get('table_name', '-')),
            ('数据总量', str(report_data.get('total_count', 0))),
            ('成功导入', str(report_data.get('success_count', 0))),
            ('失败数量', str(report_data.get('error_count', 0))),
            ('归属模块', report_data.get('module_name', '-')),
        ]

        for item_name, item_value in info_items:
            row += 1
            ws.cell(row=row, column=1, value=item_name).font = Font(name='微软雅黑', size=10)
            ws.cell(row=row, column=2, value=item_value).font = Font(name='微软雅黑', size=10)

        # 如果有错误数据，添加错误数据sheet
        error_data = report_data.get('error_data', [])
        if error_data:
            ws2 = wb.create_sheet("错误数据")
            # 写入错误数据表头
            if error_data:
                headers = list(error_data[0].keys())
                for col, header in enumerate(headers, 1):
                    cell = ws2.cell(row=1, column=col, value=header)
                    cell.font = Font(name='微软雅黑', size=11, bold=True)
                    cell.fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
                    cell.alignment = Alignment(horizontal='center', vertical='center')

                # 写入错误数据
                for row_idx, error_row in enumerate(error_data, 2):
                    for col_idx, header in enumerate(headers, 1):
                        ws2.cell(row=row_idx, column=col_idx, value=error_row.get(header, ''))

        # 调整列宽
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width

        # 保存到内存
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        # 生成文件名: 导出文件中文名+YYYY-MM-DD.xlsx
        today = datetime.now().strftime("%Y-%m-%d")
        export_file_name = f"{file_name}{today}.xlsx"

        return StreamingResponse(
            output,
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename="{export_file_name}"'}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出报告失败: {str(e)}")


@router.post("/export-error-data")
async def export_error_data(
    error_data: List[Dict[str, Any]] = Body(...),
    original_headers: List[str] = Body(...),
    file_name: str = Body("错误数据")
):
    """
    导出错误数据为Excel文件
    命名格式: 导出文件中文名+错误数据+YYYY-MM-DD
    """
    try:
        from fastapi.responses import StreamingResponse
        import openpyxl
        from openpyxl.styles import Font, Alignment, PatternFill

        # 创建Excel工作簿
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "错误数据"

        # 设置标题样式
        header_font = Font(name='微软雅黑', size=11, bold=True)
        header_fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')

        # 写入表头（使用原始文件的表头）
        for col, header in enumerate(original_headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment

        # 写入错误数据
        for row_idx, error_row in enumerate(error_data, 2):
            for col_idx, header in enumerate(original_headers, 1):
                # 尝试从error_row中获取数据，支持多种格式
                value = error_row.get(header) or error_row.get('data', {}).get(header) or ''
                ws.cell(row=row_idx, column=col_idx, value=value)

        # 添加错误原因列
        error_col = len(original_headers) + 1
        ws.cell(row=1, column=error_col, value='错误原因').font = header_font
        ws.cell(row=1, column=error_col).fill = header_fill
        ws.cell(row=1, column=error_col).alignment = header_alignment

        for row_idx, error_row in enumerate(error_data, 2):
            error_reason = error_row.get('error') or error_row.get('error_message') or ''
            ws.cell(row=row_idx, column=error_col, value=error_reason)

        # 调整列宽
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if cell.value and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width

        # 保存到内存
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        # 生成文件名: 导出文件中文名+错误数据+YYYY-MM-DD.xlsx
        today = datetime.now().strftime("%Y-%m-%d")
        export_file_name = f"{file_name}错误数据{today}.xlsx"

        return StreamingResponse(
            output,
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename="{export_file_name}"'}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出错误数据失败: {str(e)}")


# ==================== 映射优化相关API ====================

@router.post("/optimize-table-name")
async def optimize_table_name(
    chinese_name: str = Body(...)
):
    """
    优化表名映射
    使用映射优化器生成高质量的英文表名
    """
    try:
        english_name, is_optimized = mapping_optimizer.optimize_table_name(chinese_name)
        
        return {
            "chinese_name": chinese_name,
            "english_name": english_name,
            "is_optimized": is_optimized,
            "suggestions": mapping_optimizer.get_translation_suggestions(chinese_name)
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"表名优化失败: {str(e)}")


@router.post("/optimize-field-names")
async def optimize_field_names(
    field_configs: List[Dict[str, Any]] = Body(...)
):
    """
    批量优化字段名映射
    返回优化后的字段配置和需要人工处理的字段列表
    """
    try:
        processed_configs, pending_fields = mapping_optimizer.process_field_configs(field_configs)
        
        return {
            "processed_configs": processed_configs,
            "pending_fields": pending_fields,
            "total_processed": len(processed_configs),
            "pending_count": len(pending_fields)
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"字段名优化失败: {str(e)}")


@router.post("/validate-mapping-quality")
async def validate_mapping_quality(
    chinese_name: str = Body(...),
    english_name: str = Body(...)
):
    """
    验证映射质量
    检查英文字段名是否有意义、符合命名规范
    """
    try:
        is_valid, message = mapping_optimizer.validate_mapping_quality(chinese_name, english_name)
        
        return {
            "chinese_name": chinese_name,
            "english_name": english_name,
            "is_valid": is_valid,
            "message": message
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"映射质量验证失败: {str(e)}")


@router.get("/pending-fields-report")
async def get_pending_fields_report():
    """
    获取需要人工处理的字段报告
    返回现有映射中无意义的表名和字段名映射
    """
    try:
        report = mapping_optimizer.get_pending_fields_report()
        
        return {
            "total_pending": report["total_pending"],
            "pending_fields": report["pending_fields"]
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"获取待处理字段报告失败: {str(e)}")


@router.post("/get-translation-suggestions")
async def get_translation_suggestions(
    chinese_name: str = Body(...)
):
    """
    获取翻译建议
    为中文字段名提供多个翻译选项供用户选择
    """
    try:
        suggestions = mapping_optimizer.get_translation_suggestions(chinese_name)
        
        return {
            "chinese_name": chinese_name,
            "suggestions": suggestions,
            "total_suggestions": len(suggestions)
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"获取翻译建议失败: {str(e)}")
