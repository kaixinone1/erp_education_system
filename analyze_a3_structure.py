#!/usr/bin/env python3
"""
分析A3横向对折文档结构
封面（第1面）：审批意见表
内页（第2面）：主表（详细信息）
"""
from docx import Document

def analyze_a3_structure():
    template_path = r"D:\erp_thirteen\tp_education_system\backend\uploads\templates\20260208_201747_职工退休呈报表.docx"
    
    doc = Document(template_path)
    
    print("=" * 80)
    print("A3横向对折文档结构分析")
    print("=" * 80)
    
    # 分析段落 - 封面内容
    print("\n【封面段落 - 第1面】")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"  段落{i}: {repr(text)}")
    
    # 表格1 - 封面（审批意见表）
    print("\n【表格1 - 封面（审批意见表）】")
    if len(doc.tables) > 0:
        table1 = doc.tables[0]
        print(f"尺寸: {len(table1.rows)}行 x {len(table1.columns)}列")
        for ri, row in enumerate(table1.rows):
            cells_text = [cell.text.strip() for cell in row.cells]
            print(f"  行{ri}: {cells_text}")
    
    # 表格2 - 内页（主表）
    print("\n【表格2 - 内页（主表）】")
    if len(doc.tables) > 1:
        table2 = doc.tables[1]
        print(f"尺寸: {len(table2.rows)}行 x {len(table2.columns)}列")
        print("  （内容较多，省略详细输出）")
    
    print("\n" + "=" * 80)
    print("A3横向对折布局说明：")
    print("=" * 80)
    print("""
【第1面 - 封面】（对折后在外侧）
┌─────────────────────────────────────────────────────────────┐
│  编号：                                                     │
│                                                             │
│           职 工 退 休 呈 报 表                               │
│                                                             │
│                                                             │
│                                                             │
│                                                             │
│                                                             │
│                                                             │
│                                                             │
│                                                             │
│                                                             │
│                                                             │
│                                                             │
│                                                             │
│                                                             │
│  单位：枣阳市太平镇中心学校                                  │
│  姓名：王军峰                                               │
│                                                             │
│  年   月   日                                               │
