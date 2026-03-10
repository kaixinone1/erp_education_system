import psycopg2
import os
import json
import re

# Word占位符提取
def extract_docx_placeholders(file_path):
    try:
        from docx import Document
        doc = Document(file_path)
        placeholders = set()
        for para in doc.paragraphs:
            found = re.findall(r'\{\{([^}]+)\}\}', para.text)
            placeholders.update(found)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        found = re.findall(r'\{\{([^}]+)\}\}', para.text)
                        placeholders.update(found)
        return list(placeholders)
    except Exception as e:
        print(f"提取Word占位符失败: {e}")
        return []

# Excel占位符提取
def extract_xlsx_placeholders(file_path):
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, data_only=False)
        placeholders = set()
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        found = re.findall(r'\{\{([^}]+)\}\}', cell.value)
                        placeholders.update(found)
        return list(placeholders)
    except Exception as e:
        print(f"提取Excel占位符失败: {e}")
        return []

conn = psycopg2.connect(
    host='localhost',
    port='5432',
    database='taiping_education',
    user='taiping_user',
    password='taiping_password'
)
cursor = conn.cursor()

# 模板目录
templates_dir = r"d:\erp_thirteen\tp_education_system\backend\templates"

# 遍历目录中的docx和xlsx文件
for f in os.listdir(templates_dir):
    file_path = os.path.join(templates_dir, f)
    if not os.path.isfile(file_path):
        continue
    
    file_ext = f.split('.')[-1].lower()
    if file_ext not in ['docx', 'xlsx']:
        continue
    
    template_id = f.replace('.docx', '').replace('.xlsx', '')
    template_name = template_id
    file_type = file_ext
    
    # 提取占位符
    if file_ext == 'docx':
        placeholders = extract_docx_placeholders(file_path)
    else:
        placeholders = extract_xlsx_placeholders(file_path)
    
    # 保存到数据库
    cursor.execute("""
        INSERT INTO universal_templates (template_id, template_name, file_path, file_type, placeholders)
        VALUES (%s, %s, %s, %s, %s)
        ON CONFLICT (template_id) DO UPDATE SET
            file_path = EXCLUDED.file_path,
            placeholders = EXCLUDED.placeholders,
            updated_at = CURRENT_TIMESTAMP
    """, (template_id, template_name, file_path, file_type, json.dumps(placeholders)))
    
    print(f"导入: {f} - 占位符: {placeholders}")

conn.commit()
print("\n导入完成！")

# 验证
cursor.execute("SELECT COUNT(*) FROM universal_templates")
count = cursor.fetchone()[0]
print(f"通用模板表中共 {count} 个模板")

cursor.close()
conn.close()
