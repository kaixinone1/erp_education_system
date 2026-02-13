#!/usr/bin/env python3
"""
精确分析原始Word文档结构
"""
from docx import Document

def analyze():
    doc = Document(r"D:\erp_thirteen\tp_education_system\backend\uploads\templates\20260208_201747_职工退休呈报表.docx")
    
    print("=" * 80)
    print("原始Word文档结构分析")
    print("=" * 80)
    
    # 分析段落
    print("\n【段落分析】")
    for i, para in enumerate(doc.paragraphs[:10]):  # 只看前10个非空段落
        text = para.text.strip()
        if text:
            print(f"段落{i}: {repr(text)}")
    
    # 分析表格
    print("\n【表格分析】")
    for ti, table in enumerate(doc.tables):
        print(f"\n表格{ti+1}: {len(table.rows)}行 x {len(table.columns)}列")
        for ri, row in enumerate(table.rows):
            cells_text = [cell.text.strip() for cell in row.cells]
            print(f"  行{ri}: {cells_text}")

if __name__ == '__main__':
    analyze()
