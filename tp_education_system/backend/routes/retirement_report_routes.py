"""
退休呈报表相关API
"""
from fastapi import APIRouter, HTTPException, Query
from typing import List, Dict, Any, Optional
from datetime import datetime
import re
import sys
import os

# 添加 utils 目录到路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from utils.dict_utils import get_education_name

router = APIRouter(prefix="/api/retirement", tags=["retirement"])


@router.get("/template-id")
async def get_template_id_by_name(
    name: Optional[str] = Query(None, description="模板名称（可选，不传则返回最新模板）")
):
    """根据模板名称获取模板ID，如果不传名称则返回最新上传的模板"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        if name:
            # 根据名称查询
            cursor.execute("""
                SELECT id, name FROM document_templates 
                WHERE name = %s OR name LIKE %s
                ORDER BY id DESC
                LIMIT 1
            """, (name, f"%{name}%"))
        else:
            # 返回最新上传的模板
            cursor.execute("""
                SELECT id, name FROM document_templates 
                ORDER BY created_at DESC, id DESC
                LIMIT 1
            """)
        
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="未找到任何模板，请先上传模板")
        
        return {
            "status": "success",
            "template_id": row[0],
            "template_name": row[1]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"查询模板失败: {str(e)}")
    finally:
        cursor.close()
        conn.close()


DATABASE_CONFIG = {
    "host": "localhost",
    "port": "5432",
    "database": "taiping_education",
    "user": "taiping_user",
    "password": "taiping_password"
}


def get_db_connection():
    import psycopg2
    return psycopg2.connect(**DATABASE_CONFIG)


def extract_birth_date_from_id_card(id_card: str) -> Optional[str]:
    """从身份证号提取出生日期"""
    if not id_card or len(id_card) != 18:
        return None
    try:
        year = id_card[6:10]
        month = id_card[10:12]
        day = id_card[12:14]
        return f"{year}-{month}-{day}"
    except:
        return None


def extract_gender_from_id_card(id_card: str) -> Optional[str]:
    """从身份证号提取性别（奇数男，偶数女）"""
    if not id_card or len(id_card) != 18:
        return None
    try:
        # 第17位表示性别
        gender_code = int(id_card[16])
        return "男" if gender_code % 2 == 1 else "女"
    except:
        return None


def calculate_age(birth_date: str) -> int:
    """计算年龄"""
    if not birth_date:
        return 0
    try:
        birth = datetime.strptime(birth_date, "%Y-%m-%d")
        today = datetime.now()
        age = today.year - birth.year
        if (today.month, today.day) < (birth.month, birth.day):
            age -= 1
        return age
    except:
        return 0


@router.get("/teachers")
async def get_teachers():
    """获取所有教师列表"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT id, name, id_card, archive_birth_date, employment_status
            FROM teacher_basic_info
            WHERE employment_status = '在职'
            ORDER BY id
        """)
        
        rows = cursor.fetchall()
        teachers = []
        
        for row in rows:
            teacher = {
                "id": row[0],
                "name": row[1],
                "id_card": row[2],
                "archive_birth_date": row[3],
                "employment_status": row[4]
            }
            
            # 处理出生日期
            if teacher["archive_birth_date"]:
                teacher["birth_date"] = teacher["archive_birth_date"]
            else:
                teacher["birth_date"] = extract_birth_date_from_id_card(teacher["id_card"])
            
            # 处理性别
            teacher["gender"] = extract_gender_from_id_card(teacher["id_card"])
            
            # 计算年龄
            teacher["age"] = calculate_age(teacher["birth_date"])
            
            teachers.append(teacher)
        
        cursor.close()
        conn.close()
        
        return {
            "status": "success",
            "data": teachers,
            "total": len(teachers)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取教师列表失败: {str(e)}")


@router.get("/search-by-name")
async def search_teachers_by_name(
    name: Optional[str] = Query(None, description="教师姓名"),
    id: Optional[int] = Query(None, description="教师ID")
):
    """根据姓名或ID搜索教师"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        if id:
            # 根据ID查询
            cursor.execute("""
                SELECT id, name, id_card, gender, birth_date, ethnicity,
                       native_place, work_start_date, contact_phone, archive_birth_date
                FROM teacher_basic_info
                WHERE id = %s
            """, (id,))
        elif name:
            # 根据姓名模糊查询
            cursor.execute("""
                SELECT id, name, id_card, gender, birth_date, ethnicity,
                       native_place, work_start_date, contact_phone, archive_birth_date
                FROM teacher_basic_info
                WHERE name LIKE %s
                ORDER BY id
            """, (f"%{name}%",))
        else:
            raise HTTPException(status_code=400, detail="请提供姓名或ID参数")
        
        rows = cursor.fetchall()
        teachers = []
        
        for row in rows:
            teacher = {
                "teacher_id": row[0],
                "teacher_name": row[1],
                "id_card": row[2],
                "gender": row[3],
                "birth_date": row[4],
                "ethnicity": row[5],
                "native_place": row[6],
                "work_start_date": row[7],
                "contact_phone": row[8],
                "archive_birth_date": row[9]
            }
            teachers.append(teacher)
        
        cursor.close()
        conn.close()
        
        return {
            "status": "success",
            "data": teachers,
            "total": len(teachers)
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"搜索教师失败: {str(e)}")


@router.post("/collect/{teacher_id}")
async def collect_retirement_data(teacher_id: int):
    """
    汇集退休呈报表数据到中间表
    """
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # 1. 查询教师基础信息
        cursor.execute("""
            SELECT id, name, id_card, birth_date, ethnicity,
                   native_place, work_start_date, archive_birth_date
            FROM teacher_basic_info
            WHERE id = %s
        """, (teacher_id,))
        
        teacher_row = cursor.fetchone()
        if not teacher_row:
            raise HTTPException(status_code=404, detail="教师不存在")
        
        # 2. 查询最高学历信息
        cursor.execute("""
            SELECT education_level, graduation_date
            FROM teacher_education_record
            WHERE teacher_id = %s
            ORDER BY graduation_date DESC
            LIMIT 1
        """, (teacher_id,))
        
        education_row = cursor.fetchone()
        
        # 3. 处理数据
        id_card = teacher_row[2] or ''
        
        # 出生日期：优先档案出生日期，否则从身份证提取
        birth_date = teacher_row[7] or extract_birth_date_from_id_card(id_card)
        
        # 性别：从身份证提取
        gender = extract_gender_from_id_card(id_card)
        
        # 工作年限
        work_years = calculate_age(teacher_row[6]) if teacher_row[6] else 0
        
        # 4. 插入或更新中间表
        cursor.execute("""
            INSERT INTO retirement_report_data (
                teacher_id, name, id_card, gender, birth_date,
                ethnicity, education, work_start_date, work_years,
                native_place, status, created_at, updated_at
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(), NOW())
            ON CONFLICT (teacher_id) DO UPDATE SET
                name = EXCLUDED.name,
                id_card = EXCLUDED.id_card,
                gender = EXCLUDED.gender,
                birth_date = EXCLUDED.birth_date,
                ethnicity = EXCLUDED.ethnicity,
                education = EXCLUDED.education,
                work_start_date = EXCLUDED.work_start_date,
                work_years = EXCLUDED.work_years,
                native_place = EXCLUDED.native_place,
                status = 'pending',
                updated_at = NOW()
        """, (
            teacher_id,
            teacher_row[1],  # name
            id_card,
            gender,
            birth_date,
            teacher_row[4],  # ethnicity
            get_education_name(education_row[0], DATABASE_CONFIG) if education_row else None,  # education - 转换为中文
            teacher_row[6],  # work_start_date
            work_years,
            teacher_row[5],  # native_place
            'pending'
        ))
        
        conn.commit()
        
        return {
            "status": "success",
            "message": "退休呈报表数据已汇集到中间表",
            "data": {
                "teacher_id": teacher_id,
                "name": teacher_row[1],
                "gender": gender,
                "birth_date": birth_date,
                "work_years": work_years
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"数据汇集失败: {str(e)}")
    finally:
        cursor.close()
        conn.close()


@router.get("/data/list")
async def get_retirement_data_list():
    """获取退休呈报表数据列表"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            SELECT id, teacher_id, name, gender, birth_date, 
                   work_years, status, created_at
            FROM retirement_report_data
            ORDER BY created_at DESC
        """)
        
        rows = cursor.fetchall()
        result = []
        
        for row in rows:
            result.append({
                "id": row[0],
                "teacher_id": row[1],
                "name": row[2],
                "gender": row[3],
                "birth_date": str(row[4]) if row[4] else None,
                "work_years": row[5],
                "status": row[6],
                "created_at": str(row[7]) if row[7] else None
            })
        
        return {
            "status": "success",
            "data": result,
            "total": len(result)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"查询失败: {str(e)}")
    finally:
        cursor.close()
        conn.close()


@router.get("/data/detail/{teacher_id}")
async def get_retirement_data_detail(teacher_id: int):
    """获取单个教师的退休呈报表详细数据"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            SELECT * FROM retirement_report_data 
            WHERE teacher_id = %s
        """, (teacher_id,))
        
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="数据不存在")
        
        # 获取列名
        columns = [desc[0] for desc in cursor.description]
        
        # 构建字典
        result = {}
        for i, col in enumerate(columns):
            value = row[i]
            if isinstance(value, datetime):
                value = str(value)
            result[col] = value
        
        return {
            "status": "success",
            "data": result
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"查询失败: {str(e)}")
    finally:
        cursor.close()
        conn.close()


@router.put("/data/update/{teacher_id}")
async def update_retirement_data(teacher_id: int, data: Dict[str, Any]):
    """更新退休呈报表数据"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    try:
        # 构建更新SQL - 包含所有可能的字段
        allowed_fields = [
            'name', 'gender', 'id_card', 'birth_date', 'ethnicity', 'education',
            'work_start_date', 'work_years', 'native_place', 'current_address',
            'job_title', 'tech_title', 'position',
            'retirement_reason', 'retirement_address', 'retirement_date', 'retirement_type',
            'post_20140930_specialty', 'post_20140930_specialty_level', 'post_20140930_specialty_salary',
            'post_20140930_worker', 'post_20140930_worker_level', 'post_20140930_worker_salary',
            'last_promotion_date', 'last_position', 'last_position_level', 'last_position_salary',
            'last_worker', 'last_worker_level', 'last_worker_salary',
            'retirement_specialty', 'retirement_specialty_level', 'retirement_specialty_salary',
            'retirement_worker', 'retirement_worker_level', 'retirement_worker_salary',
            'work_history', 'family_support', 'pension_unit',
            'unit_opinion', 'unit_opinion_date', 'dept_opinion', 'dept_opinion_date',
            'approval_opinion', 'approval_date', 'status'
        ]
        
        updates = []
        params = []
        
        for field in allowed_fields:
            if field in data:
                updates.append(f"{field} = %s")
                params.append(data[field])
        
        if not updates:
            raise HTTPException(status_code=400, detail="没有要更新的字段")
        
        params.append(teacher_id)
        
        sql = f"""
            UPDATE retirement_report_data 
            SET {', '.join(updates)}, updated_at = NOW()
            WHERE teacher_id = %s
        """
        
        cursor.execute(sql, params)
        conn.commit()
        
        return {
            "status": "success",
            "message": "数据已更新"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=f"更新失败: {str(e)}")
    finally:
        cursor.close()
        conn.close()


@router.get("/generate-report/{teacher_id}")
async def generate_report(
    teacher_id: int,
    format: str = Query("html", description="报表格式: html, excel, word, pdf")
):
    """生成退休呈报表"""
    try:
        # 获取教师数据
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM retirement_report_data WHERE teacher_id = %s", (teacher_id,))
        row = cursor.fetchone()
        
        if not row:
            raise HTTPException(status_code=404, detail="数据不存在")
        
        # 获取列名和数据
        columns = [desc[0] for desc in cursor.description]
        data = dict(zip(columns, row))
        
        cursor.close()
        conn.close()
        
        # 根据格式生成报表
        if format == 'html':
            return generate_html_report(data)
        elif format == 'excel':
            return generate_excel_report(data)
        elif format == 'word':
            return generate_word_report(data)
        elif format == 'pdf':
            return generate_pdf_report(data)
        else:
            raise HTTPException(status_code=400, detail="不支持的格式")
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成报表失败: {str(e)}")


def generate_html_report(data: dict):
    """生成HTML格式报表"""
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>职工退休呈报表</title>
        <style>
            body {{ font-family: SimSun, serif; font-size: 14px; }}
            table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
            td, th {{ border: 1px solid #000; padding: 10px; text-align: left; }}
            .title {{ text-align: center; font-size: 22px; font-weight: bold; margin: 20px 0; }}
            .header {{ background-color: #f0f0f0; font-weight: bold; }}
        </style>
    </head>
    <body>
        <div class="title">职工退休呈报表</div>
        <table>
            <tr>
                <td class="header">姓名</td>
                <td>{data.get('name', '')}</td>
                <td class="header">性别</td>
                <td>{data.get('gender', '')}</td>
                <td class="header">身份证号</td>
                <td>{data.get('id_card', '')}</td>
            </tr>
            <tr>
                <td class="header">出生日期</td>
                <td>{data.get('birth_date', '')}</td>
                <td class="header">民族</td>
                <td>{data.get('ethnicity', '')}</td>
                <td class="header">文化程度</td>
                <td>{data.get('education', '')}</td>
            </tr>
            <tr>
                <td class="header">参加工作时间</td>
                <td>{data.get('work_start_date', '')}</td>
                <td class="header">工作年限</td>
                <td>{data.get('work_years', '')}</td>
                <td class="header">籍贯</td>
                <td>{data.get('native_place', '')}</td>
            </tr>
            <tr>
                <td class="header">现住址</td>
                <td colspan="5">{data.get('current_address', '')}</td>
            </tr>
            <tr>
                <td class="header">退休后居住地址</td>
                <td colspan="5">{data.get('retirement_address', '')}</td>
            </tr>
            <tr>
                <td class="header">退休原因</td>
                <td colspan="5">{data.get('retirement_reason', '')}</td>
            </tr>
            <tr>
                <td class="header">单位意见</td>
                <td colspan="5">{data.get('unit_opinion', '')}</td>
            </tr>
        </table>
    </body>
    </html>
    """
    
    from fastapi.responses import HTMLResponse
    return HTMLResponse(content=html_content)


def generate_excel_report(data: dict):
    """生成Excel格式报表"""
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side
    from fastapi.responses import StreamingResponse
    import io
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "职工退休呈报表"
    
    # 标题
    ws.merge_cells('A1:F1')
    ws['A1'] = '职工退休呈报表'
    ws['A1'].font = Font(size=16, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    
    # 数据行
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # 第一行
    row_data = [
        ['姓名', data.get('name', ''), '性别', data.get('gender', ''), '身份证号', data.get('id_card', '')],
        ['出生日期', data.get('birth_date', ''), '民族', data.get('ethnicity', ''), '文化程度', data.get('education', '')],
        ['参加工作时间', data.get('work_start_date', ''), '工作年限', data.get('work_years', ''), '籍贯', data.get('native_place', '')],
        ['现住址', data.get('current_address', ''), '', '', '', ''],
        ['退休后居住地址', data.get('retirement_address', ''), '', '', '', ''],
        ['退休原因', data.get('retirement_reason', ''), '', '', '', ''],
        ['单位意见', data.get('unit_opinion', ''), '', '', '', '']
    ]
    
    for i, row in enumerate(row_data, start=3):
        for j, value in enumerate(row, start=1):
            cell = ws.cell(row=i, column=j, value=value)
            cell.border = border
            cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
    
    # 合并单元格
    for row in [6, 7, 8, 9]:
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
    
    # 保存到内存
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="职工退休呈报表_{data.get("name", "" )}.xlsx"'}
    )


def generate_word_report(data: dict):
    """生成Word格式报表"""
    from docx import Document
    from docx.shared import Pt, Inches
    from fastapi.responses import StreamingResponse
    import io
    
    doc = Document()
    
    # 标题
    title = doc.add_heading('职工退休呈报表', 0)
    title.alignment = 1  # 居中
    
    # 添加表格
    table = doc.add_table(rows=7, cols=6)
    table.style = 'Table Grid'
    
    # 填充数据
    cells_data = [
        ['姓名', data.get('name', ''), '性别', data.get('gender', ''), '身份证号', data.get('id_card', '')],
        ['出生日期', data.get('birth_date', ''), '民族', data.get('ethnicity', ''), '文化程度', data.get('education', '')],
        ['参加工作时间', data.get('work_start_date', ''), '工作年限', data.get('work_years', ''), '籍贯', data.get('native_place', '')],
        ['现住址', data.get('current_address', ''), '', '', '', ''],
        ['退休后居住地址', data.get('retirement_address', ''), '', '', '', ''],
        ['退休原因', data.get('retirement_reason', ''), '', '', '', ''],
        ['单位意见', data.get('unit_opinion', ''), '', '', '', '']
    ]
    
    for i, row_data in enumerate(cells_data):
        row = table.rows[i]
        for j, value in enumerate(row_data):
            row.cells[j].text = str(value) if value else ''
    
    # 合并单元格
    for row_idx in [3, 4, 5, 6]:
        for col_idx in range(1, 6):
            table.rows[row_idx].cells[col_idx].merge(table.rows[row_idx].cells[1])
    
    # 保存到内存
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="职工退休呈报表_{data.get("name", "" )}.docx"'}
    )


def generate_pdf_report(data: dict):
    """生成PDF格式报表"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from fastapi.responses import StreamingResponse
    import io
    
    # 注册中文字体
    try:
        pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))
    except:
        pass
    
    # 创建PDF
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4)
    
    # 样式
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName='SimSun',
        fontSize=18,
        alignment=1,  # 居中
        spaceAfter=30
    )
    
    # 内容
    elements = []
    
    # 标题
    elements.append(Paragraph('职工退休呈报表', title_style))
    elements.append(Spacer(1, 20))
    
    # 表格数据
    table_data = [
        ['姓名', data.get('name', ''), '性别', data.get('gender', ''), '身份证号', data.get('id_card', '')],
        ['出生日期', data.get('birth_date', ''), '民族', data.get('ethnicity', ''), '文化程度', data.get('education', '')],
        ['参加工作时间', data.get('work_start_date', ''), '工作年限', data.get('work_years', ''), '籍贯', data.get('native_place', '')],
        ['现住址', data.get('current_address', ''), '', '', '', ''],
        ['退休后居住地址', data.get('retirement_address', ''), '', '', '', ''],
        ['退休原因', data.get('retirement_reason', ''), '', '', '', ''],
        ['单位意见', data.get('unit_opinion', ''), '', '', '', '']
    ]
    
    # 创建表格
    table = Table(table_data, colWidths=[80, 100, 80, 100, 80, 100])
    
    # 表格样式
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, -1), 'SimSun'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('BOX', (0, 0), (-1, -1), 1, colors.black),
        ('SPAN', (1, 3), (5, 3)),  # 现住址合并
        ('SPAN', (1, 4), (5, 4)),  # 退休后居住地址合并
        ('SPAN', (1, 5), (5, 5)),  # 退休原因合并
        ('SPAN', (1, 6), (5, 6)),  # 单位意见合并
    ]))
    
    elements.append(table)
    
    # 生成PDF
    doc.build(elements)
    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type='application/pdf',
        headers={'Content-Disposition': f'attachment; filename="职工退休呈报表_{data.get("name", "" )}.pdf"'}
    )
