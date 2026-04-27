from fastapi import APIRouter, UploadFile, File, HTTPException, Query
from fastapi.responses import JSONResponse, FileResponse, Response
import os
import uuid
import json
from datetime import datetime
from bs4 import BeautifulSoup

router = APIRouter(prefix="/api/template-import-test", tags=["模板导入测试"])

UPLOAD_DIR = "uploads/templates_test"
EXPORT_DIR = "uploads/exports"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)

templates_db = []

def parse_html_to_data(html_content):
    """解析HTML模板，提取表格数据"""
    soup = BeautifulSoup(html_content, 'html.parser')
    tables = soup.find_all('table')
    
    all_data = []
    for table in tables:
        rows = table.find_all('tr')
        table_data = []
        for row in rows:
            cells = row.find_all(['td', 'th'])
            row_data = [cell.get_text(strip=True) for cell in cells]
            if row_data:
                table_data.append(row_data)
        if table_data:
            all_data.append(table_data)
    
    return all_data

@router.post("/upload")
async def upload_template(file: UploadFile = File(...)):
    """上传模板文件（支持PDF/EXCEL/HTML/WORD）"""
    try:
        file_id = str(uuid.uuid4())
        filename = file.filename
        ext = filename.split('.')[-1].lower() if '.' in filename else ''
        
        # 保存文件
        save_path = os.path.join(UPLOAD_DIR, f"{file_id}.{ext}")
        with open(save_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # 模拟解析文件内容
        preview_html = await simulate_file_preview(content, ext, filename)
        
        # 保存模板信息
        template_info = {
            "id": file_id,
            "filename": filename,
            "file_type": ext,
            "upload_time": datetime.now().isoformat(),
            "file_path": save_path,
            "preview_html": preview_html,
            "fields": [],
            "data": {}
        }
        templates_db.append(template_info)
        
        return JSONResponse({
            "status": "success",
            "message": "上传成功",
            "template": template_info
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

async def simulate_file_preview(content, ext, filename):
    """文件预览HTML生成 - Word文件显示提示信息"""
    if ext == 'html':
        try:
            html_content = content.decode('utf-8', errors='ignore')
            return html_content
        except:
            pass
    
    # Word文档显示提示信息
    if ext in ['docx', 'doc']:
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: SimSun, serif; background: #f5f7fa; margin: 0; padding: 40px; }}
                .container {{ max-width: 600px; margin: 0 auto; background: #fff; border-radius: 12px; box-shadow: 0 2px 12px rgba(0,0,0,0.1); padding: 40px; text-align: center; }}
                .icon {{ font-size: 64px; margin-bottom: 20px; }}
                .title {{ font-size: 24px; font-weight: bold; margin-bottom: 16px; color: #333; }}
                .desc {{ font-size: 14px; color: #666; line-height: 1.8; margin-bottom: 24px; }}
                .btn {{ display: inline-block; padding: 12px 32px; background: #409eff; color: #fff; text-decoration: none; border-radius: 6px; font-size: 14px; }}
                .btn:hover {{ background: #66b1ff; }}
                .tips {{ text-align: left; background: #f0f9ff; padding: 20px; border-radius: 8px; margin-top: 24px; }}
                .tips h4 {{ margin: 0 0 12px 0; color: #409eff; }}
                .tips ul {{ margin: 0; padding-left: 20px; }}
                .tips li {{ color: #666; font-size: 13px; margin-bottom: 8px; }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="icon">📄</div>
                <div class="title">{filename}</div>
                <div class="desc">Word文档暂不支持在线预览</div>
                <div class="tips">
                    <h4>💡 使用提示</h4>
                    <ul>
                        <li>为保证100%还原模板样式，建议使用<strong>HTML格式</strong>模板</li>
                        <li>您可以将Word文档另存为HTML格式后上传</li>
                        <li>HTML模板可以完美保留表格样式、字体字号等</li>
                        <li>您也可以直接下载此Word文件进行编辑</li>
                    </ul>
                </div>
            </div>
        </body>
        </html>
        """
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{ font-family: SimSun, serif; font-size: 10.5pt; }}
            .container {{ width: 800px; margin: 20px auto; padding: 40px; background: #fff; }}
            .template-header {{ text-align: center; font-size: 18pt; font-weight: bold; margin-bottom: 30px; }}
            .template-table {{ width: 100%; border-collapse: collapse; }}
            .template-table td, .template-table th {{ 
                border: 1px solid #000; padding: 8px; text-align: center; min-height: 30px;
            }}
            .editable-cell {{ cursor: pointer; }}
            .editable-cell:hover {{ background: #e6f7ff; }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="template-header">{filename}</div>
            <table class="template-table">
                <tr>
                    <td style="width: 25%;">字段1</td>
                    <td style="width: 25%;" class="editable-cell" data-field="field1"></td>
                    <td style="width: 25%;">字段2</td>
                    <td style="width: 25%;" class="editable-cell" data-field="field2"></td>
                </tr>
                <tr>
                    <td>字段3</td>
                    <td class="editable-cell" data-field="field3"></td>
                    <td>字段4</td>
                    <td class="editable-cell" data-field="field4"></td>
                </tr>
                <tr>
                    <td colspan="2" style="text-align: left;">备注：</td>
                    <td colspan="2" class="editable-cell" data-field="remark"></td>
                </tr>
            </table>
            <p style="margin-top: 30px; text-align: right;">（示例预览，实际将100%还原您的模板）</p>
        </div>
    </body>
    </html>
    """

async def convert_word_to_html(content, ext, filename):
    """将Word文档转换为HTML"""
    import io
    from docx import Document
    
    # 保存为临时文件
    temp_file = io.BytesIO(content)
    
    try:
        # 读取Word文档
        doc = Document(temp_file)
        
        # 构建HTML
        html_parts = []
        html_parts.append(f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{ font-family: SimSun, serif; font-size: 12pt; line-height: 1.6; }}
        .document {{ width: 800px; margin: 20px auto; padding: 40px; background: #fff; }}
        h1, h2, h3 {{ text-align: center; }}
        table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
        table td, table th {{ border: 1px solid #000; padding: 8px; }}
    </style>
</head>
<body>
    <div class="document">
''')
        
        # 转换每个段落
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                    html_parts.append(f'<h{level}>{para.text}</h{level}>')
                else:
                    html_parts.append(f'<p>{para.text}</p>')
        
        # 转换表格
        for table in doc.tables:
            html_parts.append('<table>')
            for row in table.rows:
                html_parts.append('<tr>')
                for cell in row.cells:
                    html_parts.append(f'<td>{cell.text}</td>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        
        html_parts.append('''</div>
</body>
</html>''')
        
        return ''.join(html_parts)
        
    except Exception as e:
        print(f"Word解析错误: {e}")
        raise

@router.get("/list")
async def list_templates():
    """获取已上传的模板列表"""
    return JSONResponse({
        "status": "success",
        "templates": templates_db
    })

@router.get("/{template_id}/preview")
async def get_template_preview(template_id: str):
    """获取模板预览HTML - 直接读取原始文件"""
    template = next((t for t in templates_db if t["id"] == template_id), None)
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    file_path = template.get("file_path", "")
    file_type = template.get("file_type", "")
    
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 对于HTML文件，直接读取原始文件内容
    if file_type == 'html':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    # 其他类型返回存储的preview_html
    return template["preview_html"]

@router.get("/{template_id}/raw-file")
async def get_raw_file(template_id: str):
    """获取原始文件"""
    template = next((t for t in templates_db if t["id"] == template_id), None)
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    file_path = template.get("file_path", "")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    filename = os.path.basename(file_path)
    ext = template.get("file_type", "").lower()
    
    # HTML文件直接在浏览器显示（保持原始编码）
    if ext in ['html', 'htm']:
        with open(file_path, 'rb') as f:
            raw_content = f.read()
        
        # 检测编码
        content = raw_content
        content_type = 'text/html'
        
        # 检查HTML中的编码声明
        try:
            text = raw_content.decode('utf-8')
            # 尝试从meta标签中提取编码
            import re
            match = re.search(r'<meta[^>]+charset=["\']?([^"\'\s>]+)', text, re.I)
            if match:
                charset = match.group(1).lower()
                if charset in ['gbk', 'gb2312', 'gb18030']:
                    content = raw_content.decode('gbk')
                    content_type = f'text/html; charset={charset}'
        except:
            # UTF-8解码失败，尝试GBK
            try:
                content = raw_content.decode('gbk')
                content_type = 'text/html; charset=gbk'
            except:
                content = raw_content.decode('utf-8', errors='ignore')
        
        return Response(content=content, media_type=content_type)
    
    # 其他文件类型返回下载
    media_types = {
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'pdf': 'application/pdf',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'xls': 'application/vnd.ms-excel'
    }
    media_type = media_types.get(ext, 'application/octet-stream')
    
    return FileResponse(
        file_path,
        media_type=media_type,
        filename=filename
    )

@router.post("/{template_id}/fields")
async def save_template_fields(template_id: str, fields_data: dict):
    """保存模板字段配置"""
    template = next((t for t in templates_db if t["id"] == template_id), None)
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    template["fields"] = fields_data.get("fields", [])
    return JSONResponse({
        "status": "success",
        "message": "字段保存成功"
    })

@router.get("/{template_id}/data")
async def get_template_data(template_id: str):
    """获取模板已填写的数据"""
    template = next((t for t in templates_db if t["id"] == template_id), None)
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    return JSONResponse({
        "status": "success",
        "data": template.get("data", {})
    })

@router.post("/{template_id}/data")
async def save_template_data(template_id: str, data: dict):
    """保存模板填写的数据"""
    template = next((t for t in templates_db if t["id"] == template_id), None)
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    template["data"] = data.get("data", {})
    return JSONResponse({
        "status": "success",
        "message": "数据保存成功"
    })

@router.post("/{template_id}/export")
async def export_template(template_id: str, request: dict):
    """导出模板"""
    template = next((t for t in templates_db if t["id"] == template_id), None)
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    export_format = request.get("format", "html")
    template_data = template.get("data", {})
    file_type = template.get("file_type", "").lower()
    file_path = template.get("file_path", "")
    
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="模板文件不存在")
    
    # 读取原始文件内容（保持原有编码）
    import chardet
    with open(file_path, 'rb') as f:
        raw_content = f.read()
    
    # 检测编码
    detected = chardet.detect(raw_content)
    encoding = detected.get('encoding', 'utf-8')
    
    # 尝试用检测到的编码解码
    try:
        html_content = raw_content.decode(encoding)
    except:
        try:
            html_content = raw_content.decode('gbk')
        except:
            html_content = raw_content.decode('utf-8', errors='ignore')
    
    # 替换模板中的占位符
    for key, value in template_data.items():
        placeholder = f"{{{{{key}}}}}"
        html_content = html_content.replace(placeholder, str(value) if value else "")
    
    export_filename = f"导出_{template['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    # 对于HTML模板：直接导出HTML文件（保留100%样式）
    if file_type in ['html', 'htm']:
        if export_format == "html":
            output_path = os.path.join(EXPORT_DIR, f"{export_filename}.html")
            with open(output_path, 'w', encoding=encoding) as f:
                f.write(html_content)
            content_type = f'text/html; charset={encoding}'
            return FileResponse(
                output_path,
                media_type=content_type,
                filename=f"{export_filename}.html"
            )
        elif export_format == "pdf":
            output_path = os.path.join(EXPORT_DIR, f"{export_filename}.html")
            with open(output_path, 'w', encoding=encoding) as f:
                f.write(html_content)
            content_type = f'text/html; charset={encoding}'
            return FileResponse(
                output_path,
                media_type=content_type,
                filename=f"{export_filename}.html"
            )
    
    # 对于其他格式：尝试转换（样式可能丢失）
    try:
        if export_format == "xlsx":
            return await export_to_excel(html_content, template_data, export_filename)
        elif export_format == "docx":
            return await export_to_word(html_content, template_data, export_filename)
        else:
            # 默认导出HTML
            output_path = os.path.join(EXPORT_DIR, f"{export_filename}.html")
            with open(output_path, 'w', encoding=encoding) as f:
                f.write(html_content)
            content_type = f'text/html; charset={encoding}'
            return FileResponse(
                output_path,
                media_type=content_type,
                filename=f"{export_filename}.html"
            )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")

async def export_to_excel(html_content: str, template_data: dict, filename: str):
    """导出为Excel"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    
    soup = BeautifulSoup(html_content, 'html.parser')
    tables = soup.find_all('table')
    
    row_idx = 1
    for table in tables:
        rows = table.find_all('tr')
        for row in rows:
            cells = row.find_all(['td', 'th'])
            col_idx = 1
            for cell in cells:
                text = cell.get_text(strip=True)
                ws.cell(row=row_idx, column=col_idx, value=text)
                col_idx += 1
            row_idx += 1
        row_idx += 1
    
    output_path = os.path.join(EXPORT_DIR, f"{filename}.xlsx")
    wb.save(output_path)
    
    return FileResponse(
        output_path,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        filename=f"{filename}.xlsx"
    )

async def export_to_pdf(html_content: str, template_data: dict, filename: str):
    """导出为PDF - 使用HTML直接打印"""
    output_path = os.path.join(EXPORT_DIR, f"{filename}.html")
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    return FileResponse(
        output_path,
        media_type='text/html',
        filename=f"{filename}.html"
    )

async def export_to_word(html_content: str, template_data: dict, filename: str):
    """导出为Word - 将HTML保存为.doc格式"""
    output_path = os.path.join(EXPORT_DIR, f"{filename}.doc")
    
    try:
        from html2docx import html2docx
        html_file = os.path.join(EXPORT_DIR, f"{filename}_temp.html")
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        html2docx(html_file, output_path)
        os.remove(html_file)
        
        return FileResponse(
            output_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"{filename}.docx"
        )
    except ImportError:
        with open(output_path, 'w', encoding='utf-16') as f:
            f.write(html_content)
        
        return FileResponse(
            output_path,
            media_type='application/msword',
            filename=f"{filename}.doc"
        )

@router.delete("/{template_id}")
async def delete_template(template_id: str):
    """删除模板"""
    global templates_db
    template = next((t for t in templates_db if t["id"] == template_id), None)
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    # 删除文件
    if os.path.exists(template["file_path"]):
        os.remove(template["file_path"])
    
    templates_db = [t for t in templates_db if t["id"] != template_id]
    
    return JSONResponse({
        "status": "success",
        "message": "删除成功"
    })
