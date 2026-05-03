from docx import Document
doc = Document(r'D:\erp_thirteen\数据库信息\模板\职工退休呈报表（原版）.docx')

print('=== 段落内容 ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f'{i}: {text}')

print('\n=== 表格1完整结构 ===')
table = doc.tables[1]
print(f'表格尺寸: {len(table.rows)} 行 x {len(table.columns)} 列')
for r_idx, row in enumerate(table.rows):
    cells = []
    for c_idx, cell in enumerate(row.cells):
        text = cell.text.replace('\n', ' ').strip()[:12]
        cells.append(f'[{c_idx}]{text}')
    print(f'行{r_idx:2d}: {" | ".join(cells)}')

print('\n=== 表格0内容 ===')
table = doc.tables[0]
for r_idx, row in enumerate(table.rows):
    cells = []
    for c_idx, cell in enumerate(row.cells):
        text = cell.text.replace('\n', ' ').strip()[:20]
        cells.append(f'[{c_idx}]{text}')
    print(f'行{r_idx:2d}: {" | ".join(cells)}')