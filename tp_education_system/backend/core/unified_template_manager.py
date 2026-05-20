"""
统一模板管理器
实现：导入模板 = 网页呈现模板 = 导出模板
"""
import json
import os
from datetime import datetime
from typing import Dict, List, Any, Optional
import pandas as pd
from docx import Document


class UnifiedTemplateManager:
    """统一模板管理器"""
    
    def __init__(self, config_path: str = None):
        if config_path is None:
            config_path = os.path.join(
                os.path.dirname(__file__), 
                '..', 'config', 
                'unified_template_config.json'
            )
        
        self.config_path = config_path
        self.config = self._load_config()
    
    def _load_config(self) -> Dict:
        """加载配置文件"""
        if os.path.exists(self.config_path):
            with open(self.config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {
            "templates": {},
            "template_categories": {},
            "template_version": "1.0",
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat()
        }
    
    def _save_config(self):
        """保存配置文件"""
        self.config['updated_at'] = datetime.now().isoformat()
        with open(self.config_path, 'w', encoding='utf-8') as f:
            json.dump(self.config, f, ensure_ascii=False, indent=2)
    
    def _detect_header_row(self, df: pd.DataFrame, max_rows: int = 20) -> int:
        """
        智能检测标题行位置
        
        检测逻辑：
        1. 标题行通常包含文本而非数字
        2. 标题行的文本较短（通常是字段名）
        3. 标题行下方是数据行（包含各种类型的数据）
        """
        for row_idx in range(min(max_rows, len(df))):
            row = df.iloc[row_idx]
            
            text_count = 0
            numeric_count = 0
            empty_count = 0
            short_text_count = 0
            
            for val in row:
                if pd.isna(val) or str(val).strip() == '':
                    empty_count += 1
                elif isinstance(val, (int, float)) and not pd.isna(val):
                    numeric_count += 1
                else:
                    text_val = str(val).strip()
                    text_count += 1
                    if len(text_val) <= 20:
                        short_text_count += 1
            
            total_cells = len(row)
            non_empty = total_cells - empty_count
            
            if non_empty == 0:
                continue
            
            text_ratio = text_count / non_empty if non_empty > 0 else 0
            short_ratio = short_text_count / non_empty if non_empty > 0 else 0
            
            if row_idx < len(df) - 1:
                next_row = df.iloc[row_idx + 1]
                next_has_data = any(not pd.isna(v) and str(v).strip() != '' for v in next_row)
            else:
                next_has_data = False
            
            if text_ratio > 0.7 and short_ratio > 0.5 and next_has_data:
                return row_idx
        
        return 0
    
    def _detect_data_type(self, series: pd.Series) -> Dict:
        """
        智能检测列数据类型
        """
        non_null = series.dropna()
        
        if len(non_null) == 0:
            return {"data_type": "VARCHAR", "length": 255, "input_type": "text"}
        
        sample_values = non_null.head(20).tolist()
        
        date_count = 0
        numeric_count = 0
        text_count = 0
        max_length = 0
        
        for val in sample_values:
            str_val = str(val).strip()
            
            try:
                pd.to_datetime(val)
                date_count += 1
                continue
            except:
                pass
            
            if isinstance(val, (int, float)) and not pd.isna(val):
                numeric_count += 1
            else:
                text_count += 1
                max_length = max(max_length, len(str_val))
        
        total = len(sample_values)
        
        if date_count / total > 0.7:
            return {"data_type": "DATE", "length": None, "input_type": "date"}
        elif numeric_count / total > 0.7:
            sample_num = [v for v in sample_values if isinstance(v, (int, float))]
            has_decimal = any(isinstance(v, float) and v != int(v) for v in sample_num)
            if has_decimal:
                return {"data_type": "DECIMAL", "length": "10,2", "input_type": "number"}
            else:
                return {"data_type": "INTEGER", "length": None, "input_type": "number"}
        else:
            return {"data_type": "VARCHAR", "length": max(max_length * 2, 50), "input_type": "text"}
    
    def _analyze_template_structure(self, file_path: str) -> Dict:
        """
        智能分析Excel模板结构
        """
        df = pd.read_excel(file_path, header=None)
        
        header_row = self._detect_header_row(df)
        
        df_with_header = pd.read_excel(file_path, header=header_row)
        
        preview_rows = min(10, len(df))
        preview_data = []
        for i in range(preview_rows):
            row_data = []
            for j, val in enumerate(df.iloc[i]):
                row_data.append(str(val) if not pd.isna(val) else "")
            preview_data.append(row_data)
        
        structure = {
            "total_rows": int(len(df)),
            "total_columns": int(len(df.columns)),
            "header_row": int(header_row + 1),
            "data_start_row": int(header_row + 2),
            "data_end_row": int(len(df)),
            "sheet_name": "Sheet1",
            "preview_data": preview_data,
            "columns": []
        }
        
        for col_name in df_with_header.columns:
            col_data = df_with_header[col_name]
            type_info = self._detect_data_type(col_data)
            
            non_null = col_data.dropna()
            sample_values = non_null.head(5).tolist() if len(non_null) > 0 else []
            
            structure["columns"].append({
                "chinese_name": str(col_name),
                "data_type": type_info["data_type"],
                "length": type_info["length"],
                "input_type": type_info["input_type"],
                "sample_values": [str(v) for v in sample_values],
                "null_count": int(col_data.isna().sum()),
                "total_count": int(len(col_data))
            })
        
        return structure
    
    def parse_excel_template(self, file_path: str, template_name: str = None) -> Dict:
        """
        智能解析Excel模板文件
        
        自动识别：
        1. 标题行位置
        2. 数据类型
        3. 数据区域
        """
        if template_name is None:
            template_name = os.path.splitext(os.path.basename(file_path))[0]
        
        structure = self._analyze_template_structure(file_path)
        
        fields = []
        english_names_used = set()
        
        for idx, col_info in enumerate(structure["columns"]):
            chinese_name = col_info["chinese_name"]
            english_name = self._generate_unique_english_name(
                chinese_name, 
                english_names_used
            )
            english_names_used.add(english_name)
            
            field_config = {
                "chinese_name": chinese_name,
                "english_name": english_name,
                "column_index": idx,
                "data_type": col_info["data_type"],
                "length": col_info["length"],
                "input_type": col_info["input_type"],
                "required": col_info["null_count"] == 0,
                "display_order": idx,
                "sample_values": col_info["sample_values"]
            }
            fields.append(field_config)
        
        template_config = {
            "template_id": self._generate_template_id(template_name),
            "chinese_name": template_name,
            "english_name": self._generate_english_name(template_name),
            "version": "1.0",
            "created_at": datetime.now().isoformat(),
            "source_file": os.path.basename(file_path),
            
            "structure": structure,
            "fields": fields,
            
            "import_config": {
                "file_type": "xlsx",
                "sheet_name": structure["sheet_name"],
                "header_row": structure["header_row"],
                "start_row": structure["data_start_row"],
                "end_row": structure["data_end_row"],
                "fields": fields
            },
            
            "display_config": {
                "form_layout": "horizontal",
                "columns_per_row": 3,
                "fields": [
                    {
                        "chinese_name": field["chinese_name"],
                        "english_name": field["english_name"],
                        "display_name": field["chinese_name"],
                        "input_type": field["input_type"],
                        "placeholder": f"请输入{field['chinese_name']}",
                        "required": field["required"]
                    }
                    for field in fields
                ]
            },
            
            "export_config": {
                "file_type": "xlsx",
                "template_file": file_path,
                "sheet_name": structure["sheet_name"],
                "start_row": structure["data_start_row"],
                "fields": [
                    {
                        "chinese_name": field["chinese_name"],
                        "english_name": field["english_name"],
                        "column_index": field["column_index"]
                    }
                    for field in fields
                ]
            }
        }
        
        return template_config
    
    def parse_word_template(self, file_path: str, template_name: str = None) -> Dict:
        """
        解析Word模板文件
        
        Args:
            file_path: Word文件路径
            template_name: 模板名称（可选）
        
        Returns:
            模板配置字典
        """
        if template_name is None:
            template_name = os.path.splitext(os.path.basename(file_path))[0]
        
        doc = Document(file_path)
        placeholders = set()
        
        for paragraph in doc.paragraphs:
            import re
            matches = re.findall(r'\{([^{}]+)\}', paragraph.text)
            placeholders.update(matches)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    import re
                    matches = re.findall(r'\{([^{}]+)\}', cell.text)
                    placeholders.update(matches)
        
        fields = []
        for idx, placeholder in enumerate(sorted(placeholders)):
            field_config = {
                "chinese_name": placeholder,
                "english_name": self._generate_english_name(placeholder),
                "placeholder": f"{{{placeholder}}}",
                "data_type": "VARCHAR",
                "length": 255,
                "required": False,
                "display_order": idx
            }
            fields.append(field_config)
        
        template_config = {
            "template_id": self._generate_template_id(template_name),
            "chinese_name": template_name,
            "english_name": self._generate_english_name(template_name),
            "version": "1.0",
            "created_at": datetime.now().isoformat(),
            "source_file": os.path.basename(file_path),
            
            "fields": fields,
            
            "import_config": {
                "file_type": "docx",
                "fields": fields
            },
            
            "display_config": {
                "form_layout": "vertical",
                "fields": [
                    {
                        "chinese_name": field["chinese_name"],
                        "english_name": field["english_name"],
                        "display_name": field["chinese_name"],
                        "input_type": "text",
                        "placeholder": f"请输入{field['chinese_name']}",
                        "required": field["required"]
                    }
                    for field in fields
                ]
            },
            
            "export_config": {
                "file_type": "docx",
                "template_file": file_path,
                "fields": [
                    {
                        "chinese_name": field["chinese_name"],
                        "english_name": field["english_name"],
                        "placeholder": field["placeholder"]
                    }
                    for field in fields
                ]
            }
        }
        
        return template_config
    
    def save_template(self, template_config: Dict) -> bool:
        """
        保存模板配置
        
        Args:
            template_config: 模板配置字典
        
        Returns:
            是否保存成功
        """
        template_id = template_config["template_id"]
        self.config["templates"][template_id] = template_config
        self._save_config()
        return True
    
    def get_template(self, template_id: str) -> Optional[Dict]:
        """
        获取模板配置
        
        Args:
            template_id: 模板ID
        
        Returns:
            模板配置字典
        """
        return self.config["templates"].get(template_id)
    
    def list_templates(self, category: str = None) -> List[Dict]:
        """
        列出所有模板
        
        Args:
            category: 模板分类（可选）
        
        Returns:
            模板列表
        """
        templates = list(self.config["templates"].values())
        
        if category:
            templates = [
                t for t in templates 
                if t.get("category") == category
            ]
        
        return templates
    
    def delete_template(self, template_id: str) -> bool:
        """
        删除模板
        
        Args:
            template_id: 模板ID
        
        Returns:
            是否删除成功
        """
        if template_id in self.config["templates"]:
            del self.config["templates"][template_id]
            self._save_config()
            return True
        return False
    
    def _generate_template_id(self, template_name: str) -> str:
        """生成模板ID"""
        import hashlib
        timestamp = datetime.now().isoformat()
        unique_str = f"{template_name}_{timestamp}"
        return hashlib.md5(unique_str.encode()).hexdigest()[:12]
    
    def _generate_english_name(self, chinese_name: str) -> str:
        """
        生成英文名称（使用FieldNameManager）
        
        Args:
            chinese_name: 中文名称
        
        Returns:
            英文名称
        """
        from core.field_name_manager import FieldNameManager
        
        field_name_manager = FieldNameManager()
        english_name = field_name_manager.get_english_name(chinese_name)
        
        return english_name
    
    def _generate_unique_english_name(self, chinese_name: str, used_names: set) -> str:
        """
        生成唯一的英文名称
        
        Args:
            chinese_name: 中文名称
            used_names: 已使用的英文名称集合
        
        Returns:
            唯一的英文名称
        """
        base_english_name = self._generate_english_name(chinese_name)
        
        if base_english_name not in used_names:
            return base_english_name
        
        counter = 1
        while f"{base_english_name}_{counter}" in used_names:
            counter += 1
        
        return f"{base_english_name}_{counter}"
