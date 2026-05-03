from docx import Document
doc = Document(r'D:\erp_thirteen\数据库信息\模板\职工退休呈报表.docx')

print('='*80)
print('模板分析报告')
print('='*80)

# 提取所有占位符
placeholders = set()

def extract_placeholders(text):
    import re
    matches = re.findall(r'\{\{(\w+)\}\}', text)
    return matches

print('\n一、段落中的占位符：')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        phs = extract_placeholders(text)
        if phs:
            print(f'  段落{i}: {text[:60]}...')
            for ph in phs:
                print(f'    → 占位符: {{{{ {ph} }}}}')
                placeholders.add(ph)

print('\n二、表格中的占位符：')
for t_idx, table in enumerate(doc.tables):
    print(f'\n  表格{t_idx} ({len(table.rows)}行 x {len(table.columns)}列):')
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            phs = extract_placeholders(text)
            if phs:
                for ph in phs:
                    print(f'    行{r_idx}列{c_idx}: {{{{{ph}}}}}')
                    placeholders.add(ph)

print('\n' + '='*80)
print('三、占位符清单（共{}个）:'.format(len(placeholders)))
print('='*80)
for ph in sorted(placeholders):
    print(f'  - {ph}')