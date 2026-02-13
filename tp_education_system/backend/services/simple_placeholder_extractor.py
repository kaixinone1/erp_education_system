"""
简单健壮的占位符提取服务
支持多种文件格式
"""
import re
import os
from typing import List, Dict, Any


def clean_placeholder_name(name: str) -> str:
    """清理占位符名称，去除HTML标签"""
    # 去除HTML标签
    name = re.sub(r'<[^>]+>', '', name)
    # 去除HTML实体
    name = re.sub(r'&[^;]+;', '', name)
    # 去除多余空白
    name = ' '.join(name.split())
    return name.strip()


def is_valid_field_name(name: str) -> bool:
    """验证字段名是否有效"""
    name = name.strip()
    if not name:
        return False
    # 排除纯数字
    if name.isdigit():
        return False
    # 排除HTML条件注释相关内容
    html_comment_keywords = ['if gte', 'if lte', 'endif', 'if !support', 'if !vml']
    if any(kw in name.lower() for kw in html_comment_keywords):
        return False
    # 排除CSS代码
    css_keywords = ['mso-', 'style', 'font-', 'border', 'padding', 'margin', 
                    'color:', 'background', 'text-align', 'white-space', 'lang=',
                    'font-family', 'font-size', 'kerning', 'mso-', 'page:', 'margin-']
    if any(kw in name.lower() for kw in css_keywords):
        return False
    # 排除包含特殊字符的
    if any(c in name for c in [':', ';', '{', '}', '<', '>', '/', '\\', '!', '=']):
        return False
    # 排除太长的
    if len(name) > 50:
        return False
    # 至少包含一个中文字符或字母
    if not re.search(r'[\u4e00-\u9fa5a-zA-Z]', name):
        return False
    return True


def extract_from_text(content: str) -> List[str]:
    """从文本中提取所有可能的占位符"""
    placeholders = []
    
    # 模式1: {字段名} - 支持HTML标签包裹
    pattern1 = r'\{([^{}]+)\}'
    matches1 = re.findall(pattern1, content)
    
    # 模式2: {{字段名}}
    pattern2 = r'\{\{([^{}]+)\}\}'
    matches2 = re.findall(pattern2, content)
    
    # 模式3: [字段名]
    pattern3 = r'\[([^\[\]]+)\]'
    matches3 = re.findall(pattern3, content)
    
    # 合并所有匹配
    all_matches = matches1 + matches2 + matches3
    
    # 去重并保持顺序
    seen = set()
    for match in all_matches:
        # 清理HTML标签
        cleaned = clean_placeholder_name(match)
        if is_valid_field_name(cleaned) and cleaned not in seen:
            seen.add(cleaned)
            placeholders.append(cleaned)
    
    return placeholders


def read_file_with_encoding(file_path: str) -> str:
    """尝试多种编码读取文件"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin-1']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    # 如果都失败，使用latin-1（不会报错但可能显示乱码）
    with open(file_path, 'r', encoding='latin-1') as f:
        return f.read()


def extract_from_html(file_path: str) -> List[Dict[str, Any]]:
    """从HTML文件中提取占位符"""
    fields = []
    
    try:
        content = read_file_with_encoding(file_path)
        
        # 直接提取所有占位符
        placeholders = extract_from_text(content)
        
        for ph in placeholders:
            fields.append({
                'name': ph,
                'label': ph,
                'field_type': 'text',
                'source': 'html'
            })
        
    except Exception as e:
        print(f"提取HTML字段失败: {e}")
    
    return fields


def extract_from_word(file_path: str) -> List[Dict[str, Any]]:
    """从Word文件中提取占位符"""
    fields = []
    
    try:
        from docx import Document
        doc = Document(file_path)
        
        # 提取所有文本
        all_text = []
        for para in doc.paragraphs:
            all_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        
        content = '\n'.join(all_text)
        placeholders = extract_from_text(content)
        
        for ph in placeholders:
            fields.append({
                'name': ph,
                'label': ph,
                'field_type': 'text',
                'source': 'word'
            })
        
    except Exception as e:
        print(f"提取Word字段失败: {e}")
    
    return fields


def extract_from_excel(file_path: str) -> List[Dict[str, Any]]:
    """从Excel文件中提取占位符"""
    fields = []
    
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        
        all_text = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value:
                        all_text.append(str(cell.value))
        
        content = '\n'.join(all_text)
        placeholders = extract_from_text(content)
        
        for ph in placeholders:
            fields.append({
                'name': ph,
                'label': ph,
                'field_type': 'text',
                'source': 'excel'
            })
        
    except Exception as e:
        print(f"提取Excel字段失败: {e}")
    
    return fields


def extract_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    """从PDF文件中提取占位符"""
    fields = []
    
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            all_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    all_text.append(text)
        
        content = '\n'.join(all_text)
        placeholders = extract_from_text(content)
        
        for ph in placeholders:
            fields.append({
                'name': ph,
                'label': ph,
                'field_type': 'text',
                'source': 'pdf'
            })
        
    except Exception as e:
        print(f"提取PDF字段失败: {e}")
    
    return fields


def extract_fields(file_path: str, file_ext: str) -> List[Dict[str, Any]]:
    """通用字段提取入口"""
    ext = file_ext.lower()
    
    if ext in ['.html', '.htm']:
        return extract_from_html(file_path)
    elif ext == '.docx':
        return extract_from_word(file_path)
    elif ext in ['.xlsx', '.xls']:
        return extract_from_excel(file_path)
    elif ext == '.pdf':
        return extract_from_pdf(file_path)
    else:
        # 默认按文本处理
        return extract_from_html(file_path)


# 向后兼容
extract_html_fields = extract_from_html
extract_word_fields = extract_from_word
extract_excel_fields = extract_from_excel
extract_pdf_fields = extract_from_pdf
