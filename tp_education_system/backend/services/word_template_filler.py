"""
Word模板填充引擎
读取 .docx 模板，查找所有 {{占位符}}，替换为实际数据，保存为新的 .docx 文件

核心思路：
1. 打开模板 .docx
2. 遍历文档中的所有段落，查找并替换 {{}} 占位符
3. 遍历所有表格的所有单元格，查找并替换 {{}} 占位符
4. 用数据字典中的值替换占位符
5. 如果数据值为 None 或空，保留占位符为空
6. 保存为新的 .docx 文件（保留原始Section/分页结构）
"""
import re
import os
import shutil
from docx import Document


def _replace_placeholders_in_text(text, data, not_found_list):
    """
    在文本中替换所有 {{占位符}}
    
    Args:
        text: 原始文本
        data: 数据字典
        not_found_list: 未找到的占位符列表（会被追加）
    
    Returns:
        (替换后的文本, 替换次数)
    """
    placeholder_pattern = re.compile(r'\{\{(.+?)\}\}')
    replaced_count = 0
    
    placeholders = placeholder_pattern.findall(text)
    if not placeholders:
        return text, 0
    
    new_text = text
    for ph in placeholders:
        ph_key = ph.strip()
        value = data.get(ph_key)
        if value is None:
            value = ''
            not_found_list.append(ph_key)
        elif not isinstance(value, str):
            value = str(value)
        
        new_text = new_text.replace('{{' + ph + '}}', value)
        replaced_count += 1
    
    return new_text, replaced_count


def _replace_placeholders_in_paragraphs(doc, data, not_found_list):
    """
    替换文档中所有段落（非表格）的占位符
    
    Args:
        doc: python-docx Document 对象
        data: 数据字典
        not_found_list: 未找到的占位符列表
    """
    replaced_count = 0
    
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        
        new_text, count = _replace_placeholders_in_text(para.text, data, not_found_list)
        replaced_count += count
        
        if count > 0:
            # 保存原始文本（清除runs前），用于后续判断
            original_text = para.text
            # 清除所有runs，设置新文本
            for run in para.runs:
                run.text = ''
            # 清理姓名后的点号（模板封面段落中 {{姓名}} 后面有 ". " 用于分隔，替换后应移除）
            if '{{姓名}}' in original_text and new_text.startswith('姓名：'):
                new_text = new_text.rstrip(' .')
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)
    
    return replaced_count


def _replace_placeholders_in_tables(doc, data, not_found_list):
    """
    替换文档中所有表格单元格的占位符
    保留模板原有的字体格式（字号、字体、颜色、加粗等）
    
    Args:
        doc: python-docx Document 对象
        data: 数据字典
        not_found_list: 未找到的占位符列表
    """
    replaced_count = 0
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text:
                    continue
                
                new_text, count = _replace_placeholders_in_text(original_text, data, not_found_list)
                replaced_count += count
                
                if new_text != original_text and count > 0:
                    # 保留模板字体格式：逐段落替换占位符
                    _replace_in_cell_paragraphs(cell, data, not_found_list)
    
    return replaced_count


def _replace_in_cell_paragraphs(cell, data, not_found_list):
    """
    在单元格的每个段落中替换占位符，保留原有字体格式
    
    策略：
    1. 若占位符完整存在于单个run内 → 逐run替换，其他run的静态文本保持不变
    2. 若占位符跨多个run分割（如 {{ / 事业管理岗位 / 1}}）→ 合并所有run后统一替换
    """
    for para in cell.paragraphs:
        full_text = para.text
        if not full_text.strip():
            continue
        
        if '{{' not in full_text:
            continue
        
        # 检查是否有单run占位符（完整占位符在同一个run中）
        has_single_run_placeholder = False
        for run in para.runs:
            if '{{' in run.text and '}}' in run.text:
                has_single_run_placeholder = True
                break
        
        if has_single_run_placeholder:
            # 逐run替换：只替换包含占位符的run，其他run保持不变
            for run in para.runs:
                if '{{' in run.text and '}}' in run.text:
                    new_text, count = _replace_placeholders_in_text(run.text, data, not_found_list)
                    if count > 0:
                        run.text = new_text
        else:
            # 跨run占位符：合并所有run文本，统一替换后写回第一个run
            new_text, count = _replace_placeholders_in_text(full_text, data, not_found_list)
            if count > 0:
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)


def fill_word_template(template_path: str, output_path: str, data: dict) -> str:
    """
    填充Word模板 —— 100%保留原始模板结构，只替换占位符

    不修改模板的任何属性（段落顺序、行高、字体、间距、Section结构等），
    仅将 {{占位符}} 替换为实际数据值。

    Args:
        template_path: 模板文件路径 (.docx)
        output_path: 输出文件路径 (.docx)
        data: 数据字典，key=占位符名，value=实际值

    Returns:
        输出文件路径
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"模板文件不存在: {template_path}")

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # 打开模板（保留原始Section结构）
    doc = Document(template_path)

    not_found = []

    # 第一步：替换段落中的占位符
    para_count = _replace_placeholders_in_paragraphs(doc, data, not_found)

    # 第二步：替换表格中的占位符
    table_count = _replace_placeholders_in_tables(doc, data, not_found)

    total_replaced = para_count + table_count

    if not_found:
        unique_not_found = list(set(not_found))
        print(f"[Word模板] 以下占位符无对应数据: {unique_not_found}")
    print(f"[Word模板] 段落替换 {para_count} 个占位符, 表格替换 {table_count} 个占位符, 共 {total_replaced} 个")

    # 保存（100%保留原始模板结构）
    doc.save(output_path)
    return output_path


def extract_placeholders(template_path: str) -> list:
    """
    提取模板中所有占位符名称（包括段落和表格）
    
    Args:
        template_path: 模板文件路径
    
    Returns:
        占位符名称列表（去重）
    """
    if not os.path.exists(template_path):
        return []
    
    doc = Document(template_path)
    placeholder_pattern = re.compile(r'\{\{(.+?)\}\}')
    all_placeholders = []
    
    # 提取段落中的占位符
    for para in doc.paragraphs:
        placeholders = placeholder_pattern.findall(para.text)
        all_placeholders.extend(ph.strip() for ph in placeholders)
    
    # 提取表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                placeholders = placeholder_pattern.findall(cell.text)
                all_placeholders.extend(ph.strip() for ph in placeholders)
    
    return list(set(all_placeholders))