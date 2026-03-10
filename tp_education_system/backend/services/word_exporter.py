"""
Word导出服务
读取原始Word模板，填充数据，生成完整的Word文档
"""
import os
import tempfile
from typing import Dict, Optional
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class WordExporter:
    """Word导出器"""
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def export_report(self, template_path: str, data: Dict, teacher_id: str) -> Optional[str]:
        """
        导出报表为Word
        读取原始Word模板，填充数据
        
        Args:
            template_path: Word模板文件路径
            data: 填充数据
            teacher_id: 教师ID
            
        Returns:
            Word文件路径
        """
        try:
            # 读取原始Word文件
            doc = Document(template_path)
            
            # 填充表格1（审批意见区域）
            self._fill_table1(doc.tables[0], data)
            
            # 填充表格2（详细信息）
            self._fill_table2(doc.tables[1], data)
            
            # 保存文件 - 使用模板文件名生成输出文件名
            template_name = os.path.basename(template_path)
            output_filename = f"{template_name}_{teacher_id}_已填充.docx"
            output_path = os.path.join(self.temp_dir, output_filename)
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            print(f"导出Word失败: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _fill_table1(self, table, data: Dict):
        """填充表格1 - 审批意见区域"""
        # 第1行：呈报单位意见
        name = data.get('姓名', '')
        table.rows[0].cells[1].text = f"经研究，同意 {name} 同志按以下第（    ）条办理"
        
        # 第2行：主管部门审查意见
        table.rows[1].cells[1].text = "\n同意呈报\n              "
        
        # 第3行：退休一次性补贴审批意见
        subsidy_text = f"根据鄂人社发【2017】8号文件规定，同意 {name} 同志发放独生子女费        元，教育特殊贡献奖        元，"
        table.rows[2].cells[1].text = subsidy_text
    
    def _fill_table2(self, table, data: Dict):
        """填充表格2 - 详细信息（28行 x 15列）"""
        
        # 第1行：姓名、性别、出生年月
        self._set_cell_text(table, 0, 1, data.get('姓名', ''))
        self._set_cell_text(table, 0, 6, data.get('性别', ''))
        self._set_cell_text(table, 0, 12, data.get('出生年月', data.get('出生日期', '')))
        
        # 第2行：民族、文化程度、是否独生子女
        self._set_cell_text(table, 1, 1, data.get('民族', ''))
        self._set_cell_text(table, 1, 6, data.get('文化程度', ''))
        self._set_cell_text(table, 1, 12, data.get('是否独生子女', ''))
        
        # 第3行：入党年月、职务、技术职称
        self._set_cell_text(table, 2, 1, data.get('入党年月', ''))
        self._set_cell_text(table, 2, 6, data.get('职务', ''))
        self._set_cell_text(table, 2, 12, data.get('技术职称', ''))
        
        # 第4行：参加工作时间、工作年限
        self._set_cell_text(table, 3, 3, data.get('参加工作时间', ''))
        self._set_cell_text(table, 3, 9, str(data.get('工作年限', '')))
        
        # 第5行：籍贯、现在住址
        self._set_cell_text(table, 4, 1, data.get('籍贯', ''))
        self._set_cell_text(table, 4, 8, data.get('现在住址', data.get('现住址', '')))
        
        # 第6行：工作简历（标题行，不填充）
        
        # 第7行：工作简历表头（不填充）
        
        # 第8-13行：工作简历内容（6行）
        work_resume = data.get('工作简历', [])
        if isinstance(work_resume, list):
            for i, resume in enumerate(work_resume[:6]):
                row_idx = 7 + i
                if row_idx < 13:
                    self._set_cell_text(table, row_idx, 0, resume.get('自何年何月', ''))
                    self._set_cell_text(table, row_idx, 2, resume.get('至何年何月', ''))
                    self._set_cell_text(table, row_idx, 5, resume.get('在何单位任何职', ''))
                    self._set_cell_text(table, row_idx, 11, resume.get('证明人及其住址', ''))
        
        # 第14行：退休原因
        self._set_cell_text(table, 13, 1, data.get('退休原因', ''))
        
        # 第15行：供养直系亲属
        self._set_cell_text(table, 14, 1, data.get('供养直系亲属', data.get('直系亲属供养情况', '')))
        
        # 第16行：退休后居住地址
        self._set_cell_text(table, 15, 1, data.get('退休后居住地址', ''))
        
        # 第17-20行：2014年9月30日工资信息
        # 机关工人
        self._set_cell_text(table, 16, 4, data.get('2014机关工人技术等级', ''))
        self._set_cell_text(table, 16, 7, data.get('2014机关工人级别薪级', ''))
        
        # 事业管理
        self._set_cell_text(table, 17, 4, data.get('2014事业管理岗位', ''))
        self._set_cell_text(table, 17, 7, data.get('2014事业管理对应原职务', ''))
        self._set_cell_text(table, 17, 13, data.get('2014事业管理薪级', ''))
        
        # 事业专技
        self._set_cell_text(table, 18, 4, data.get('2014事业专技岗位', ''))
        self._set_cell_text(table, 18, 7, data.get('2014事业专技对应原职务', ''))
        self._set_cell_text(table, 18, 13, data.get('2014事业专技薪级', ''))
        
        # 事业工勤
        self._set_cell_text(table, 19, 4, data.get('2014事业工勤岗位', ''))
        self._set_cell_text(table, 19, 7, data.get('2014事业工勤对应技术等级', ''))
        self._set_cell_text(table, 19, 13, data.get('2014事业工勤薪级', ''))
        
        # 第21-24行：最后一次职务升降时
        # 机关工人
        self._set_cell_text(table, 20, 4, data.get('最后机关工人技术等级', ''))
        self._set_cell_text(table, 20, 7, data.get('最后机关工人级别薪级', ''))
        
        # 事业管理
        self._set_cell_text(table, 21, 4, data.get('最后事业管理岗位', ''))
        self._set_cell_text(table, 21, 7, data.get('最后事业管理对应原职务', ''))
        self._set_cell_text(table, 21, 13, data.get('最后事业管理薪级', ''))
        
        # 事业专技
        self._set_cell_text(table, 22, 4, data.get('最后事业专技岗位', ''))
        self._set_cell_text(table, 22, 7, data.get('最后事业专技对应原职务', ''))
        self._set_cell_text(table, 22, 13, data.get('最后事业专技薪级', ''))
        
        # 事业工勤
        self._set_cell_text(table, 23, 4, data.get('最后事业工勤岗位', ''))
        self._set_cell_text(table, 23, 7, data.get('最后事业工勤对应技术等级', ''))
        self._set_cell_text(table, 23, 13, data.get('最后事业工勤薪级', ''))
        
        # 第25-28行：退休时
        # 机关工人
        self._set_cell_text(table, 24, 4, data.get('退休机关工人技术等级', ''))
        self._set_cell_text(table, 24, 7, data.get('退休机关工人级别薪级', ''))
        
        # 事业管理
        self._set_cell_text(table, 25, 4, data.get('退休事业管理岗位', ''))
        self._set_cell_text(table, 25, 7, data.get('退休事业管理对应原职务', ''))
        self._set_cell_text(table, 25, 13, data.get('退休事业管理薪级', ''))
        
        # 事业专技
        self._set_cell_text(table, 26, 4, data.get('退休事业专技岗位', ''))
        self._set_cell_text(table, 26, 7, data.get('退休事业专技对应原职务', ''))
        self._set_cell_text(table, 26, 13, data.get('退休事业专技薪级', ''))
        
        # 事业工勤
        self._set_cell_text(table, 27, 4, data.get('退休事业工勤岗位', ''))
        self._set_cell_text(table, 27, 7, data.get('退休事业工勤对应技术等级', ''))
        self._set_cell_text(table, 27, 13, data.get('退休事业工勤薪级', ''))
    
    def _set_cell_text(self, table, row_idx, col_idx, text):
        """设置单元格文本"""
        try:
            if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(text) if text is not None else ''
        except Exception as e:
            print(f"设置单元格[{row_idx},{col_idx}]失败: {e}")
