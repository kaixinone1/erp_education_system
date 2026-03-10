"""
模板导出器
将数据填充到各种格式的模板中
"""
import os
import re
import shutil
from typing import Dict, List, Optional
from docx import Document


class WordExporter:
    """Word文档导出器"""
    
    @staticmethod
    def export(template_path: str, output_path: str, data: Dict) -> str:
        """
        导出Word文档
        
        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            data: 填充数据
            
        Returns:
            输出文件路径
        """
        # 复制模板文件
        shutil.copy(template_path, output_path)
        
        # 打开复制的文件
        doc = Document(output_path)
        
        # 替换段落中的占位符
        for para in doc.paragraphs:
            for key, value in data.items():
                placeholder = f'{{{{{key}}}}}'
                if placeholder in para.text:
                    # 判断是否有有效值
                    if value is not None and str(value).strip() != '':
                        # 有值，进行替换
                        para.clear()
                        new_text = para.text.replace(placeholder, str(value))
                        para.add_run(new_text)
                    # 无值时保留原占位符，不做任何替换
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = f'{{{{{key}}}}}'
                        if placeholder in cell.text:
                            # 判断是否有有效值
                            if value is not None and str(value).strip() != '':
                                # 有值，进行替换
                                new_text = cell.text.replace(placeholder, str(value))
                                cell.text = new_text
                            # 无值时保留原占位符，不做任何替换
        
        # 保存文件
        doc.save(output_path)
        
        return output_path


class ExcelExporter:
    """Excel文档导出器"""
    
    @staticmethod
    def export(template_path: str, output_path: str, data: Dict) -> str:
        """
        导出Excel文档
        
        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            data: 填充数据
            
        Returns:
            输出文件路径
        """
        from openpyxl import load_workbook
        
        # 复制模板文件
        shutil.copy(template_path, output_path)
        
        # 打开工作簿
        wb = load_workbook(output_path)
        ws = wb.active
        
        # 遍历所有单元格，替换占位符
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    for key, value in data.items():
                        placeholder = f'{{{{{key}}}}}'
                        if placeholder in cell.value:
                            # 判断是否有有效值
                            if value is not None and str(value).strip() != '':
                                # 有值，进行替换
                                cell.value = cell.value.replace(placeholder, str(value))
                            # 无值时保留原占位符，不做任何替换
        
        # 保存文件
        wb.save(output_path)
        
        return output_path


class PDFExporter:
    """PDF文档导出器"""
    
    @staticmethod
    def export(template_path: str, output_path: str, data: Dict) -> str:
        """
        导出PDF文档
        
        Args:
            template_path: 模板文件路径（PDF或HTML）
            output_path: 输出文件路径
            data: 填充数据
            
        Returns:
            输出文件路径
        """
        # 对于PDF，我们需要重新生成
        # 这里使用HTML作为中间格式，然后转换为PDF
        from playwright.sync_api import sync_playwright
        
        # 读取模板文件（假设是HTML格式）
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        html_content = None
        for encoding in encodings:
            try:
                with open(template_path, 'r', encoding=encoding) as f:
                    html_content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if html_content is None:
            raise Exception("无法读取模板文件")
        
        # 替换占位符
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            # 判断是否有有效值
            if value is not None and str(value).strip() != '':
                # 有值，进行替换
                html_content = html_content.replace(placeholder, str(value))
            # 无值时保留原占位符，不做任何替换
        
        # 使用Playwright生成PDF
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            
            # 加载HTML内容
            page.set_content(html_content)
            
            # 等待页面加载完成
            page.wait_for_load_state('networkidle')
            
            # 生成PDF
            page.pdf(
                path=output_path,
                format='A4',
                margin={
                    'top': '10mm',
                    'right': '10mm',
                    'bottom': '10mm',
                    'left': '10mm'
                },
                print_background=True
            )
            
            browser.close()
        
        return output_path


class HTMLExporter:
    """HTML文档导出器"""
    
    @staticmethod
    def export(template_path: str, output_path: str, data: Dict) -> str:
        """
        导出HTML文档
        
        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            data: 填充数据
            
        Returns:
            输出文件路径
        """
        # 读取模板文件
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        content = None
        for encoding in encodings:
            try:
                with open(template_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise Exception("无法读取模板文件")
        
        # 替换占位符
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            content = content.replace(placeholder, str(value))
        
        # 写入输出文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return output_path


class TemplateExporter:
    """模板导出器工厂类"""
    
    @staticmethod
    def export(template_path: str, output_path: str, data: Dict) -> str:
        """
        根据文件格式自动选择导出器
        
        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            data: 填充数据
            
        Returns:
            输出文件路径
        """
        ext = os.path.splitext(template_path)[1].lower()
        
        if ext == '.docx':
            return WordExporter.export(template_path, output_path, data)
        elif ext == '.xlsx':
            return ExcelExporter.export(template_path, output_path, data)
        elif ext == '.pdf':
            return PDFExporter.export(template_path, output_path, data)
        elif ext in ['.html', '.htm']:
            return HTMLExporter.export(template_path, output_path, data)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
