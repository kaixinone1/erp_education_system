"""
报表生成服务
支持从模板生成Word文档，配置页面设置
"""
import json
import os
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class ReportService:
    """报表生成服务"""
    
    def __init__(self):
        self.config_path = os.path.join(
            os.path.dirname(os.path.dirname(__file__)), 
            'config', 
            'report_definitions.json'
        )
        self.reports_config = self._load_config()
    
    def _load_config(self) -> Dict:
        """加载报表配置"""
        if os.path.exists(self.config_path):
            with open(self.config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {"reports": []}
    
    def get_report_config(self, report_id: str) -> Optional[Dict]:
        """获取报表配置"""
        for report in self.reports_config.get("reports", []):
            if report["id"] == report_id:
                return report
        return None
    
    def generate_report(self, report_id: str, data: Dict[str, Any], output_path: str) -> str:
        """
        生成报表
        
        Args:
            report_id: 报表ID
            data: 填充数据
            output_path: 输出文件路径
            
        Returns:
            生成的文件路径
        """
        config = self.get_report_config(report_id)
        if not config:
            raise ValueError(f"未找到报表配置: {report_id}")
        
        template_file = config.get("template_file")
        page_settings = config.get("page_settings", {})
        data_mapping = config.get("data_mapping", {})
        
        # 如果有模板文件，使用模板
        template_path = os.path.join(
            os.path.dirname(os.path.dirname(__file__)),
            template_file
        ) if template_file else None
        
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
        else:
            # 创建新文档
            doc = Document()
        
        # 设置页面
        self._setup_page(doc, page_settings)
        
        # 填充数据
        self._fill_data(doc, data, data_mapping)
        
        # 保存文档
        doc.save(output_path)
        
        return output_path
    
    def _setup_page(self, doc: Document, page_settings: Dict):
        """设置页面格式"""
        section = doc.sections[0]
        
        # 纸张大小
        paper_size = page_settings.get("paper_size", "A4")
        if paper_size == "A4":
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
        elif paper_size == "A3":
            section.page_width = Inches(11.69)
            section.page_height = Inches(16.54)
        
        # 纸张方向
        orientation = page_settings.get("orientation", "portrait")
        if orientation == "landscape":
            section.orientation = 1  # WD_ORIENT.LANDSCAPE
            # 交换宽高
            section.page_width, section.page_height = section.page_height, section.page_width
        
        # 页边距（转换为英寸）
        section.top_margin = Inches(page_settings.get("margin_top", 2.54) / 2.54)
        section.bottom_margin = Inches(page_settings.get("margin_bottom", 2.54) / 2.54)
        section.left_margin = Inches(page_settings.get("margin_left", 3.17) / 2.54)
        section.right_margin = Inches(page_settings.get("margin_right", 3.17) / 2.54)
    
    def _fill_data(self, doc: Document, data: Dict, data_mapping: Dict):
        """填充数据到文档"""
        # 遍历所有段落
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                text = run.text
                # 替换占位符 {{field_name}}
                for cn_field, en_field in data_mapping.items():
                    placeholder = f"{{{{{en_field}}}}}"
                    if placeholder in text:
                        value = data.get(en_field, "")
                        text = text.replace(placeholder, str(value))
                run.text = text
        
        # 遍历所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            text = run.text
                            for cn_field, en_field in data_mapping.items():
                                placeholder = f"{{{{{en_field}}}}}"
                                if placeholder in text:
                                    value = data.get(en_field, "")
                                    text = text.replace(placeholder, str(value))
                            run.text = text
    
    def create_template(self, report_id: str, output_path: str):
        """
        创建报表模板文件
        
        Args:
            report_id: 报表ID
            output_path: 模板输出路径
        """
        config = self.get_report_config(report_id)
        if not config:
            raise ValueError(f"未找到报表配置: {report_id}")
        
        page_settings = config.get("page_settings", {})
        data_mapping = config.get("data_mapping", {})
        
        # 创建新文档
        doc = Document()
        
        # 设置页面
        self._setup_page(doc, page_settings)
        
        # 添加标题
        title = doc.add_heading(config.get("name", "报表"), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加说明
        doc.add_paragraph(f"说明: {config.get('description', '')}")
        doc.add_paragraph()
        
        # 添加数据字段说明
        doc.add_heading("数据字段", level=2)
        for cn_field, en_field in data_mapping.items():
            p = doc.add_paragraph()
            p.add_run(f"{cn_field}: ").bold = True
            p.add_run(f"{{{{{en_field}}}}}")
        
        # 保存模板
        doc.save(output_path)
        
        return output_path


# 全局实例
report_service = ReportService()
