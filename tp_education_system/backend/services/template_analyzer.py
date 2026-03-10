"""
模板分析器
分析各种格式的模板文件，提取页面设置和占位符
"""
import re
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field


@dataclass
class PageSettings:
    """页面设置"""
    paper_size: str = 'A4'           # A3/A4/A5/Letter
    orientation: str = 'portrait'    # portrait/landscape
    width_cm: float = 21.0
    height_cm: float = 29.7
    margin_left_cm: float = 2.5
    margin_right_cm: float = 2.5
    margin_top_cm: float = 2.5
    margin_bottom_cm: float = 2.5


@dataclass
class Placeholder:
    """占位符"""
    name: str                        # 占位符名称，如：{{姓名}}
    format_type: str                 # word/excel/pdf/html
    # Word/Excel位置
    table_index: Optional[int] = None
    row_index: Optional[int] = None
    cell_index: Optional[int] = None
    # PDF位置
    page_num: Optional[int] = None
    x_pos: Optional[float] = None
    y_pos: Optional[float] = None
    # HTML位置
    css_selector: Optional[str] = None


@dataclass
class TemplateInfo:
    """模板信息"""
    template_id: str
    file_format: str                 # docx/xlsx/pdf/html
    source_file: str
    page_settings: PageSettings = field(default_factory=PageSettings)
    placeholders: List[Placeholder] = field(default_factory=list)


class WordAnalyzer:
    """Word文档分析器"""
    
    @staticmethod
    def analyze(file_path: str, template_id: str) -> TemplateInfo:
        """
        分析Word文档
        
        Args:
            file_path: Word文件路径
            template_id: 模板ID
            
        Returns:
            TemplateInfo对象
        """
        from docx import Document
        
        doc = Document(file_path)
        
        # 提取页面设置
        section = doc.sections[0]
        width_cm = section.page_width.cm
        height_cm = section.page_height.cm
        
        # 判断纸张大小和方向
        is_landscape = width_cm > height_cm
        if max(width_cm, height_cm) >= 40:
            paper_size = 'A3'
        else:
            paper_size = 'A4'
        
        page_settings = PageSettings(
            paper_size=paper_size,
            orientation='landscape' if is_landscape else 'portrait',
            width_cm=width_cm,
            height_cm=height_cm,
            margin_left_cm=section.left_margin.cm,
            margin_right_cm=section.right_margin.cm,
            margin_top_cm=section.top_margin.cm,
            margin_bottom_cm=section.bottom_margin.cm
        )
        
        # 提取占位符
        placeholders = []
        
        # 遍历段落
        for para in doc.paragraphs:
            found = re.findall(r'\{\{([^}]+)\}\}', para.text)
            for ph in found:
                placeholders.append(Placeholder(
                    name=ph,
                    format_type='word'
                ))
        
        # 遍历表格
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    found = re.findall(r'\{\{([^}]+)\}\}', cell.text)
                    for ph in found:
                        placeholders.append(Placeholder(
                            name=ph,
                            format_type='word',
                            table_index=table_idx,
                            row_index=row_idx,
                            cell_index=cell_idx
                        ))
        
        return TemplateInfo(
            template_id=template_id,
            file_format='docx',
            source_file=file_path,
            page_settings=page_settings,
            placeholders=placeholders
        )


class ExcelAnalyzer:
    """Excel文档分析器"""
    
    @staticmethod
    def analyze(file_path: str, template_id: str) -> TemplateInfo:
        """
        分析Excel文档
        
        Args:
            file_path: Excel文件路径
            template_id: 模板ID
            
        Returns:
            TemplateInfo对象
        """
        from openpyxl import load_workbook
        
        wb = load_workbook(file_path)
        ws = wb.active
        
        # 提取页面设置
        page_setup = ws.page_setup
        
        # 判断纸张大小和方向
        paper_size = 'A4'
        orientation = 'portrait'
        
        if page_setup.paperSize:
            if page_setup.paperSize == 8:  # A3
                paper_size = 'A3'
            elif page_setup.paperSize == 9:  # A4
                paper_size = 'A4'
        
        if page_setup.orientation:
            if page_setup.orientation == 'landscape':
                orientation = 'landscape'
        
        # 默认尺寸
        if paper_size == 'A3':
            width_cm = 42.01 if orientation == 'landscape' else 29.70
            height_cm = 29.70 if orientation == 'landscape' else 42.01
        else:  # A4
            width_cm = 29.70 if orientation == 'landscape' else 21.0
            height_cm = 21.0 if orientation == 'landscape' else 29.70
        
        page_settings = PageSettings(
            paper_size=paper_size,
            orientation=orientation,
            width_cm=width_cm,
            height_cm=height_cm
        )
        
        # 提取占位符
        placeholders = []
        
        for row_idx, row in enumerate(ws.iter_rows(), start=1):
            for col_idx, cell in enumerate(row, start=1):
                if cell.value and isinstance(cell.value, str):
                    found = re.findall(r'\{\{([^}]+)\}\}', cell.value)
                    for ph in found:
                        placeholders.append(Placeholder(
                            name=ph,
                            format_type='excel',
                            row_index=row_idx - 1,  # 0-based
                            cell_index=col_idx - 1
                        ))
        
        return TemplateInfo(
            template_id=template_id,
            file_format='xlsx',
            source_file=file_path,
            page_settings=page_settings,
            placeholders=placeholders
        )


class PDFAnalyzer:
    """PDF文档分析器"""
    
    @staticmethod
    def analyze(file_path: str, template_id: str) -> TemplateInfo:
        """
        分析PDF文档
        
        Args:
            file_path: PDF文件路径
            template_id: 模板ID
            
        Returns:
            TemplateInfo对象
        """
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        page = reader.pages[0]
        
        # 提取页面设置
        width_pt = float(page.mediabox.width)
        height_pt = float(page.mediabox.height)
        
        # 转换为厘米 (1 pt = 0.0352778 cm)
        width_cm = width_pt * 0.0352778
        height_cm = height_pt * 0.0352778
        
        # 判断纸张大小和方向
        is_landscape = width_cm > height_cm
        if max(width_cm, height_cm) >= 40:
            paper_size = 'A3'
        else:
            paper_size = 'A4'
        
        page_settings = PageSettings(
            paper_size=paper_size,
            orientation='landscape' if is_landscape else 'portrait',
            width_cm=width_cm,
            height_cm=height_cm
        )
        
        # 提取占位符（从文本中提取）
        placeholders = []
        
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text:
                found = re.findall(r'\{\{([^}]+)\}\}', text)
                for ph in found:
                    placeholders.append(Placeholder(
                        name=ph,
                        format_type='pdf',
                        page_num=page_num
                    ))
        
        return TemplateInfo(
            template_id=template_id,
            file_format='pdf',
            source_file=file_path,
            page_settings=page_settings,
            placeholders=placeholders
        )


class HTMLAnalyzer:
    """HTML文档分析器"""
    
    @staticmethod
    def analyze(file_path: str, template_id: str) -> TemplateInfo:
        """
        分析HTML文档
        
        Args:
            file_path: HTML文件路径
            template_id: 模板ID
            
        Returns:
            TemplateInfo对象
        """
        from bs4 import BeautifulSoup
        
        # 读取HTML文件
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        content = None
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise Exception("无法读取HTML文件")
        
        soup = BeautifulSoup(content, 'html.parser')
        
        # 提取页面设置（从CSS @page规则）
        paper_size = 'A4'
        orientation = 'portrait'
        width_cm = 21.0
        height_cm = 29.7
        
        # 查找CSS中的@page规则
        style_tags = soup.find_all('style')
        for style in style_tags:
            if style.string:
                # 查找size属性
                size_match = re.search(r'size:\s*(\w+)', style.string, re.IGNORECASE)
                if size_match:
                    size_val = size_match.group(1).upper()
                    if size_val in ['A3', 'A4', 'A5']:
                        paper_size = size_val
                
                # 查找orientation
                if 'landscape' in style.string.lower():
                    orientation = 'landscape'
        
        # 根据纸张大小设置尺寸
        if paper_size == 'A3':
            width_cm = 42.01 if orientation == 'landscape' else 29.70
            height_cm = 29.70 if orientation == 'landscape' else 42.01
        elif paper_size == 'A4':
            width_cm = 29.70 if orientation == 'landscape' else 21.0
            height_cm = 21.0 if orientation == 'landscape' else 29.70
        
        page_settings = PageSettings(
            paper_size=paper_size,
            orientation=orientation,
            width_cm=width_cm,
            height_cm=height_cm
        )
        
        # 提取占位符
        placeholders = []
        
        # 从整个HTML内容中提取占位符
        found = re.findall(r'\{\{([^}]+)\}\}', content)
        for ph in found:
            placeholders.append(Placeholder(
                name=ph,
                format_type='html'
            ))
        
        return TemplateInfo(
            template_id=template_id,
            file_format='html',
            source_file=file_path,
            page_settings=page_settings,
            placeholders=placeholders
        )


class TemplateAnalyzer:
    """模板分析器工厂类"""
    
    @staticmethod
    def analyze(file_path: str, template_id: str) -> TemplateInfo:
        """
        根据文件格式自动选择分析器
        
        Args:
            file_path: 文件路径
            template_id: 模板ID
            
        Returns:
            TemplateInfo对象
        """
        import os
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.docx':
            return WordAnalyzer.analyze(file_path, template_id)
        elif ext == '.xlsx':
            return ExcelAnalyzer.analyze(file_path, template_id)
        elif ext == '.pdf':
            return PDFAnalyzer.analyze(file_path, template_id)
        elif ext in ['.html', '.htm']:
            return HTMLAnalyzer.analyze(file_path, template_id)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
