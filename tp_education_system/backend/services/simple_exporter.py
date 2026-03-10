"""
简单导出器 - 只保留核心功能
复制模板文件，替换占位符，返回填充后的文件
"""
import os
import shutil
import re
from docx import Document


def replace_text_in_paragraph(paragraph, placeholder, replacement):
    """
    在段落中替换文本，保留格式
    找到包含占位符的run，只替换那个run的文本
    """
    if placeholder not in paragraph.text:
        return
    
    # 遍历所有runs，找到包含占位符的run
    for run in paragraph.runs:
        if placeholder in run.text:
            # 只替换这个run的文本，保留格式
            run.text = run.text.replace(placeholder, str(replacement))
            return
    
    # 如果占位符跨越多个runs，需要合并处理
    # 获取完整文本
    full_text = paragraph.text
    if placeholder in full_text:
        new_text = full_text.replace(placeholder, str(replacement))
        # 保存第一个run的格式
        if paragraph.runs:
            first_run = paragraph.runs[0]
            # 清除所有runs
            for run in paragraph.runs:
                run.text = ""
            # 在第一个run中设置新文本
            first_run.text = new_text


def replace_text_in_cell(cell, placeholder, replacement):
    """
    在单元格中替换文本，保留格式
    """
    if placeholder not in cell.text:
        return
    
    # 遍历单元格中的所有段落
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, placeholder, replacement)


def export_word(template_path: str, output_path: str, data: dict) -> str:
    """
    导出Word文件
    复制模板，替换{{字段名}}为实际数据，保留格式
    """
    # 复制模板文件
    shutil.copy(template_path, output_path)
    
    # 打开文件
    doc = Document(output_path)
    
    # 替换所有{{字段名}} - 在段落中
    for para in doc.paragraphs:
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            replace_text_in_paragraph(para, placeholder, str(value))
    
    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f'{{{{{key}}}}}'
                    replace_text_in_cell(cell, placeholder, str(value))
    
    # 保存
    doc.save(output_path)
    return output_path


def export_excel(template_path: str, output_path: str, data: dict) -> str:
    """
    导出Excel文件
    复制模板，替换{{字段名}}为实际数据
    """
    from openpyxl import load_workbook
    
    # 复制模板文件
    shutil.copy(template_path, output_path)
    
    # 打开工作簿
    wb = load_workbook(output_path)
    ws = wb.active
    
    # 替换所有单元格中的{{字段名}}
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                for key, value in data.items():
                    placeholder = f'{{{{{key}}}}}'
                    if placeholder in cell.value:
                        cell.value = cell.value.replace(placeholder, str(value))
    
    # 保存
    wb.save(output_path)
    return output_path


def export_html(template_path: str, output_path: str, data: dict) -> str:
    """
    导出HTML文件
    复制模板，替换{{字段名}}为实际数据
    """
    # 读取模板
    with open(template_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    
    # 替换{{字段名}}
    for key, value in data.items():
        placeholder = f'{{{{{key}}}}}'
        content = content.replace(placeholder, str(value))
    
    # 写入输出文件
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return output_path


def export_pdf(template_path: str, output_path: str, data: dict) -> str:
    """
    导出PDF文件
    如果是HTML模板，先填充再转PDF；如果是PDF模板，暂不支持填充
    """
    from playwright.sync_api import sync_playwright
    
    # 读取模板（假设是HTML）
    with open(template_path, 'r', encoding='utf-8', errors='ignore') as f:
        html_content = f.read()
    
    # 替换{{字段名}}
    for key, value in data.items():
        placeholder = f'{{{{{key}}}}}'
        html_content = html_content.replace(placeholder, str(value))
    
    # 使用Playwright生成PDF
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(html_content)
        page.wait_for_load_state('networkidle')
        page.pdf(path=output_path, format='A4', print_background=True)
        browser.close()
    
    return output_path
