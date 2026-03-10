"""
最终的Word导出器 - 最简单直接的方式
"""
import shutil
from docx import Document


def export_word(template_path: str, output_path: str, data: dict) -> str:
    """
    导出Word文件
    1. 复制模板
    2. 替换{{字段名}}
    3. 保存
    """
    # 复制模板
    shutil.copy(template_path, output_path)
    
    # 打开文档
    doc = Document(output_path)
    
    # 替换所有{{字段名}}
    for para in doc.paragraphs:
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(value))
    
    # 替换表格中的{{字段名}}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f'{{{{{key}}}}}'
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))
    
    # 保存
    doc.save(output_path)
    return output_path
