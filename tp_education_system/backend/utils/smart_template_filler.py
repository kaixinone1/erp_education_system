"""
智能模板填充器 - 自动识别表格结构并填充数据
"""
from docx import Document
from docx.shared import Pt
from typing import Dict, Any, List, Tuple
import re
import os

# 字段映射配置 - 中文表头到英文字段的映射
FIELD_MAPPINGS = {
    # 基本信息
    "姓名": "teacher_name",
    "教师姓名": "teacher_name",
    "性别": "gender",
    "出生年月": "birth_date",
    "出生日期": "birth_date",
    "身份证号": "id_card",
    "身份证号码": "id_card",
    "民族": "ethnicity",
    "籍贯": "native_place",
    "联系电话": "contact_phone",
    "电话": "contact_phone",

    # 教育信息
    "文化程度": "education",
    "学历": "education",
    "毕业学校": "graduation_school",
    "学校": "graduation_school",
    "专业": "major",
    "学位": "degree",

    # 工作信息
    "参加工作时间": "work_start_date",
    "工作时间": "work_start_date",
    "工作单位": "work_unit",
    "单位": "work_unit",
    "职务": "position",
    "职称": "title",
    "技术职称": "title",
    "工作年限": "work_years",

    # 退休信息
    "退休日期": "retirement_date",
    "年龄": "age",

    # 其他
    "是否独生子女": "is_only_child",
    "入党年月": "join_party_date",
    "现在住址": "current_address",
    "退休原因": "retirement_reason",
}


def normalize_text(text: str) -> str:
    """标准化文本 - 去除空格和换行"""
    if not text:
        return ""
    return re.sub(r'\s+', '', text.strip())


def find_field_mapping(chinese_header: str) -> str:
    """
    根据中文表头查找对应的英文字段
    支持模糊匹配
    """
    normalized = normalize_text(chinese_header)

    # 精确匹配
    if normalized in FIELD_MAPPINGS:
        return FIELD_MAPPINGS[normalized]

    # 模糊匹配 - 检查是否包含关键字
    for cn, en in FIELD_MAPPINGS.items():
        if cn in normalized or normalized in cn:
            return en

    return None


def analyze_table_structure(table) -> List[Tuple[int, int, str]]:
    """
    分析表格结构，返回单元格位置及其对应的字段名
    返回: [(行号, 列号, 字段名), ...]
    """
    field_positions = []

    # 遍历表格的所有单元格
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            text = normalize_text(cell.text)

            # 跳过空单元格
            if not text:
                continue

            # 尝试匹配字段
            field_name = find_field_mapping(text)
            if field_name:
                field_positions.append((row_idx, col_idx, field_name))

    return field_positions


def find_data_cells(table, field_positions: List[Tuple[int, int, str]]) -> Dict[str, Tuple[int, int]]:
    """
    找到需要填充数据的单元格位置
    通常数据在表头右侧或下方
    """
    data_cells = {}
    rows = len(table.rows)
    cols = len(table.rows[0].cells) if rows > 0 else 0

    for row_idx, col_idx, field_name in field_positions:
        # 策略1: 检查右侧单元格
        if col_idx + 1 < cols:
            right_cell = table.rows[row_idx].cells[col_idx + 1]
            if not normalize_text(right_cell.text) or normalize_text(right_cell.text) != normalize_text(table.rows[row_idx].cells[col_idx].text):
                data_cells[field_name] = (row_idx, col_idx + 1)
                continue

        # 策略2: 检查下方单元格
        if row_idx + 1 < rows:
            below_cell = table.rows[row_idx + 1].cells[col_idx]
            if not normalize_text(below_cell.text) or normalize_text(below_cell.text) != normalize_text(table.rows[row_idx].cells[col_idx].text):
                data_cells[field_name] = (row_idx + 1, col_idx)
                continue

    return data_cells


def fill_table_smart(table, data: Dict[str, Any]) -> int:
    """
    智能填充表格
    返回填充的单元格数量
    """
    filled_count = 0

    # 分析表格结构
    field_positions = analyze_table_structure(table)

    if not field_positions:
        # 如果没有识别到表头，尝试直接查找包含字段名的单元格
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = normalize_text(cell.text)
                # 检查是否是字段名（在FIELD_MAPPINGS中）
                if text in FIELD_MAPPINGS:
                    field_name = FIELD_MAPPINGS[text]
                    # 尝试填充右侧或下方单元格
                    if col_idx + 1 < len(row.cells):
                        target_cell = row.cells[col_idx + 1]
                        if field_name in data and data[field_name]:
                            target_cell.text = str(data[field_name])
                            filled_count += 1

        return filled_count

    # 找到数据单元格位置
    data_cells = find_data_cells(table, field_positions)

    # 填充数据
    for field_name, (row_idx, col_idx) in data_cells.items():
        if field_name in data and data[field_name]:
            try:
                cell = table.rows[row_idx].cells[col_idx]
                # 只填充空单元格或占位符单元格
                current_text = normalize_text(cell.text)
                if not current_text or current_text in ['', '.', ' ']:
                    cell.text = str(data[field_name])
                    filled_count += 1
            except IndexError:
                continue

    return filled_count


def fill_paragraphs_smart(doc, data: Dict[str, Any]) -> int:
    """
    智能填充段落中的占位符
    识别类似 "姓名：____" 或 "姓名：" 后面有空格的模式
    """
    filled_count = 0

    # 段落填充模式
    patterns = [
        (r'(姓名|教师姓名)[：:\s]*([\s\.]+|$)', 'teacher_name'),
        (r'(性别)[：:\s]*([\s\.]+|$)', 'gender'),
        (r'(出生年月|出生日期)[：:\s]*([\s\.]+|$)', 'birth_date'),
        (r'(身份证号|身份证号码)[：:\s]*([\s\.]+|$)', 'id_card'),
        (r'(单位)[：:\s]*([\s\.]+|$)', 'work_unit'),
    ]

    for para in doc.paragraphs:
        text = para.text
        for pattern, field_name in patterns:
            if re.search(pattern, text) and field_name in data and data[field_name]:
                # 替换文本
                new_text = re.sub(
                    pattern,
                    lambda m: f"{m.group(1)}：{data[field_name]}",
                    text
                )
                if new_text != text:
                    para.text = new_text
                    filled_count += 1
                break

    return filled_count


def smart_fill_template(template_path: str, data: Dict[str, Any], output_path: str = None) -> str:
    """
    智能填充模板

    Args:
        template_path: 模板文件路径
        data: 填充数据
        output_path: 输出文件路径（可选）

    Returns:
        填充后的文件路径
    """
    # 加载模板
    doc = Document(template_path)

    total_filled = 0

    # 1. 填充表格
    for table in doc.tables:
        filled = fill_table_smart(table, data)
        total_filled += filled
        print(f"表格填充了 {filled} 个单元格")

    # 2. 填充段落
    filled = fill_paragraphs_smart(doc, data)
    total_filled += filled
    print(f"段落填充了 {filled} 处")

    # 生成输出路径
    if not output_path:
        base_name = os.path.basename(template_path)
        name, ext = os.path.splitext(base_name)
        output_path = os.path.join(
            os.path.dirname(template_path),
            f"filled_{name}{ext}"
        )

    # 保存文档
    doc.save(output_path)
    print(f"\n智能填充完成！")
    print(f"共填充 {total_filled} 处")
    print(f"输出文件: {output_path}")

    return output_path


# 测试函数
if __name__ == "__main__":
    # 从数据库获取真实数据，不使用硬编码测试数据
    print("请从数据库获取真实数据进行测试")
