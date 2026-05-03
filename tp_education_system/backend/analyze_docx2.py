from docx import Document
doc = Document(r'D:\erp_thirteen\数据库信息\模板\职工退休呈报表.doc')

print('=== 段落内容 ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f'{i}: {text}')

print('\n=== 表格内容 ===')
for t_idx, table in enumerate(doc.tables):
    print(f'\n表格 {t_idx}: {len(table.rows)} 行 x {len(table.columns)} 列')
    for r_idx, row in enumerate(table.rows[:10]):
        cells = []
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.replace('\n', ' ').strip()
            if text:
                cells.append(f'[{c_idx}]{text[:20]}')
        if cells:
            print(f'  行{r_idx}: {" | ".join(cells)}')