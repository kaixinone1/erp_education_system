from docx import Document
doc = Document(r'D:\erp_thirteen\数据库信息\模板\职工退休呈报表.docx')

print('=== 段落内容 ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f'{i}: {text}')

print('\n=== 表格0内容 ===')
table = doc.tables[0]
print(f'尺寸: {len(table.rows)} 行 x {len(table.columns)} 列')
for r_idx, row in enumerate(table.rows):
    cells = []
    for c_idx, cell in enumerate(row.cells):
        text = cell.text.replace('\n', ' ').strip()
        if text:
            cells.append(f'[{c_idx}]{text[:25]}')
    if cells:
        print(f'  行{r_idx}: {" | ".join(cells)}')

print('\n=== 表格1内容 ===')
table = doc.tables[1]
print(f'尺寸: {len(table.rows)} 行 x {len(table.columns)} 列')
for r_idx, row in enumerate(table.rows):
    cells = []
    for c_idx, cell in enumerate(row.cells):
        text = cell.text.replace('\n', ' ').strip()
        if text:
            cells.append(f'[{c_idx}]{text[:20]}')
    if cells:
        print(f'  行{r_idx}: {" | ".join(cells)}')