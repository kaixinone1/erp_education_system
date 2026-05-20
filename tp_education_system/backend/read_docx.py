from docx import Document

doc_path = r'D:\erp_thirteen\系统说明文件类\备注信息统计汇总方式.docx'
doc = Document(doc_path)

print('=' * 60)
print('备注信息统计汇总方式文档内容：')
print('=' * 60)

for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

print('\n' + '=' * 60)
print('表格内容：')
print('=' * 60)

for i, table in enumerate(doc.tables):
    print(f'\n表格 {i+1}：')
    for row in table.rows:
        row_text = [cell.text.strip() for cell in row.cells]
        print(' | '.join(row_text))
